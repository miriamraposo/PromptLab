# api.py 


# --- 1. Imports de la librería estándar de Python ---
import os
import io
import logging
import json
import re
import time
import uuid
import httpx
import threading
from functools import wraps
from io import StringIO, BytesIO
from urllib.parse import urlparse, unquote
from typing import Optional, Dict
from collections import deque
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeout
import base64        
import binascii 
import tempfile
import shutil
from pathlib import Path
from io import BytesIO
from PIL import Image as PillowImage 
import werkzeug
from threading import Lock
from datetime import datetime
from PIL import Image
from apscheduler.schedulers.background import BackgroundScheduler
import atexit
from uuid import uuid4
import mimetypes
import csv
# --- 2. Imports de librerías de terceros (las que instalaste con pip) ---
import numpy as np
import joblib
import requests
import pandas as pd
import fitz
from dotenv import load_dotenv
from flask import Flask, request, jsonify, g, send_file
from flask import Response
import torch
from flask_cors import CORS
from supabase import create_client, Client
from gotrue.errors import AuthApiError
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename
from werkzeug.exceptions import BadRequest, InternalServerError
from docx import Document

#from rq import Queue
#from rq.job import Job
#from redis import Redis

# Inicializar dotenv
load_dotenv()
from utils.translators import translate_es_en 

from services.project_service import ProjectService
from services.data_service import DataService
from services.backend_nlp import TextAnalysisService
from services.asistente_de_limpieza_orquestador import  get_preliminary_analysis, get_column_details, duplicar_columna_wrapper, cleaning_action,check_dataset_quality
from services import asistente_de_limpieza_orquestador as orquestador
from services.load_data import robust_read_csv, apply_cleaning_action, drop_rows
from utils.supabase_handler import supabase_handler
from services.backend_nlp import TextAnalysisService
from services.visualize import generate_plot 
from services.backend_predictivo import PredictionService
from services.model_manager import ModelManager
from services.prompt_lab_service import PromptLabService
from services.history_service import HistoryService
from services.faq_service import FAQService
from services.TabularClusteringService import TabularClusteringService
from vision_processor import analisis_completo_de_imagen
from vision_processor import orquestar_edicion_avanzada, extraer_datos_estructurados_con_gemini, procesar_imagen_completa ,generar_imagen_desde_texto, crear_meme,  extraer_color_dominante, analizar_contenido_imagen_google,download_image_from_url
import json 
from utils.security import encrypt_text, decrypt_text 
from services.vision_backend import VisionPredictionService


# --- 1. CONFIGURACIÓN INICIAL ---


logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] - %(message)s')
logger = logging.getLogger(__name__) 

# --- 2. CARGA DE VARIABLES DE ENTORNO ---

SUPABASE_URL = os.environ.get("SUPABASE_URL")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY")
SUPABASE_SERVICE_KEY = os.environ.get("SUPABASE_SERVICE_KEY")
SECRET_KEY = os.environ.get("FLASK_SECRET_KEY", "desarrollo-super-secreto-cambiar-en-prod")
FRONTEND_URL = os.environ.get("FRONTEND_URL", "http://localhost:5173")
N8N_WEBHOOK_URL = os.getenv("N8N_WEBHOOK_URL", "https://tu-n8n.com/webhook/export")
print("ENCRYPTION_KEY cargada:", os.environ.get("ENCRYPTION_KEY") is not None)
print("---------------------------------")

if not all([SUPABASE_URL, SUPABASE_KEY, SUPABASE_SERVICE_KEY]):
    # Hacemos que el error sea más obvio
    raise ValueError("Error CRÍTICO: Una o más variables de entorno de Supabase no se cargaron. Revisa tu archivo .env y la ruta de carga.")



def clean_for_json(obj):
    """
    Convierte NaN, NaT y tipos incompatibles con JSON a valores válidos.
    """
    if isinstance(obj, pd.DataFrame):
        return obj.where(pd.notnull(obj), None).to_dict(orient="records")
    elif isinstance(obj, pd.Series):
        return obj.where(pd.notnull(obj), None).to_list()
    elif isinstance(obj, (np.floating, np.integer, np.bool_)):
        return obj.item()
    elif isinstance(obj, (list, tuple)):
        return [clean_for_json(x) for x in obj]
    elif isinstance(obj, dict):
        return {k: clean_for_json(v) for k, v in obj.items()}
    elif pd.isna(obj):
        return None
    else:
        return obj
    

# --- 3. INICIALIZACIÓN DE APP Y SERVICIOS ---
# (Sin cambios, todo perfecto)
app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024 
app.config["SECRET_KEY"] = SECRET_KEY


origins = [
        os.environ.get("FRONTEND_URL", "http://localhost:5173"),
        "http://127.0.0.1:5173"
    ]

    # Añadimos la URL de ngrok si está definida en el .env
ngrok_url = os.environ.get("NGROK_URL")
if ngrok_url:
    origins.append(ngrok_url)
    print(f"CORS: Permitiendo acceso desde la URL de ngrok: {ngrok_url}")


# Si tienes una URL de producción, la añadirías aquí también.

# 2. Aplicamos la configuración de CORS con la lista de orígenes
CORS(
    app,
    resources={r"/api/*": {"origins": origins}}, # <-- Usamos la lista
    methods=["GET", "POST", "PUT", "DELETE", "PATCH", "OPTIONS"],
    allow_headers=["Content-Type", "Authorization"],
    supports_credentials=True
)


_llm_manager_instance = None
_llm_manager_lock = threading.Lock()

def get_llm_manager():
    global _llm_manager_instance
    with _llm_manager_lock:
        if _llm_manager_instance is None:
            logger.info("Inicializando ModelManager por primera vez...")
            _llm_manager_instance = ModelManager()
            
            _llm_manager_instance.setup_api_keys(
                google_api_key=os.getenv("GOOGLE_API_KEY")
            )
            
            # ---> AÑADE ESTA LÍNEA <---
            # Después de configurar las APIs, precarga los modelos de visión.
            _llm_manager_instance.preload_vision_models()
            
            logger.info("✅ ModelManager listo, APIs configuradas y modelos de visión precargados.")
        return _llm_manager_instance

logger.info("Inicializando servicios de la aplicación...")
model_manager = get_llm_manager()

cliente_public: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
supabase_admin: Client = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)
project_service = ProjectService(db_client=supabase_admin, storage_handler=supabase_handler)
dataset_service = DataService(db_client=supabase_admin, storage_handler=supabase_handler)
text_analysis_service = TextAnalysisService()
prediction_service = PredictionService()
executor = ThreadPoolExecutor(max_workers=5)
history_service = HistoryService() 
prompt_lab_service = PromptLabService(model_manager=model_manager, history_service=history_service)
vision_prediction_service = VisionPredictionService()
tabular_clustering_service = TabularClusteringService()

# 2. Inmediatamente después, configuramos las dependencias.
tabular_clustering_service.configure(
    dataset_service=dataset_service,
    supabase_handler=supabase_handler
)

VALID_PROVIDERS = {"google", "openai", "aws", "azure"}

logger.info("Conectando FAQService con PromptLabService...")
faq_service = FAQService(prompt_lab_service=prompt_lab_service)
logger.info("✅ FAQService listo y conectado.")

EXECUTION_TIMEOUT_SEC = 180

MAX_IMAGE_SIZE_MB = 10  # Tamaño máximo de 10 MB por imagen
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'webp'}

def is_allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS



class ValidationError(Exception):
    pass

@app.errorhandler(ValidationError)
def handle_validation_error(error):
    response = jsonify({"success": False, "error": str(error)})
    response.status_code = 400
    return response


def token_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # --- VUELVE A AÑADIR ESTE BLOQUE ---
        if request.method == 'OPTIONS':
            # Esto intercepta la petición de comprobación del navegador
            # y le dice "Sí, tienes permiso para preguntar".
            # La respuesta real a la petición OPTIONS la construirá Flask-CORS
            # con las cabeceras correctas (Allow-Headers, etc.).
            # Devolver una respuesta OK simple desde aquí es una forma segura de hacerlo.
            return jsonify({'status': 'ok'}), 200

        auth_header = request.headers.get("Authorization")
        if not auth_header or not auth_header.startswith("Bearer "):
            return jsonify({"success": False, "error": "Token de autorización 'Bearer' no encontrado."}), 401
        
        token = auth_header.split(" ")[1]
        try:
            user_data = cliente_public.auth.get_user(token)
            current_user = user_data.user
            if not current_user: raise AuthApiError("Usuario no válido.", 401)
        except AuthApiError as e:
            return jsonify({"success": False, "error": f"Token inválido o expirado: {e.message}"}), 401
        
        return f(current_user, *args, **kwargs)
    return decorated_function





# --- 5. ENDPOINTS DE LA API ---

@app.route("/api/health")
def health_check():
    return jsonify({"status": "ok", "message": "Backend PromptLab funcionando."})

# === RUTAS DE GESTIÓN DE PROYECTOS (CRUD) ===

@app.route("/api/projects", methods=["GET"])
@token_required
def get_projects(current_user):
    result = project_service.get_projects_by_user(current_user.id)
    status_code = 200 if result.get("success") else 400
    return jsonify(result), status_code

# === AQUÍ ESTÁ LA VERSIÓN REFINADA ===
@app.route("/api/projects", methods=["POST"])
@token_required
def create_project(current_user):
    """Crea un nuevo proyecto para el usuario autenticado, con validación robusta."""
    data = request.get_json()
    
    # Un solo bloque de validación que utiliza tu manejador de errores.
    if not data or not isinstance(data.get("project_name"), str) or not data.get("project_name", "").strip():
        raise ValidationError("El campo 'project_name' es obligatorio y debe ser un texto válido.")
    
    project_name = data["project_name"].strip()
        
    result = project_service.create_project(
        user_id=current_user.id,
        project_name=project_name,
        project_type=data.get("project_type", "indefinido")
    )
    status_code = 201 if result.get("success") else 400
    return jsonify(result), status_code


@app.route("/api/projects/<string:project_id>", methods=["GET"])
@token_required
def get_project(current_user, project_id):
    result = project_service.get_project_by_id(project_id, current_user.id)
    return jsonify(result), 200 if result.get("success") else 404



# === RUTAS DE GESTIÓN DE DATASETS ===

@app.route("/api/projects/<string:project_id>", methods=["PUT"])
@token_required
def rename_project_endpoint(current_user, project_id):
    """
    Endpoint para renombrar un proyecto. Recibe un JSON con 'project_name'.
    """
    data = request.get_json()
    
    if not data or not data.get("project_name"):
        return jsonify({"success": False, "error": "El nuevo nombre del proyecto ('project_name') es obligatorio."}), 400
    
    new_name = data["project_name"]
    
    # Llama a la función del servicio que ya tienes y que funciona bien
    result = project_service.rename_project(
        project_id=project_id,
        user_id=current_user.id,
        new_name=new_name
    )
    
    status_code = 200 if result.get("success") else 400
    return jsonify(result), status_code



@app.route("/api/projects/<string:project_id>/datasets/<string:dataset_id>", methods=["DELETE"])
@token_required
def delete_dataset_endpoint(current_user, project_id, dataset_id):
    # Renombrado para evitar conflicto de nombres con la función del servicio
    result = dataset_service.delete_dataset(dataset_id, current_user.id)
    status_code = 200 if result["success"] else 400
    return jsonify(result), status_code

@app.route("/api/projects/<string:project_id>", methods=["DELETE"])
@token_required
def delete_project(current_user, project_id):
    """
    Elimina un proyecto específico y todos sus datasets asociados.
    """
    # La función del servicio ya necesita el dataset_service, así que se lo pasamos
    result = project_service.delete_project(
        project_id=project_id, 
        user_id=current_user.id, 
        dataset_service=dataset_service
    )
    
    status_code = 200 if result.get("success") else 400
    return jsonify(result), status_code


@app.route("/api/projects/<string:project_id>/datasets", methods=["GET"])
@token_required
def list_datasets_in_project(current_user, project_id):
    """
    Lista todos los datasets de un proyecto específico.
    """
    result = dataset_service.list_datasets(project_id, current_user.id)
    
    # Aplicamos la conversión a camelCase que ya tienes
    if result.get("success") and isinstance(result.get("data"), list):
        result["data"] = [camelize_dict(item) for item in result["data"]]

    return jsonify(result), 200 if result["success"] else 400



# === RUTAS DE GESTIÓN DE DATASETS ===

@app.route("/api/projects/<string:project_id>/datasets/upload", methods=["POST"])
@token_required
def upload_dataset(current_user, project_id):
    """
    Recibe un archivo y delega toda la lógica de guardado al DataService.
    """
    try:
        if 'file' not in request.files:
            return jsonify({"success": False, "error": "No se encontró el archivo en la petición."}), 400

        file = request.files['file']
        if not file or file.filename == '':
            return jsonify({"success": False, "error": "No se seleccionó ningún archivo válido."}), 400
        
        # Simplemente pasamos el archivo al servicio, que se encargará de todo
        result = dataset_service.create_dataset_from_file(
            user_id=current_user.id,
            project_id=project_id,
            file=file
        )
        
        status_code = 201 if result.get("success") else 400
        return jsonify(result), status_code

    except Exception as e:
        logger.error(f"Error fatal en la subida de archivo: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500
    
    
@app.route("/api/datasets/<string:dataset_id>/details", methods=["GET"])
@token_required
def get_dataset_details(current_user, dataset_id):
    """
    Devuelve metadatos clave de un dataset, como su tipo y nombre.
    """
    try:
        # Consulta a Supabase para obtener solo la información que necesitamos
        response = supabase_admin.table('datasets') \
            .select('dataset_name, dataset_type, project_id') \
            .eq('dataset_id', dataset_id) \
            .eq('user_id', current_user.id) \
            .single() \
            .execute()

        if not response.data:
            return jsonify({"success": False, "error": "Dataset no encontrado"}), 404
        
        return jsonify({"success": True, "data": response.data}), 200

    except Exception as e:
        logger.error(f"Error en get_dataset_details para dataset {dataset_id}: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor"}), 500

def to_camel_case(snake_str):
    components = snake_str.split('_')
    if len(components) == 1:
        return snake_str
    return components[0] + ''.join(x.title() for x in components[1:])


def camelize_dict(data):
    if isinstance(data, dict):
        return {to_camel_case(k): camelize_dict(v) for k, v in data.items()}
    elif isinstance(data, list):
        return [camelize_dict(item) for item in data]
    return data


def is_valid_url(url):
    try:
        result = urlparse(url)
        return all([result.scheme, result.netloc])
    except:
        return False


@app.route("/api/projects/<string:project_id>/datasets/import/api", methods=["POST"])
@token_required
def import_dataset_api(current_user, project_id):
    data = request.get_json()
    if not data:
        return jsonify({"success": False, "error": "Datos JSON obligatorios."}), 400

    dataset_name = data.get("dataset_name")
    api_url = data.get("api_url")
    params = data.get("params")
    headers = data.get("headers")
    data_path = data.get("data_path")
    max_records = data.get("max_records", 10000)

    if not dataset_name or not api_url:
        return jsonify({"success": False, "error": "dataset_name y api_url son obligatorios."}), 400

    if not is_valid_url(api_url):
        return jsonify({"success": False, "error": "api_url no es una URL válida."}), 400

    result = dataset_service.import_from_api_json(
        user_id=current_user.id,
        project_id=project_id,
        dataset_name=dataset_name,
        api_url=api_url,
        params=params,
        headers=headers,
        data_path=data_path,
        max_records=max_records
    )

    status_code = 201 if result.get("success") else 400
    return jsonify(result), status_code



@app.route("/api/projects/<string:project_id>/datasets/<string:dataset_id>", methods=["PATCH"])
@token_required
def rename_dataset_endpoint(current_user, project_id, dataset_id):
    """
    Endpoint para renombrar un dataset. Recibe un JSON con 'dataset_name'.
    """
    data = request.get_json()

    # Validación del input
    if not data or not data.get("dataset_name") or not data.get("dataset_name").strip():
        return jsonify({"success": False, "error": "El nuevo nombre del dataset ('dataset_name') es obligatorio y no puede estar vacío."}), 400

    new_name = data["dataset_name"].strip()

    # Llamada a la nueva función del servicio
    result = dataset_service.rename_dataset(
        dataset_id=dataset_id,
        user_id=current_user.id,
        new_name=new_name
    )

    status_code = 200 if result.get("success") else 400
    if result.get("success"):
         # Convertimos el resultado a camelCase para consistencia con el frontend
        result["data"] = camelize_dict(result["data"])
        
    return jsonify(result), status_code

# ================================================================
#    ENDPOINT: EDITAR COLUMNAS DE UN DATASET (ELIMINAR/RENOMBRAR)
# ================================================================
@app.route('/api/datasets/<string:dataset_id>/edit-columns', methods=['POST'])
@token_required
def edit_columns(current_user, dataset_id):
    """
    Aplica acciones de edición de columnas (eliminar o renombrar) a un dataset.
    """
    try:
        # --- PASO 1: CARGAR INFO DEL DATASET ---
        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        if not dataset_info or not dataset_info.get("success"):
            return jsonify({"success": False, "error": "Dataset no encontrado o acceso denegado"}), 404

        dataset_data = dataset_info["data"]
        storage_path = dataset_data.get("storage_path")
        if not storage_path:
            return jsonify({"success": False, "error": "Falta la ruta de almacenamiento del dataset"}), 500

        # --- PASO 1b: LEER ARCHIVO DIRECTAMENTE CON HANDLER ---
        df = supabase_handler.load_file_as_dataframe(user_id=current_user.id, path=storage_path)
        if df is None or df.empty:
            return jsonify({"success": False, "error": "Error al leer el archivo o archivo vacío."}), 400
        print("\n--- DEBUG: ENDPOINT /edit-columns ---")
        print(f"[ENTRADA] Shape del DataFrame: {df.shape}")
        print(f"[ENTRADA] Duplicados: {df.duplicated().sum()}")
        print(f"[ENTRADA] Columnas: {df.columns.tolist()}")
        

        # --- PASO 2: RECIBIR INSTRUCCIONES DEL FRONTEND ---
        request_data = request.get_json()
        action = request_data.get("action")
        params = request_data.get("params", {})

        if not action:
            return jsonify({"success": False, "error": "Falta el campo 'action' en el cuerpo del request"}), 400

        # --- PASO 3: EJECUTAR LA ACCIÓN ---
        cleaning_result = apply_cleaning_action(df, action, params)

        if not cleaning_result.get("success"):
            return jsonify(cleaning_result), 400

        df_limpio = pd.DataFrame(cleaning_result["dataframe"])
        print("Columnas del DF LIMPIO (después de la acción):", df_limpio.columns.tolist())
        print(f"\n--- DEBUG: ENDPOINT /edit-columns (DESPUÉS) ---")
        print(f"[SALIDA] Shape del DataFrame: {df_limpio.shape}")
        print(f"[SALIDA] Duplicados: {df_limpio.duplicated().sum()}")
        print(f"[SALIDA] Columnas: {df_limpio.columns.tolist()}")
        print(f"--- FIN DEBUG ---\n")

        # --- PASO 3b: GENERAR ANÁLISIS ---
        analysis_result = get_preliminary_analysis(df_limpio)

        # --- PASO 4: GUARDAR RESULTADO EN STORAGE ---
        try:
            supabase_handler.save_dataframe_to_storage(
                user_id=current_user.id,
                path=storage_path,
                df=df_limpio
            )
            print(f"--- DEBUG STORAGE: Dataset actualizado en {storage_path}")
        except Exception as storage_error:
            logger.error(f"Error al guardar dataset actualizado: {storage_error}", exc_info=True)
            return jsonify({"success": False, "error": "Error al guardar dataset actualizado."}), 500

        # --- PASO 5: DEVOLVER RESULTADO ---
        if analysis_result.get("success"):
            analysis_data = analysis_result["data"]  # La data con la estructura "plana"
            
            # --- INICIO DEL BLOQUE DE TRANSFORMACIÓN ---
            formatted_diagnostics = {
                "summary": {
                    "totalRows": analysis_data["project_summary"]["rows"],
                    "totalColumns": analysis_data["project_summary"]["columns"],
                    "duplicates": analysis_data["duplicates_summary"],
                    "nanValues": {
                        "totalCount": analysis_data["nulls_summary"]["total_count"],
                        "percentage": analysis_data["nulls_summary"]["total_percentage"]
                    }
                },
                "columnDetails": analysis_data["nulls_summary"]["by_column"],
                "columns_info": analysis_data["columns_info"]
            }
            # --- FIN DEL BLOQUE ---

            # Construcción de respuesta final bien estructurada
            final_response_data = {
                "diagnostics": formatted_diagnostics,
                "previewData": df_limpio.head(50).fillna("N/A").to_dict(orient="records"),
                "cleaning_message": cleaning_result.get("message", "Transformación aplicada con éxito.")
            }

            return jsonify({
                "success": True,
                "data": final_response_data
            })

        else:
            return jsonify({
                "success": False,
                "error": "Error al generar análisis del dataset modificado."
            }), 500

    except Exception as e:
        logger.error(f"[edit_columns] Error en el endpoint edit-columns: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor"}), 500


@app.route("/api/datasets/<string:dataset_id>/content", methods=["PUT"])
@token_required
def update_dataset_content(current_user, dataset_id):
    try:
        data = request.get_json()
        new_content = data.get("newContent")

        if new_content is None:
            return jsonify({"success": False, "error": "Falta el campo 'newContent'."}), 400

        # --- LÓGICA DIRECTA Y CORRECTA ---
        # No llamamos a un servicio, hacemos la operación simple aquí.
        # Actualizamos la columna 'contenido_texto_editado' y nada más.
        response = supabase_admin.table('datasets') \
            .update({'contenido_texto_editado': new_content}) \
            .eq('dataset_id', dataset_id) \
            .eq('user_id', current_user.id) \
            .execute()

        if not response.data:
             return jsonify({"success": False, "error": "No se pudo actualizar el documento."}), 404

        logger.info(f"Contenido editado para el dataset {dataset_id} guardado en DB.")
        
        return jsonify({"success": True, "message": "Contenido del documento guardado con éxito."}), 200

    except Exception as e:
        logger.error(f"Error en update_dataset_content para dataset {dataset_id}: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno al guardar el contenido."}), 500


# ================================================================
#    ENDPOINT: OBTENER CONTENIDO DE TEXTO DE UN DATASET
# ================================================================
@app.route('/api/datasets/<string:dataset_id>/get-text-content', methods=['GET'])
@token_required
def get_dataset_text_content(current_user, dataset_id):
    """
    Devuelve el contenido textual de un dataset.
    - Si existe texto previamente editado, lo devuelve.
    - Si no, convierte el archivo original (parquet, xlsx, csv) a texto plano CSV.
    """
    try:
        # --- 1. Consulta segura: traer path y texto editado ---
        response = (
            supabase_admin.table('datasets')
            .select('storage_path, contenido_texto_editado')
            .eq('dataset_id', dataset_id)
            .eq('user_id', current_user.id)
            .single()
            .execute()
        )

        if not response.data:
            logger.warning(f"[404] Dataset no encontrado o acceso denegado. ID: {dataset_id}, User: {current_user.id}")
            return jsonify({"success": False, "error": "Dataset no encontrado o acceso denegado"}), 404

        dataset_data = response.data

        # --- 2. Devolver contenido editado si existe ---
        if dataset_data.get('contenido_texto_editado') is not None:
            logger.info(f"[200] Devolviendo contenido editado para dataset {dataset_id}")
            return jsonify({
                "success": True,
                "textContent": dataset_data['contenido_texto_editado']
            }), 200

        # --- 3. Convertir archivo original a CSV usando handler ---
        storage_path = dataset_data.get("storage_path")
        if not storage_path:
            logger.error(f"[500] Falta storage_path para dataset {dataset_id}")
            return jsonify({"success": False, "error": "Ruta del archivo no disponible"}), 500

        csv_text_content = supabase_handler.load_file_as_csv(user_id=current_user.id, path=storage_path)
        if csv_text_content is None:
            logger.error(f"[500] No se pudo cargar o convertir archivo a CSV. Dataset {dataset_id}")
            return jsonify({"success": False, "error": "No se pudo cargar o convertir el archivo."}), 400

        logger.info(f"[200] Archivo convertido exitosamente a CSV para dataset {dataset_id}")

        # --- 4. Devolver CSV como texto ---
        return jsonify({
            "success": True,
            "textContent": csv_text_content
        }), 200

    except Exception as e:
        logger.critical(f"[500] Error inesperado en get-text-content para dataset {dataset_id}: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor"}), 500

    
    # --- guardar cambios edicion tabulares---

@app.route('/api/datasets/<string:dataset_id>/save-text-content', methods=['POST'])
@token_required
def save_dataset_text_content(current_user, dataset_id):
    try:
        # --- 1. Obtener info y el contenido de la petición (esto se queda igual) ---
        request_data = request.get_json()
        new_content_str = request_data.get('textContent')
        if not new_content_str:
            return jsonify({"success": False, "error": "El campo 'textContent' es requerido."}), 400

        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        if not dataset_info or not dataset_info.get("success"):
            return jsonify({"success": False, "error": "Dataset no encontrado o acceso denegado."}), 404
        
        storage_path = dataset_info["data"].get("storage_path")
        file_ext = os.path.splitext(storage_path)[-1].lower().replace('.', '')

        # --- 2. Convertir el texto editado en un DataFrame ---
        # El frontend nos envía un string en formato CSV, así que lo leemos.
        try:
            df_modificado = pd.read_csv(StringIO(new_content_str))
            logger.info(f"Texto editado convertido a DataFrame. {df_modificado.shape[0]} filas encontradas.")
        except Exception as e:
            logger.error(f"Error al parsear el CSV del editor: {e}")
            return jsonify({"success": False, "error": "El formato del texto editado no es un CSV válido."}), 400

        # --- 3. Convertir el DataFrame al formato de archivo original y guardarlo en Supabase ---
        # Esto SOBRESCRIBE el archivo viejo con los nuevos datos.
        try:
            buffer_out = BytesIO()
            if file_ext == 'parquet':
                df_modificado.to_parquet(buffer_out, index=False)
            elif file_ext == 'xlsx':
                df_modificado.to_excel(buffer_out, index=False)
            else: # Asumimos CSV como default
                df_modificado.to_csv(buffer_out, index=False)
            
            buffer_out.seek(0)
            
            # Subimos el buffer para sobrescribir el archivo en el Storage
            supabase_admin.storage.from_("proyectos-usuarios").update(
                path=storage_path,
                file=buffer_out.getvalue(),
                file_options={
                    "content-type": "application/octet-stream",
                    "cache-control": "0",
                    "upsert": "true"
                }
                )
            logger.info(f"Archivo '{storage_path}' sobrescrito con éxito en Supabase Storage.")

        except Exception as e:
            logger.error(f"Error al guardar el archivo modificado en Supabase: {e}")
            return jsonify({"success": False, "error": "No se pudo guardar el archivo modificado."}), 500

        # --- 4. [LA MAGIA] Generar un nuevo diagnóstico y devolverlo ---
    
        nuevo_diagnostico_result = orquestador.get_preliminary_analysis(df_modificado)
        if not nuevo_diagnostico_result.get("success"):
            logger.error("Archivo guardado, pero falló el re-análisis.")
            return jsonify({"success": False, "error": "Archivo guardado, pero no se pudo generar el nuevo diagnóstico."}), 500

        # --- 5. Construir y enviar la respuesta rica en datos ---
        analysis_data = nuevo_diagnostico_result["data"]
        
        # Aseguramos que el preview también sea seguro para JSON
        df_preview_temp = df_modificado.head(100).copy()
        df_preview_object = df_preview_temp.astype(object)
        df_preview_filled = df_preview_object.fillna("N/A")
        preview_data = df_preview_filled.to_dict(orient="records")


        formatted_diagnostics = {
            "summary": {
               "totalRows": analysis_data["project_summary"]["rows"],
               "totalColumns": analysis_data["project_summary"]["columns"],
               "duplicates": analysis_data["duplicates_summary"],
               "nanValues": {
               "totalCount": analysis_data["nulls_summary"]["total_count"],
               "percentage": analysis_data["nulls_summary"]["total_percentage"]
                }
                },
            "columnDetails": analysis_data["nulls_summary"]["by_column"], # <-- La clave correcta
            "columns_info": analysis_data["columns_info"] 
          }


        response_data = {
            "diagnostics": formatted_diagnostics,
            "previewData": preview_data
        }
        
        return jsonify({
            "success": True, 
            "message": "Cambios guardados y aplicados al dataset principal.",
            "data": response_data # Devolvemos los datos frescos
        })

    except Exception as e:
        logger.exception(f"Error inesperado en save_text_content para dataset {dataset_id}")
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500



@app.route('/api/datasets/<string:dataset_id>/get-dataset-content', methods=['GET'])
@token_required
def get_dataset_promptlab_content(current_user, dataset_id):
    """
    Devuelve el contenido textual o tabular de un dataset.
    - Para documentos (pdf, docx, txt, json): devuelve texto.
    - Para archivos tabulares (csv, xlsx, parquet): devuelve JSON (filas).
    Priorizamos el contenido editado si existe (para cualquier tipo de dataset).
    """
    BUCKET_NAME = 'proyectos-usuarios'

    try:
        # 1. Obtener información del dataset
        response = (
            supabase_admin.table('datasets')
            .select('storage_path, contenido_texto_editado, dataset_type')
            .eq('dataset_id', dataset_id)
            .eq('user_id', current_user.id)
            .single()
            .execute()
        )

        if not response.data:
            logger.warning(f"Dataset {dataset_id} no encontrado para el usuario {current_user.id}")
            return jsonify({"success": False, "error": "Dataset no encontrado"}), 404

        dataset_data = response.data
        dataset_type = dataset_data.get("dataset_type", "unknown").lower()
        storage_path = dataset_data.get("storage_path")

        if not storage_path:
            logger.error(f"Dataset {dataset_id} no tiene ruta de archivo (storage_path)")
            return jsonify({"success": False, "error": "Ruta de archivo no disponible"}), 500

        # 2. Si hay contenido editado, devolverlo SIEMPRE (independiente del tipo)
        contenido_editado = dataset_data.get('contenido_texto_editado')
        if contenido_editado is not None:
            return jsonify({
                "success": True,
                "textContent": contenido_editado.strip(),
                "datasetType": dataset_type
            })

        # 3. Descargar archivo del storage
        try:
            file_bytes = supabase_admin.storage.from_(BUCKET_NAME).download(storage_path)
        except Exception as storage_err:
            logger.error(
                f"Error al descargar archivo desde storage para dataset {dataset_id}: {storage_err}",
                exc_info=True
            )
            return jsonify({"success": False, "error": "No se pudo acceder al archivo"}), 500

        result = {}

        # 4. Procesar según tipo de dataset
        try:
            if dataset_type == 'text':
                for enc in ['utf-8', 'latin-1', 'windows-1252']:
                    try:
                        result["textContent"] = file_bytes.decode(enc)
                        break
                    except UnicodeDecodeError:
                        continue
                if not result.get("textContent"):
                    raise ValueError("No se pudo decodificar el archivo de texto")

            elif dataset_type == 'pdf':
                with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                    result["textContent"] = "\n".join(page.get_text("text") for page in doc)

            elif dataset_type == 'docx':
                doc = Document(BytesIO(file_bytes))
                result["textContent"] = "\n".join(p.text for p in doc.paragraphs)

            elif dataset_type == 'json':
                data = json.loads(file_bytes.decode('utf-8'))
                result["textContent"] = json.dumps(data, indent=2, ensure_ascii=False)

            # --- TABULARES ---
            elif dataset_type == 'csv':
                df = pd.read_csv(BytesIO(file_bytes))
                result["tableContent"] = df.head(100).to_dict(orient="records")

            elif dataset_type == 'xlsx':
                df = pd.read_excel(BytesIO(file_bytes))
                result["tableContent"] = df.head(100).to_dict(orient="records")

            elif dataset_type == 'parquet':
                df = pd.read_parquet(BytesIO(file_bytes))
                result["tableContent"] = df.head(100).to_dict(orient="records")

            else:
                logger.warning(f"Tipo de archivo '{dataset_type}' no soportado para dataset {dataset_id}")
                return jsonify({
                    "success": False,
                    "error": f"Tipo de archivo '{dataset_type}' no soportado por este endpoint."
                }), 400

            result["datasetType"] = dataset_type
            return jsonify({"success": True, **result})

        except Exception as processing_err:
            logger.error(
                f"Error procesando archivo del dataset {dataset_id}: {processing_err}",
                exc_info=True
            )
            return jsonify({"success": False, "error": "Error al procesar el archivo"}), 500

    except Exception as e:
        logger.critical(
            f"Error inesperado en get-dataset-content para dataset {dataset_id}",
            exc_info=True
        )
        return jsonify({"success": False, "error": "Error interno del servidor"}), 500




@app.route('/api/datasets/<string:dataset_id>/get-document-text-content', methods=['GET'])
@token_required
def get_document_text_content(current_user, dataset_id):
    """
    Devuelve el contenido textual de un dataset de tipo documento (pdf, docx, json, txt).
    Prioriza el texto editado si existe. Si no, extrae desde el archivo original.
    """
    BUCKET_NAME = 'proyectos-usuarios'

    try:
        # 1. Obtener información del dataset
        response = supabase_admin.table('datasets') \
            .select('storage_path, contenido_texto_editado, dataset_type') \
            .eq('dataset_id', dataset_id) \
            .eq('user_id', current_user.id) \
            .single() \
            .execute()

        if not response.data:
            logger.warning(f"Dataset {dataset_id} no encontrado para el usuario {current_user.id}")
            return jsonify({"success": False, "error": "Dataset no encontrado"}), 404

        dataset_data = response.data

        # 2. Si hay contenido editado, devolverlo
        contenido_editado = dataset_data.get('contenido_texto_editado')
        if contenido_editado is not None:
            return jsonify({"success": True, "textContent": contenido_editado.strip()})

        # 3. Extraer texto del archivo original
        storage_path = dataset_data.get("storage_path")
        dataset_type = dataset_data.get("dataset_type", "unknown").lower()

        if not storage_path:
            logger.error(f"Dataset {dataset_id} no tiene ruta de archivo (storage_path)")
            return jsonify({"success": False, "error": "Ruta de archivo no disponible"}), 500

        try:
            file_bytes = supabase_admin.storage.from_(BUCKET_NAME).download(storage_path)
        except Exception as storage_err:
            logger.error(f"Error al descargar archivo desde storage para dataset {dataset_id}: {storage_err}", exc_info=True)
            return jsonify({"success": False, "error": "No se pudo acceder al archivo"}), 500

        text_content = ""

        try:
            if dataset_type == 'text':
                for enc in ['utf-8', 'latin-1', 'windows-1252']:
                    try:
                        text_content = file_bytes.decode(enc)
                        break
                    except UnicodeDecodeError:
                        continue
                if not text_content:
                    raise ValueError("No se pudo decodificar el archivo de texto")

            elif dataset_type == 'pdf':
                with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                    text_content = "\n".join(page.get_text("text") for page in doc)

            elif dataset_type == 'docx':
                doc = Document(BytesIO(file_bytes))
                text_content = "\n".join(p.text for p in doc.paragraphs)

            elif dataset_type == 'json':
                try:
                    data = json.loads(file_bytes.decode('utf-8'))
                    text_content = json.dumps(data, indent=2, ensure_ascii=False)
                except Exception as e:
                    raise ValueError("Error procesando archivo JSON") from e

            else:
                logger.warning(f"Tipo de archivo '{dataset_type}' no soportado para dataset {dataset_id}")
                return jsonify({
                    "success": False,
                    "error": f"Tipo de archivo '{dataset_type}' no soportado por este endpoint."
                }), 400

            return jsonify({"success": True, "textContent": text_content.strip(), "datasetType": dataset_type})

        except Exception as processing_err:
            logger.error(f"Error procesando archivo del dataset {dataset_id}: {processing_err}", exc_info=True)
            return jsonify({"success": False, "error": "Error al procesar el archivo"}), 500

    except Exception as e:
        logger.critical(f"Error inesperado en get-document-text-content para dataset {dataset_id}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor"}), 500



# ==========================================================
# ENDPOINTS DE ANÁLISIS (VERSIÓN FINAL Y CONSISTENTE)
# ==========================================================

# --- Endpoint para Resumen ---
@app.route("/api/analysis/summary", methods=["POST"])
@token_required
def analyze_summary(current_user):
    try:
        # 1. Validar input
        data = request.get_json()
        text = data.get("text", "").strip()
        # La validación de texto vacío ahora la hará el servicio, pero es bueno tenerla aquí también
        if not text:
            return jsonify({"error": "El cuerpo de la petición debe contener un campo 'text' no vacío."}), 400

        # 2. Llamar al método público del servicio
        result = text_analysis_service.do_summary(text)
        
        # 3. Devolver respuesta exitosa
        response_data = {"analysis_type": "summary", "result": result}
        return jsonify(response_data), 200

    # 4. MANEJO DE ERRORES ESPECÍFICOS
    except ValueError as e:
        # Captura errores como "texto vacío" o "no hay oraciones".
        # Estos son errores del cliente.
        return jsonify({"error": str(e)}), 422  # 422: Unprocessable Entity

    except NotImplementedError as e:
        # Captura el error si 'sumy' no está instalado.
        # Es un error de configuración del servidor.
        return jsonify({"error": str(e)}), 501  # 501: Not Implemented

    except RuntimeError as e:
        # Captura el fallo del algoritmo de resumen.
        # Este es un error interno del servidor.
        logger.error(f"Error de ejecución en /api/analysis/summary: {e}", exc_info=True)
        return jsonify({"error": "El algoritmo de resumen no pudo procesar el texto."}), 500
    
    except Exception as e:
        # Captura cualquier otro error inesperado.
        logger.error(f"Error inesperado en /api/analysis/summary: {e}", exc_info=True)
        return jsonify({"error": "Ocurrió un error interno al generar el resumen."}), 500



@app.route("/api/analysis/sentiment", methods=["POST"])
@token_required
def analyze_sentiment(current_user):
    # ¡Este ya estaba perfecto! No necesita cambios.
    try:
        data = request.get_json()
        text = data.get("text", "").strip()
        if not text:
            return jsonify({"error": "Texto vacío o inválido."}), 400

        result = text_analysis_service.do_sentiment(text)
        
        response_data = {"analysis_type": "sentiment", "result": result}
        return jsonify(response_data), 200

    except NotImplementedError as e:
        return jsonify({"error": str(e)}), 501
    except Exception as e:
        logger.error(f"Error en /api/analysis/sentiment: {e}", exc_info=True)
        return jsonify({"error": "Ocurrió un error interno al analizar el sentimiento."}), 500
    

# --- Endpoint para Entidades ---
@app.route("/api/analysis/entities", methods=["POST"])
@token_required
def analyze_entities(current_user):
    # ¡Perfecto!
    try:
        data = request.get_json()
        text = data.get("text", "").strip()
        if not text:
            return jsonify({"error": "Texto vacío o inválido."}), 400

        result = text_analysis_service.do_entities(text)

        response_data = {"analysis_type": "entities", "result": result}
        return jsonify(response_data), 200

    except Exception as e:
        logger.error(f"Error en /api/analysis/entities: {e}", exc_info=True)
        return jsonify({"error": "Ocurrió un error interno al extraer entidades."}), 500

# --- Endpoint para Tópicos ---

@app.route("/api/analysis/topics", methods=["POST"])
@token_required
def analyze_topics(current_user):
    try:
        data = request.get_json()
        text = data.get("text", "").strip()
        if not text:
            return jsonify({"success": False, "error": "El texto no puede estar vacío."}), 400

        # --- VALIDACIÓN AÑADIDA AQUÍ ---
        # Definimos un umbral mínimo de palabras. 25 es un número razonable para empezar.
        MIN_WORDS_THRESHOLD = 25
        word_count = len(text.split())

        if word_count < MIN_WORDS_THRESHOLD:
            # ¡Este es el mensaje claro que quieres enviar!
            error_message = f"El texto es demasiado corto. Se necesitan al menos {MIN_WORDS_THRESHOLD} palabras para el análisis de tópicos."
            # Usamos 422 porque la solicitud es válida, pero no procesable.
            return jsonify({"success": False, "error": error_message}), 422
        # --- FIN DE LA VALIDACIÓN ---

        # Si pasamos la validación, llamamos al servicio.
        result = text_analysis_service.do_topics(text)
        
        # Envolvemos la respuesta exitosa para consistencia
        response_data = {"success": True, "analysis_type": "topics", "result": result}
        return jsonify(response_data), 200

    except ValueError as e:
        # Este catch ahora sirve como un seguro por si la librería falla por otra razón
        logger.warning(f"ValueError en do_topics a pesar de la validación: {e}")
        return jsonify({"success": False, "error": "No se pudo procesar el texto para el análisis de tópicos."}), 422
    except Exception as e:
        logger.error(f"Error en /api/analysis/topics: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Ocurrió un error interno al modelar tópicos."}), 500
    
    
    
# --- Endpoint para Estadísticas ---
@app.route("/api/analysis/statistics", methods=["POST"])
@token_required
def analyze_statistics(current_user):
    # ¡Perfecto!
    try:
        data = request.get_json()
        text = data.get("text", "").strip()
        if not text:
            return jsonify({"error": "Texto vacío o inválido."}), 400

        result = text_analysis_service.do_statistics(text)

        response_data = {"analysis_type": "statistics", "result": result}
        return jsonify(response_data), 200

    except Exception as e:
        logger.error(f"Error en /api/analysis/statistics: {e}", exc_info=True)
        return jsonify({"error": "Ocurrió un error interno al extraer estadísticas."}), 500

# --- Endpoint para Cláusulas ---
@app.route("/api/analysis/clauses", methods=["POST"])
@token_required
def analyze_clauses(current_user):
    # ¡Perfecto!
    try:
        data = request.get_json()
        text = data.get("text", "").strip()
        if not text:
            return jsonify({"error": "Texto vacío o inválido."}), 400

        result = text_analysis_service.do_clauses(text)

        response_data = {"analysis_type": "clauses", "result": result}
        return jsonify(response_data), 200

    except Exception as e:
        logger.error(f"Error en /api/analysis/clauses: {e}", exc_info=True)
        return jsonify({"error": "Ocurrió un error interno al extraer cláusulas."}), 500
    
    
# ==========================================================
# ENDPOINTS DE ANÁLISIS POR PERFIL
# ==========================================================

# --- Endpoint para Perfil Legal ---
@app.route("/api/analysis/profile/legal", methods=["POST"])
@token_required
def analyze_profile_legal(current_user):
    try:
        # 1. Validar input (igual que siempre)
        data = request.get_json()
        text = data.get("text", "").strip()
        if not text:
            return jsonify({"error": "El cuerpo de la petición debe contener 'text' no vacío."}), 400

        # 2. Llamar al método público del perfil específico
        result = text_analysis_service.do_legal_analysis(text)
        
        # 3. Devolver respuesta exitosa
        # Nota: El 'analysis_type' ya viene dentro del 'result' que devuelve el servicio.
        return jsonify(result), 200

    # 4. Manejo de errores (puedes ser tan específico o genérico como necesites)
    except ValueError as e:
        # Errores de datos no procesables
        return jsonify({"error": str(e)}), 422
    except Exception as e:
        # Cualquier otro error interno
        logger.error(f"Error inesperado en /api/analysis/profile/legal: {e}", exc_info=True)
        return jsonify({"error": "Ocurrió un error interno al realizar el análisis legal."}), 500


# --- Endpoint para Perfil de Marketing ---
@app.route("/api/analysis/profile/marketing", methods=["POST"])
@token_required
def analyze_profile_marketing(current_user):
    try:
        data = request.get_json()
        text = data.get("text", "").strip()
        if not text:
            return jsonify({"error": "El cuerpo de la petición debe contener 'text' no vacío."}), 400

        # Llamada al método de perfil de marketing
        result = text_analysis_service.do_marketing_analysis(text)
        
        return jsonify(result), 200

    except ValueError as e:
        return jsonify({"error": str(e)}), 422
    except Exception as e:
        logger.error(f"Error inesperado en /api/analysis/profile/marketing: {e}", exc_info=True)
        return jsonify({"error": "Ocurrió un error interno al realizar el análisis de marketing."}), 500


# --- Endpoint para Perfil de Estilo de Escritura ---
@app.route("/api/analysis/profile/writing", methods=["POST"])
@token_required
def analyze_profile_writing(current_user):
    try:
        data = request.get_json()
        text = data.get("text", "").strip()
        if not text:
            return jsonify({"error": "El cuerpo de la petición debe contener 'text' no vacío."}), 400

        # Llamada al método de perfil de escritura
        result = text_analysis_service.do_writing_style_analysis(text)
        
        return jsonify(result), 200

    except ValueError as e:
        return jsonify({"error": str(e)}), 422
    except Exception as e:
        logger.error(f"Error inesperado en /api/analysis/profile/writing: {e}", exc_info=True)
        return jsonify({"error": "Ocurrió un error interno al realizar el análisis de estilo."}), 500
    


    # ================================================================
#    ENDPOINT: CLEAN ACTION SOBRE UNA COLUMNA DEL DATASET
# ================================================================

@app.route('/api/datasets/<string:dataset_id>/columns/clean_action', methods=['POST'])
@token_required
def clean_column_action_transformer(current_user, dataset_id):
    try:
        payload = request.get_json()
        action = payload.get('action')
        params = payload.get('params', {})

        if not action:
            return jsonify({"success": False, "error": "No se especificó acción."}), 400

        # --- 1. Obtener dataset ---
        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        if not dataset_info or not dataset_info.get("success"):
            return jsonify({"success": False, "error": "Dataset no encontrado."}), 404

        storage_path = dataset_info["data"].get("storage_path")
        if not storage_path:
            return jsonify({"success": False, "error": "Ruta del archivo no disponible."}), 500

        # --- 2. Cargar DataFrame ---
        df = supabase_handler.load_file_as_dataframe(user_id=current_user.id, path=storage_path)
        if df is None or df.empty:
            return jsonify({"success": False, "error": "No se pudo procesar el archivo."}), 400

        # --- 3. Ejecutar acción ---
        result = cleaning_action(df, action, params)
        if not result.get("success"):
            return jsonify(result), 400

        # --- 3.1 Validar DataFrame limpio ---
        df_clean = result.get("cleaned_dataframe")
        if df_clean is None:
            df_clean_data = result.get("dataframe")
            if df_clean_data is not None:
                df_clean = pd.DataFrame(df_clean_data)
            else:
                return jsonify({"success": False, "error": "Respuesta inesperada del servicio de limpieza."}), 500

        # --- 4. Guardar DataFrame ---
        supabase_handler.save_dataframe_to_storage(user_id=current_user.id, path=storage_path, df=df_clean)

        # --- 5. Re-análisis ---
        df_clean_safe = df_clean.fillna(np.nan).replace([np.nan], [None])
        analysis_result = get_preliminary_analysis(df_clean_safe)
        if not analysis_result.get("success"):
            return jsonify({"success": False, "error": "No se pudo re-analizar el dataset."}), 500

        # --- 6. Formatear respuesta ---
        analysis_data = analysis_result["data"]
        formatted_diagnostics = {
            "summary": {
                "totalRows": analysis_data["project_summary"]["rows"],
                "totalColumns": analysis_data["project_summary"]["columns"],
                "duplicates": analysis_data["duplicates_summary"],
                "nanValues": {
                    "totalCount": analysis_data["nulls_summary"]["total_count"],
                    "percentage": analysis_data["nulls_summary"]["total_percentage"]
                }
            },
            "columnDetails": analysis_data["nulls_summary"]["by_column"],
            "columns_info": analysis_data["columns_info"]
        }

        final_response_data = {
            "diagnostics": formatted_diagnostics,
            "previewData": df_clean_safe.head(100).to_dict(orient="records"),
            "cleaning_message": result.get("message", "Acción de limpieza aplicada correctamente.")
        }

        return jsonify({"success": True, "data": final_response_data}), 200

    except Exception as e:
        logger.error(f"Error en endpoint /columns/clean_action: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno en el servidor."}), 500


# ================================================================
#    ENDPOINT: CLEAN DATASET (Acciones rápidas de limpieza)
# ================================================================

@app.route('/api/datasets/<string:dataset_id>/clean', methods=['POST'])
@token_required
def clean_dataset(current_user, dataset_id):
    try:
        payload = request.get_json()
        action = payload.get('action')
        params = payload.get('params', {})

        action_map = {
            "general_delete_duplicates": "general_drop_duplicates",
            "general_delete_rows_with_nulls": "general_drop_na_rows"
        }
        
        backend_action = action_map.get(action)
        if not backend_action:
            return jsonify({"success": False, "error": f"La acción '{action}' no es válida."}), 400

        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        if not dataset_info or not dataset_info.get("success"):
            return jsonify({"success": False, "error": "Dataset no encontrado."}), 404

        storage_path = dataset_info["data"].get("storage_path")
        if not storage_path:
            return jsonify({"success": False, "error": "Ruta del archivo no disponible."}), 500

        df = supabase_handler.load_file_as_dataframe(user_id=current_user.id, path=storage_path)
        if df is None or df.empty:
            return jsonify({"success": False, "error": "No se pudo cargar el dataset."}), 400
        
        # --- INICIO DE DEPURACIÓN (PASO 1) ---
        print(f"\n--- DEBUG: CLEAN DATASET - ANTES DE LIMPIAR ---")
        print(f"Acción a ejecutar: {backend_action}")
        print(f"Shape del DataFrame original: {df.shape}")
        print(f"Duplicados en DataFrame original: {df.duplicated().sum()}")
        # --- FIN DE DEPURACIÓN ---
        
        result = orquestador.cleaning_action(df, backend_action, params)
        
        # --- INICIO DE DEPURACIÓN (PASO 2) ---
        print(f"\n--- DEBUG: CLEAN DATASET - DESPUÉS DEL ORQUESTADOR ---")
        print(f"El orquestador devolvió success: {result.get('success')}")
        if 'cleaned_dataframe' in result:
            df_clean_temp = result["cleaned_dataframe"]
            print(f"Shape del DataFrame limpio: {df_clean_temp.shape}")
            print(f"Duplicados en DataFrame limpio: {df_clean_temp.duplicated().sum()}")
        else:
            print("¡ALERTA! El resultado del orquestador NO contiene la clave 'cleaned_dataframe'.")
            print(f"Resultado completo: {result}")
        # --- FIN DE DEPURACIÓN ---

        if not result.get("success"):
            return jsonify({"success": False, "error": result.get("error", "Error en el orquestador.")}), 500

        df_clean = result["cleaned_dataframe"]
        new_diagnostics = result["new_diagnostics"]
        message = result.get("message", "Acción completada exitosamente.")

        # Aquí vamos a asumir que el guardado funciona, pero la depuración nos lo confirmará.
        supabase_handler.save_dataframe_to_storage(user_id=current_user.id, path=storage_path, df=df_clean)
        print("--- DEBUG: DataFrame limpio enviado a Supabase Handler para guardado. ---\n")

        response_data = {
            "diagnostics": {
                "summary": {
                    "totalRows": new_diagnostics["project_summary"]["rows"],
                    "totalColumns": new_diagnostics["project_summary"]["columns"],
                    "duplicates": new_diagnostics["duplicates_summary"],
                    "nanValues": {
                        "totalCount": new_diagnostics["nulls_summary"]["total_count"],
                        "percentage": new_diagnostics["nulls_summary"]["total_percentage"]
                    }
                },
                "columnDetails": new_diagnostics["nulls_summary"]["by_column"],
                "columns_info": new_diagnostics.get("columns_info", {})
            },
            "previewData": df_clean.head(100).fillna("N/A").to_dict(orient="records")
        }

        return jsonify({"success": True, "message": message, "data": response_data}), 200

    except Exception as e:
        logger.error(f"Error general en clean_dataset: {e}", exc_info=True)
     
        return jsonify({"success": False, "error": "Ocurrió un error interno durante la limpieza."}), 500



@app.route('/api/datasets/<string:dataset_id>/columns/clean', methods=['POST'])
@token_required
def clean_column_from_inspector(current_user, dataset_id):
    try:
        payload = request.get_json()
        action = payload.get('action')
        params = payload.get('params', {})
        columna_afectada = params.get('columna')

        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        storage_path = dataset_info["data"].get("storage_path")
        df = supabase_handler.load_file_as_dataframe(user_id=current_user.id, path=storage_path)
        

        if df is None:
            raise ValueError("No se pudo leer el archivo tabular.")

        result = orquestador.cleaning_action(df, action, params)
        if not result.get("success"):
            return jsonify(result), 400

        # ✅ Fix para evitar evaluar DataFrame como booleano
        cleaned_df_from_result = result.get("cleaned_dataframe")
        if cleaned_df_from_result is not None:
            df_clean = cleaned_df_from_result
        else:
            df_clean = pd.DataFrame(result.get("dataframe", {}))

        if df_clean.empty:
            raise ValueError("El orquestador no devolvió un DataFrame válido.")

        message = result.get("message", "Acción completada.")

        # Guardar en Supabase
        file_ext = os.path.splitext(storage_path)[-1].lower().replace('.', '')
        buffer_out = BytesIO()
        if file_ext == 'xlsx':
            df_clean.to_excel(buffer_out, index=False)
        elif file_ext == 'parquet':
            df_clean.to_parquet(buffer_out, index=False)
        else:
            df_clean.to_csv(buffer_out, index=False)
        buffer_out.seek(0)

        supabase_admin.storage.from_("proyectos-usuarios").update(
            path=storage_path,
            file=buffer_out.getvalue(),
            file_options={
                "content-type": "application/octet-stream",
                "cache-control": "0",
                "upsert": "true"
            }
        )

        updated_column_details = None
        if columna_afectada:
            details_result = orquestador.get_column_details(df_clean, columna_afectada)
            if details_result.get("success"):
                updated_column_details = details_result["data"]

        analysis_result = orquestador.get_preliminary_analysis(df_clean)
        analysis_data = analysis_result["data"] if analysis_result.get("success") else {}

        formatted_diagnostics = {
            "summary": {
                "totalRows": analysis_data.get("project_summary", {}).get("rows", 0),
                "totalColumns": analysis_data.get("project_summary", {}).get("columns", 0),
                "duplicates": analysis_data.get("duplicates_summary", {}),
                "nanValues": {
                    "totalCount": analysis_data.get("nulls_summary", {}).get("total_count", 0),
                    "percentage": analysis_data.get("nulls_summary", {}).get("total_percentage", 0)
                }
            },
            "columnDetails": analysis_data.get("nulls_summary", {}).get("by_column", {}),
            "columns_info": analysis_data.get("columns_info", {})
        }

        final_response_data = {
            "diagnostics": formatted_diagnostics,
            "previewData": df_clean.head(100).fillna("N/A").to_dict(orient="records"),
            "updatedColumnDetails": updated_column_details
        }

        return jsonify({"success": True, "message": message, "data": final_response_data}), 200

    except Exception as e:
        logger.error(f"Error en endpoint /columns/clean: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Ocurrió un error interno."}), 500

    
@app.route('/api/datasets/<string:dataset_id>/diagnose', methods=['GET'])
@token_required
def diagnose_dataset(current_user, dataset_id):
    try:
        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        if not dataset_info or not dataset_info.get("success"):
            return jsonify({"success": False, "error": "Dataset not found or access denied"}), 404

        dataset_data = dataset_info["data"]
        storage_path = dataset_data.get("storage_path")
        dataset_type_from_db = dataset_data.get("dataset_type", "unknown").lower().strip()

        if not storage_path:
            return jsonify({"success": False, "error": "Dataset record is incomplete"}), 500

        real_extension = (
            os.path.basename(storage_path).split('.')[-1].lower().strip()
            if '.' in os.path.basename(storage_path)
            else 'unknown'
        )

        BUCKET_NAME = 'proyectos-usuarios'
        file_content_bytes = supabase_admin.storage.from_(BUCKET_NAME).download(storage_path)

        previewData = {}
        diagnostics = {}
        df = None

        if dataset_type_from_db == 'tabular':
            try:
                df = supabase_handler.load_file_as_dataframe(user_id=current_user.id, path=storage_path)
                if df is None:
                   return jsonify({"success": False, "error": "No se pudo leer el archivo. Formato no soportado o archivo corrupto."}), 400
                print(f"\n--- DEBUG: DIAGNOSE DATASET ID {dataset_id} ---")
                print(f"[1] DataFrame COMPLETO: {df.shape[0]} filas, {df.shape[1]} columnas")
                print(f"[2] Duplicados encontrados: {df.duplicated().sum()}")

                analysis_result = get_preliminary_analysis(df)
                print(f"[3] Resultado análisis preliminar: {json.dumps(analysis_result, indent=2)[:1000]}")
                print("--- FIN DEBUG ---\n")

                if not analysis_result.get("success"):
                    return jsonify(analysis_result), 500

                analysis_data = analysis_result["data"]
                diagnostics = {
                    "summary": {
                        "totalRows": analysis_data["project_summary"]["rows"],
                        "totalColumns": analysis_data["project_summary"]["columns"],
                        "duplicates": analysis_data["duplicates_summary"],
                        "nanValues": {
                            "totalCount": analysis_data["nulls_summary"]["total_count"],
                            "percentage": analysis_data["nulls_summary"]["total_percentage"]
                        }
                    },
                    "columnDetails": analysis_data["nulls_summary"]["by_column"],
                    "columns_info": analysis_data["columns_info"] 
                }
                
                
                df_preview_temp = df.head(100).copy()
                df_preview_object = df_preview_temp.astype(object)
                df_preview_filled = df_preview_object.fillna("N/A")
                previewData = df_preview_filled.to_dict(orient="records") 

            except Exception as read_error:
                logger.warning(f"Failed to read tabular file for dataset {dataset_id}: {read_error}")
                return jsonify({"success": False, "error": str(read_error)}), 400

        elif dataset_type_from_db == 'text':
            try:
                text_content = None
                encodings_to_try = ['utf-8', 'latin-1', 'windows-1252']
                for enc in encodings_to_try:
                    try:
                        text_content = file_content_bytes.decode(enc)
                        logger.info(f"Archivo de texto decodificado con {enc}")
                        break
                    except UnicodeDecodeError:
                        continue

                if text_content is None:
                    return jsonify({"success": False, "error": "No se pudo decodificar el archivo de texto."}), 400

                diagnostics = {
                    "characters": len(text_content),
                    "lines": text_content.count('\n') + 1,
                    "size_bytes": len(file_content_bytes)
                }
                previewData = text_content

            except Exception as e:
                logger.warning(f"Error reading text file for dataset {dataset_id}: {e}")
                return jsonify({"success": False, "error": "Error leyendo el archivo de texto."}), 400

        elif dataset_type_from_db == 'pdf':
            try:
                full_text = ""
                with fitz.open(stream=file_content_bytes, filetype="pdf") as doc:
                    diagnostics = {
                        "pages": doc.page_count,
                        "size_bytes": len(file_content_bytes)
                    }
                    for page in doc:
                        full_text += page.get_text("text") + "\n"
                previewData = full_text.strip()
            except Exception as e:
                logger.warning(f"Failed to read PDF for dataset {dataset_id}: {e}")
                return jsonify({"success": False, "error": "Failed to read PDF file."}), 400

        elif dataset_type_from_db == 'docx':
            try:
                # --- INICIO DEL BLOQUE DE DEPURACIÓN ---
                print("\n--- DEBUG DOCX: Intentando abrir el archivo .docx ---")
                doc = Document(BytesIO(file_content_bytes))
                print(f"--- DEBUG DOCX: Archivo abierto. Número de párrafos encontrados: {len(doc.paragraphs)}")

                for i, p in enumerate(doc.paragraphs):
                    print(f"--- DEBUG DOCX: Párrafo {i}: '{p.text[:100]}...'")
                # --- FIN DEL BLOQUE DE DEPURACIÓN ---

                full_text = "\n".join(p.text for p in doc.paragraphs)
                print(f"--- DEBUG DOCX: Texto completo extraído (primeros 200 chars): '{full_text[:200]}'")

                diagnostics = {
                    "paragraphs": len(doc.paragraphs),
                    "size_bytes": len(file_content_bytes)
                }
                previewData = full_text.strip()
            except Exception as e:
                logger.warning(f"Failed to read DOCX for dataset {dataset_id}: {e}")
                return jsonify({"success": False, "error": "Failed to read DOCX file."}), 400

        elif dataset_type_from_db == 'json':
            try:
                data = json.loads(file_content_bytes.decode('utf-8'))
                diagnostics = {
                    "top_level_keys": len(data.keys()) if isinstance(data, dict) else 0,
                    "size_bytes": len(file_content_bytes)
                }
                previewData = json.dumps(data, indent=2)
            except Exception as e:
                logger.warning(f"Failed to read JSON for dataset {dataset_id}: {e}")
                return jsonify({"success": False, "error": "Failed to read JSON file."}), 400

        else:
            return jsonify({"success": False, "error": f"Category '{dataset_type_from_db}' not supported."}), 400

        return jsonify({
            "success": True,
            "data": {
                "fileType": dataset_type_from_db,
                "diagnostics": diagnostics,
                "previewData": previewData,
                "storage_path": storage_path 
            }
        })

    except Exception as e:
        logger.error(f"Error in diagnose_dataset for ID {dataset_id}: {e}", exc_info=True)
        return jsonify({"success": False, "error": "An internal server error occurred."}), 500

        

@app.route('/api/datasets/<string:dataset_id>/columns-summary', methods=['GET'])
@token_required
def get_columns_summary(current_user, dataset_id):
    try:
        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        storage_path = dataset_info["data"].get("storage_path")

        df = supabase_handler.load_file_as_dataframe(user_id=current_user.id, path=storage_path)
        if df is None:
            return jsonify({"success": False, "error": "No se pudo cargar el archivo (formato no soportado o corrupto)."}), 400

        result = orquestador.get_analysis_for_visualization(df)
        if result.get("success"):
            return jsonify(result)
        else:
            return jsonify(result), 400

    except Exception as e:
        logger.error(f"Error en endpoint /columns-summary: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500
    


@app.route('/api/datasets/<string:dataset_id>/visualize', methods=['POST'])
@token_required
def get_visualization(current_user, dataset_id):
    try:
        payload = request.get_json()
        plot_type = payload.get('plot_type')
        params = payload.get('params', {})

        if not plot_type:
            return jsonify({"success": False, "error": "Falta 'plot_type'."}), 400

        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        storage_path = dataset_info["data"].get("storage_path")

        df = supabase_handler.load_file_as_dataframe(user_id=current_user.id, path=storage_path)
        if df is None:
            return jsonify({"success": False, "error": "No se pudo cargar el archivo (formato no soportado o corrupto)."}), 400

        result = generate_plot(df, plot_type, params)
        if result.get("success"):
            return jsonify(result)
        else:
            return jsonify(result), 400

    except Exception as e:
        logger.error(f"Error en /visualize: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500

        

@app.route("/api/datasets/<string:dataset_id>/column-details", methods=["GET"])
@token_required
def get_column_details_endpoint(current_user, dataset_id):
    column_name = request.args.get("column")
    if not column_name:
        return jsonify({"success": False, "error": "Falta el parámetro 'column'."}), 400

    try:
        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        storage_path = dataset_info["data"].get("storage_path")

        df = supabase_handler.load_file_as_dataframe(user_id=current_user.id, path=storage_path)
        if df is None:
            return jsonify({"success": False, "error": "No se pudo cargar el archivo (formato no soportado o corrupto)."}), 400

        # Solo delega el análisis
        result = get_column_details(df, column_name)
        status_code = 200 if result.get("success") else 400
        return jsonify(result), status_code

    except Exception as e:
        logger.error(f"Error general en /column-details para dataset {dataset_id}: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500


# --- Lista de acciones permitidas SOLO para creación de columnas ---
ALLOWED_CREATE_ACTIONS = [
    "general_create_calculated_column",
    "general_create_column_by_value"
]

@app.route('/api/datasets/<string:dataset_id>/columns/create', methods=['POST'])
@token_required
def create_column_from_production(current_user, dataset_id):
    """
    Endpoint de producción para creación de columnas.
    Solo acepta acciones definidas en ALLOWED_CREATE_ACTIONS.
    """
    try:
        # --- 1. Leer payload y validar acción ---
        payload = request.get_json()
        action = payload.get('action')
        params = payload.get('params', {})

        if action not in ALLOWED_CREATE_ACTIONS:
            return jsonify({
                "success": False,
                "error": f"Acción '{action}' no permitida en este endpoint."
            }), 400

        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        storage_path = dataset_info["data"].get("storage_path")
        

        df = supabase_handler.load_file_as_dataframe(
            user_id=current_user.id,
            path=storage_path
        )

        print(f"\n--- DEBUG: ENDPOINT /columns/create ---")
        print(f"[ENTRADA] Shape del DataFrame: {df.shape}")
        print(f"[ENTRADA] Duplicados: {df.duplicated().sum()}")

        if df is None:
            raise ValueError("No se pudo leer el archivo tabular. Podría estar corrupto o en un formato no soportado.")

        # --- 2. Ejecutar la acción de creación de columna ---
        result = orquestador.cleaning_action(df, action, params)
        if not result.get("success"):
            return jsonify(result), 400

        df_clean = result["cleaned_dataframe"]
        message = result.get("message", "Columna creada correctamente.")

        # --- 3. Guardar archivo actualizado ---
        file_ext = os.path.splitext(storage_path)[-1].lower().replace('.', '')
        buffer_out = BytesIO()

        if file_ext == 'xlsx':
            df_clean.to_excel(buffer_out, index=False)
        elif file_ext == 'parquet':
            df_clean.to_parquet(buffer_out, index=False)
        else:
            df_clean.to_csv(buffer_out, index=False)
        buffer_out.seek(0)

        supabase_admin.storage.from_("proyectos-usuarios").update(
            path=storage_path,
            file=buffer_out.getvalue(),
            file_options={
               "content-type": "application/octet-stream",
               "cache-control": "0",
               "upsert": "true"
            }
        )

        # --- 4. Diagnóstico general ---
        analysis_result = orquestador.get_preliminary_analysis(df_clean)

        if not analysis_result.get("success"):
            return jsonify({"success": False, "error": "Error generando diagnóstico."}), 500
        

        general_diagnostics = {
            "summary": analysis_result["data"]["project_summary"],
            "duplicates": analysis_result["data"]["duplicates_summary"],
            "nanValues": analysis_result["data"]["nulls_summary"],
            "columnDetails": analysis_result["data"]["nulls_summary"]["by_column"],
            "columns_info": analysis_result["data"]["columns_info"]
        }

        df_preview_temp = df_clean.head(100).copy()
        df_preview_object = df_preview_temp.astype(object)
        df_preview_filled = df_preview_object.fillna("N/A")
        preview_data = df_preview_filled.to_dict(orient="records")

        # --- 5. Respuesta final ---
        response_data = {
            "fileType": "tabular",
            "diagnostics": general_diagnostics,
            "previewData": preview_data
        }

        return jsonify({
            "success": True,
            "message": message,
            "data": response_data
        })

    except Exception as e:
        logger.error(f"Error en endpoint /columns/create: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Ocurrió un error interno."}), 500

@app.route('/api/datasets/<string:dataset_id>/columns/duplicate', methods=['POST'])
@token_required
def duplicar_columna_endpoint(current_user, dataset_id):
    """
    Endpoint de producción para duplicar una columna.
    Sigue el patrón robusto de devolver el estado completo del dataset.
    """
    try:
        # --- 1. Obtener payload ---
        payload = request.get_json()
        params = {
            "columna_original": payload.get("columna_original"),
            "nuevo_nombre_columna": payload.get("nuevo_nombre_columna")
        }

        if not params["columna_original"] or not params["nuevo_nombre_columna"]:
            return jsonify({"success": False, "error": "Parámetros 'columna_original' y 'nuevo_nombre_columna' requeridos."}), 400

        # --- 2. Cargar el DF desde Supabase ---
        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        if not dataset_info.get("success"):
            return jsonify(dataset_info), 404
        
        storage_path = dataset_info["data"].get("storage_path")
        df = supabase_handler.load_file_as_dataframe(user_id=current_user.id, path=storage_path)

        if df is None:
            return jsonify({"success": False, "error": "No se pudo cargar el dataset."}), 400

        # --- Depuración de ENTRADA ---
        print("\n--- DEBUG: ENDPOINT /columns/duplicate ---")
        print(f"[ENTRADA] Shape del DataFrame: {df.shape}")
        print(f"[ENTRADA] Duplicados: {df.duplicated().sum()}")

        # --- 3. Ejecutar la acción usando el orquestador ---
        action = "general_duplicate_column"
        result = orquestador.cleaning_action(df, action, params)
        
        if not result.get("success"):
            return jsonify(result), 400

        # --- CORRECCIÓN: definir df_modificado antes de usarlo ---
        df_modificado = result["cleaned_dataframe"]

        # --- Depuración de SALIDA ---
        print(f"[SALIDA] Shape del DataFrame modificado: {df_modificado.shape}")
        print(f"[SALIDA] Duplicados: {df_modificado.duplicated().sum()}")
        print(f"--- FIN DEBUG ---\n")

        mensaje = result.get("message", "Columna duplicada correctamente.")

        # --- 4. Guardar archivo actualizado ---
        file_ext = os.path.splitext(storage_path)[-1].lower().replace('.', '')
        buffer_out = BytesIO()

        if file_ext == 'xlsx':
            df_modificado.to_excel(buffer_out, index=False)
        elif file_ext == 'parquet':
            df_modificado.to_parquet(buffer_out, index=False)
        else:  # csv por default
            df_modificado.to_csv(buffer_out, index=False)
        buffer_out.seek(0)
        
        supabase_admin.storage.from_("proyectos-usuarios").update(
            path=storage_path,
            file=buffer_out.getvalue(),
            file_options={
                "content-type": "application/octet-stream",
                "cache-control": "0",
                "upsert": "true"
            }
        )

        # --- 5. Generar diagnóstico y preview ---
        new_diagnostics = result["new_diagnostics"]

        general_diagnostics = {
            "summary": new_diagnostics["project_summary"],
            "duplicates": new_diagnostics["duplicates_summary"],
            "nanValues": new_diagnostics["nulls_summary"],
            "columnDetails": new_diagnostics["nulls_summary"]["by_column"],
            "columns_info": new_diagnostics["columns_info"]
        }

        df_preview_temp = df_modificado.head(100).copy()
        df_preview_object = df_preview_temp.astype(object)
        df_preview_filled = df_preview_object.fillna("N/A")
        preview_data = df_preview_filled.to_dict(orient="records")

        # --- 6. Construir la respuesta final ---
        response_data = {
            "fileType": "tabular",
            "diagnostics": general_diagnostics,
            "previewData": preview_data
        }

        return jsonify({
            "success": True,
            "message": mensaje,
            "data": response_data
        }), 200

    except Exception as e:
        logger.error(f"Error fatal en endpoint /duplicate: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Ocurrió un error interno en el servidor."}), 500


# ===================================================================
#          ENDPOINT PARA CHEQUEO DE CALIDAD DEL DATASET
# ===================================================================

@app.route("/api/datasets/<string:dataset_id>/check-quality", methods=["GET"])
@token_required
def get_dataset_quality_check(current_user, dataset_id):
    """
    Realiza un chequeo de calidad sobre un dataset específico.
    Evalúa duplicados, nulos y outliers.
    """

    logger.info(f"[QUALITY CHECK] Usuario={current_user.id} | Dataset={dataset_id} | Iniciando chequeo de calidad")

    try:
        # ---------------------------------------------------------------
        # PASO 1: Obtener metadata del dataset
        # ---------------------------------------------------------------
        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        if not dataset_info or not dataset_info.get("success"):
            error_msg = dataset_info.get("error", "Dataset no encontrado o acceso denegado.")
            logger.warning(f"[QUALITY CHECK] Fallo al obtener info del dataset {dataset_id}: {error_msg}")
            return jsonify({
                "success": False,
                "error": error_msg,
                "data": {"is_clean": None, "issues": {}}
            }), 404

        storage_path = dataset_info["data"].get("storage_path")
        if not storage_path:
            logger.error(f"[QUALITY CHECK] Dataset {dataset_id} no tiene storage_path en su metadata.")
            return jsonify({
                "success": False,
                "error": "Metadata incompleta (sin ruta de archivo).",
                "data": {"is_clean": None, "issues": {}}
            }), 400

        # ---------------------------------------------------------------
        # PASO 2: Cargar el dataset como DataFrame
        # ---------------------------------------------------------------
        df = supabase_handler.load_file_as_dataframe(user_id=current_user.id, path=storage_path)
        if df is None:
            logger.warning(f"[QUALITY CHECK] No se pudo cargar el DataFrame desde {storage_path}")
            return jsonify({
                "success": False,
                "error": "Archivo no cargable o formato no soportado.",
                "data": {"is_clean": None, "issues": {}}
            }), 400

        if df.empty:
            logger.info(f"[QUALITY CHECK] Dataset {dataset_id} vacío. No hay problemas de calidad que analizar.")
            return jsonify({
                "success": True,
                "message": "Dataset vacío. No hay problemas de calidad que analizar.",
                "data": {"is_clean": True, "issues": {}}
            }), 200

        # ---------------------------------------------------------------
        # PASO 3: Ejecutar función de chequeo de calidad (con copia)
        # ---------------------------------------------------------------
        quality_result = orquestador.check_dataset_quality(df.copy())
        # Aseguramos estructura de 'data' aunque algo falle internamente
        if "data" not in quality_result:
            quality_result["data"] = {"is_clean": None, "issues": {}}

        # ---------------------------------------------------------------
        # PASO 4: Respuesta
        # ---------------------------------------------------------------
        return jsonify(quality_result), 200

    except Exception as e:
        logger.error(f"[QUALITY CHECK] Error fatal en dataset {dataset_id}: {e}", exc_info=True)
        return jsonify({
            "success": False,
            "error": "Ocurrió un error interno grave en el servidor.",
            "data": {"is_clean": None, "issues": {}}
        }), 500




@app.route('/api/models/<string:model_id>/predict', methods=['POST'])
@token_required
def predict_endpoint(current_user, model_id):
    """
    Usa un modelo entrenado para realizar predicciones sobre nuevos datos
    y guarda el resultado en el historial.
    """
    try:
        # 1. Obtener datos de entrada
        input_data = request.get_json()
        if not input_data:
            return jsonify({"success": False, "error": "No se proporcionaron datos de entrada."}), 400

        # 2. Buscar metadata del modelo
        model_record = supabase_admin.table('trained_models') \
            .select('model_storage_path') \
            .eq('id', model_id) \
            .eq('user_id', current_user.id) \
            .single() \
            .execute()

        if not model_record.data:
            return jsonify({"success": False, "error": "Modelo no encontrado o acceso denegado."}), 404

        model_storage_path = model_record.data.get('model_storage_path')
        if not model_storage_path:
            return jsonify({"success": False, "error": "Ruta de modelo inválida en metadata."}), 500

        # 3. Cargar modelo de forma segura
        model_bytes = supabase_handler.load_file(
            user_id=current_user.id,
            path=model_storage_path
        )
        if model_bytes is None:
            return jsonify({"success": False, "error": "No se pudo descargar el modelo."}), 500

        # 4. Realizar predicción
        prediction_result = prediction_service.make_prediction(
            serialized_model_bytes=model_bytes,
            input_data=input_data
        )

        if not prediction_result.get("success"):
            return jsonify(prediction_result), 400

        # 5. Guardar predicción en historial (sin cortar el flujo si falla)
        try:
            supabase_admin.table('predictions_history').insert({
                "model_id": model_id,
                "user_id": current_user.id,
                "input_data": input_data,
                "prediction_output": prediction_result.get("prediction_data", {})
            }).execute()
        except Exception as history_error:
            logging.warning(f"[Predict History] No se pudo guardar en historial: {history_error}")

        # 6. Respuesta final
        return jsonify({
            "success": True,
            "message": "Predicción generada con éxito.",
            **prediction_result
        }), 200

    except Exception as e:
        logging.error(f"[Predict Endpoint] Error inesperado: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor durante la predicción."}), 500






# ===================================================================
#          ENDPOINT PARA CHEQUEO DE CALIDAD DEL DATASET
# ===================================================================
@app.route('/api/datasets/<string:dataset_id>/analyze-problem', methods=['POST'])
@token_required
def analyze_problem_endpoint(current_user, dataset_id):
    """
    NUEVA VERSIÓN: Analiza una columna objetivo y devuelve sus características
    fundamentales para que el frontend pueda guiar al usuario, detectando correctamente
    si es numérica o categórica según el contenido real.
    """
    if request.method == 'OPTIONS':
        return jsonify({'status': 'ok'}), 200

    logger = app.logger
    logger.info(f"[ANALYZE_TARGET] Usuario={current_user.id} | Dataset={dataset_id} | Analizando columna objetivo")

    try:
        # --- 1. VALIDAR Y OBTENER LA COLUMNA OBJETIVO ---
        config = request.get_json(silent=True)
        if not config or "target_column" not in config:
            return jsonify({"success": False, "error": "Debe especificar 'target_column'."}), 400
        target_col = config["target_column"]

        # --- 2. CARGAR EL DATAFRAME COMPLETO ---
        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        storage_path = dataset_info["data"].get("storage_path")
        df = supabase_handler.load_file_as_dataframe(user_id=current_user.id, path=storage_path)

        if df is None or df.empty:
            return jsonify({"success": False, "error": "No se pudo cargar el dataset o está vacío."}), 400
        if target_col not in df.columns:
            return jsonify({"success": False, "error": f"La columna '{target_col}' no existe."}), 400

        # --- 3. DETECTAR EL TIPO REAL DE LA COLUMNA ---
        target_series = df[target_col].dropna()
        unique_count = int(target_series.nunique())
        dtype_real = "categorical"  # valor por defecto

        if not target_series.empty:
            try:
                # Detectar si todos los valores pueden ser enteros
                as_float = target_series.astype(float)
                if (as_float == as_float.astype(int)).all():
                    dtype_real = "numeric"
                else:
                    dtype_real = "numeric"  # si tiene decimales
            except Exception:
                dtype_real = "categorical"

        logger.info(f"[ANALYZE_TARGET] Columna='{target_col}', DtypeReal='{dtype_real}', UniqueCount={unique_count}")

        # --- 4. DEVOLVER LA INFORMACIÓN ---
        return jsonify({
            "success": True,
            "data": {
                "target_column": target_col,
                "dtype": str(df[target_col].dtype),  # dtype original en Pandas
                "dtype_real": dtype_real,            # tipo real detectado
                "unique_count": unique_count,
            }
        })

    except Exception as e:
        logger.error(f"[ANALYZE_TARGET] Error fatal: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno durante el análisis de la columna."}), 500



@app.route('/api/datasets/<string:dataset_id>/columns', methods=['GET'])
@token_required
def get_dataset_columns(current_user, dataset_id):
    """
    Devuelve una lista simple con los nombres y tipos de las columnas de un dataset.
    Optimizado para ser rápido y usado por selectores en el frontend.
    """
    try:
        # 1️⃣ Validar parámetros básicos
        if not dataset_id or not isinstance(dataset_id, str):
            return jsonify({"success": False, "error": "El parámetro 'dataset_id' es inválido."}), 400

        # 2️⃣ Obtener información del dataset
        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        if not dataset_info.get("success"):
            return jsonify(dataset_info), 404

        storage_path = dataset_info["data"].get("storage_path")
        if not storage_path:
            return jsonify({"success": False, "error": "Ruta de almacenamiento no encontrada."}), 400

        # 3️⃣ Cargar SOLO la primera fila para leer nombres y tipos de columnas
        try:
            df = supabase_handler.load_file_as_dataframe(
                user_id=current_user.id,
                path=storage_path,
                nrows=1  # evita cargar todo el dataset
            )
        except FileNotFoundError:
            return jsonify({"success": False, "error": "El archivo del dataset no existe."}), 404
        except Exception as load_err:
            app.logger.error(f"Error cargando cabeceras del dataset {dataset_id}: {load_err}", exc_info=True)
            return jsonify({"success": False, "error": "No se pudo leer la cabecera del archivo."}), 500

        if df is None or df.empty:
            return jsonify({"success": False, "error": "Dataset vacío o sin columnas."}), 400

        # 4️⃣ Construir respuesta optimizada para el frontend
        columns_data = []
        for col, dtype in df.dtypes.items():
            columns_data.append({
                "name": str(col),
                "type": str(dtype)
            })

        # 5️⃣ Respuesta final
        return jsonify({
            "success": True,
            "columns": columns_data
        })

    except Exception as e:
        # Log detallado del error para depuración
        app.logger.error(f"Error inesperado en get_dataset_columns para dataset {dataset_id}: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno al obtener las columnas."}), 500


# ===================================================================
#          ENDPOINT PARA PREPARAR FEATURES
# ===================================================================
@app.route('/api/datasets/<string:dataset_id>/prepare-features', methods=['POST'])
@token_required
def prepare_features_endpoint(current_user, dataset_id):
    """
    Recibe el target y el tipo de problema, y devuelve el análisis de las features,
    detectando correctamente las columnas categóricas y numéricas, y devolviendo
    la información detallada que el frontend espera.
    """
    if request.method == 'OPTIONS':
        return jsonify({'status': 'ok'}), 200

    logger = app.logger
    logger.info(f"[PREPARE_FEATURES] Usuario={current_user.id} | Dataset={dataset_id} | Preparando features")

    try:
        # --- 1. VALIDAR INPUT ---
        config = request.get_json(silent=True)
        if not config:
            return jsonify({"success": False, "error": "Faltan parámetros en la request."}), 400

        target_col = config.get("target_column")
        problem_type = config.get("problem_type")
        if not target_col or not problem_type:
            return jsonify({"success": False, "error": "Faltan 'target_column' o 'problem_type'."}), 400

        # --- 2. CARGAR DATAFRAME ---
        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        storage_path = dataset_info["data"].get("storage_path")
        df = supabase_handler.load_file_as_dataframe(user_id=current_user.id, path=storage_path)

        if df is None or df.empty:
            return jsonify({"success": False, "error": "Dataset no encontrado o vacío."}), 404
        if target_col not in df.columns:
            return jsonify({"success": False, "error": f"La columna '{target_col}' no existe."}), 400

        # --- 3. SEPARAR X e y ---
        y = df[target_col]
        X = df.drop(columns=[target_col])

        # --- 4. ANALIZAR FEATURES ---
        categorical_feature_names = X.select_dtypes(include=['object', 'category']).columns.tolist()

        detailed_categorical_features = []
        for col_name in categorical_feature_names:
            unique_count = X[col_name].nunique()
            detailed_categorical_features.append({
                "name": col_name,
                "unique_count": int(unique_count)
            })

        # --- 5. CALCULAR IMPORTANCIA DE FEATURES ---
        feature_importance = prediction_service._calculate_feature_importance(X, y, problem_type)

        # --- 6. NUEVA LÓGICA: columnas con orden predefinido ---
        predefined_ordinal_keys = [k.lower() for k in prediction_service.ORDINAL_MAPS.keys()]

        categorical_features_with_predefined_order = [
            f['name'] for f in detailed_categorical_features if f['name'].lower() in predefined_ordinal_keys
        ]

        # --- 7. RESPUESTA ---
        return jsonify({
            "success": True,
            "data": {
                "categorical_features": detailed_categorical_features,
                "feature_importance": feature_importance,
                "predefined_ordinal_columns": categorical_features_with_predefined_order
            }
        })

    except Exception as e:
        logger.error(f"[PREPARE_FEATURES] Error fatal: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno al preparar características."}), 500




@app.route('/api/datasets/<string:dataset_id>/train-model', methods=['POST'])
@token_required
def train_model_endpoint(current_user, dataset_id):
    logger = app.logger
    logger.info(f"[TRAIN_MODEL] Usuario={current_user.id} Dataset={dataset_id}")
    start_time = time.time()

    try:
        # --- PASO 1: OBTENER CONFIGURACIÓN (SIN LIMPIEZA) ---
        config = request.get_json(silent=True)
        if not config:
            return jsonify({"success": False, "error": "No se recibió configuración JSON."}), 400

        # Obtenemos los datos tal cual vienen
        target_col = config.get("target_col")
        model_name = config.get("model_name")
        project_name = config.get("project_name")
        model_display_name = config.get("model_display_name")
        problem_type_from_request = config.get("problem_type")
        encoding_strategies = config.get("encoding_strategies", {})
        use_smote = bool(config.get("use_smote", False))
        use_cv = bool(config.get("use_cv", False))
        is_experiment_only = bool(config.get("is_experiment_only", False))

        # --- PASO 2: VALIDACIONES ---
        if not target_col:
            return jsonify({"success": False, "error": "'target_col' es obligatorio."}), 400
        if not model_name:
            return jsonify({"success": False, "error": "'model_name' es obligatorio."}), 400
        if not isinstance(encoding_strategies, dict):
            return jsonify({"success": False, "error": "'encoding_strategies' debe ser un objeto JSON."}), 400
        if not is_experiment_only:
            if not project_name:
                return jsonify({"success": False, "error": "'project_name' es obligatorio."}), 400
            if not model_display_name:
                return jsonify({"success": False, "error": "'model_display_name' es obligatorio."}), 400

        # --- PASO 3: CARGAR DATASET (SIN LIMPIEZA DE COLUMNAS) ---
        dataset_info_result = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        if not dataset_info_result.get("success"):
            return jsonify(dataset_info_result), 404

        dataset_data = dataset_info_result["data"]
        project_id = dataset_data.get("project_id")
        storage_path = dataset_data.get("storage_path")

        if not project_id or not storage_path:
            return jsonify({"success": False, "error": "Metadata incompleta del dataset."}), 400

        try:
            df_fresh = supabase_handler.load_file_as_dataframe(user_id=current_user.id, path=storage_path)
        except Exception as load_err:
            logger.error(f"[TRAIN_MODEL] Error cargando dataset: {load_err}", exc_info=True)
            return jsonify({"success": False, "error": "Error cargando dataset desde storage."}), 500

        if df_fresh is None or df_fresh.empty:
            return jsonify({"success": False, "error": "Dataset vacío o no cargado."}), 400

        MAX_ROWS = 1_000_000
        if len(df_fresh) > MAX_ROWS:
            return jsonify({"success": False, "error": f"Dataset demasiado grande (máx {MAX_ROWS} filas)."}), 400

        # --- PASO 4: ENTRENAR MODELO ---
        training_response = prediction_service.train_model(
            df=df_fresh.copy(deep=True),
            target_col=target_col,
            model_name=model_name,
            encoding_strategies=encoding_strategies,
            use_smote=use_smote,
            use_cv=use_cv,
            problem_type=problem_type_from_request
        )

        if not training_response.get("success"):
            return jsonify(training_response), 400

        training_data = training_response["data"]

        # --- CASO EXPERIMENTO ---
        if is_experiment_only:
            if 'artifacts' in training_data:
                del training_data['artifacts']
            elapsed = round(time.time() - start_time, 2)
            logger.info(f"[TRAIN_MODEL] Experimento completado en {elapsed}s (sin guardar)")
            return jsonify({
                "success": True,
                "message": "Experimento completado con éxito.",
                "results": training_data,
                "elapsed_time_sec": elapsed,
                "is_experiment": True
            }), 200

        # --- PASO 5: GUARDAR MODELO EN STORAGE ---
        artifacts_to_save = training_data.pop('artifacts', None)
        if not artifacts_to_save:
            return jsonify({"success": False, "error": "No se generaron artefactos del modelo."}), 500

        buffer = BytesIO()
        joblib.dump(artifacts_to_save, buffer)
        model_bytes = buffer.getvalue()

        model_id = str(uuid.uuid4())
        model_storage_path = f"{current_user.id}/{project_id}/models/{model_id}.joblib"

        try:
            supabase_admin.storage.from_("proyectos-usuarios").upload(
                path=model_storage_path,
                file=model_bytes,
                file_options={"content-type": "application/octet-stream"}
            )
            logger.info(f"[TRAIN_MODEL] Modelo guardado en Storage: {model_storage_path}")
        except Exception as storage_error:
            logger.error(f"[TRAIN_MODEL] Error guardando modelo en storage: {storage_error}", exc_info=True)
            return jsonify({"success": False, "error": "Error al guardar el modelo."}), 500

        # --- PASO 6: GUARDAR METADATA EN BD ---
        feature_cols = [col for col in df_fresh.columns if col != target_col]

        new_model_record = {
            "id": model_id,
            "project_id": project_id,
            "user_id": current_user.id,
            "model_name": model_name,
            "source_dataset_id": dataset_id,
            "target_col": target_col,
            "feature_cols": feature_cols,
            "model_storage_path": model_storage_path,
            "evaluation_results": training_data,
            "project_name": project_name,
            "model_display_name": model_display_name
        }

        try:
            supabase_admin.table("trained_models").insert(new_model_record).execute()
            logger.info(f"[TRAIN_MODEL] Metadata del modelo {model_id} guardada en BD.")

        # --- MANEJO DE DUPLICADOS ---
        except Exception as db_error:
            # Comprobamos si es error de duplicado
            if hasattr(db_error, 'code') and db_error.code == '23505':
                logger.warning(f"[TRAIN_MODEL] Intento de guardar modelo con nombre duplicado: {model_display_name}")
                supabase_admin.storage.from_("proyectos-usuarios").remove([model_storage_path])
                return jsonify({
                    "success": False, 
                    "error": f"El nombre del modelo '{model_display_name}' ya existe en el proyecto '{project_name}'. Por favor, elige un nombre único."
                }), 409
            # Otros errores
            logger.error(f"[TRAIN_MODEL] Fallo guardando metadata BD: {db_error}", exc_info=True)
            supabase_admin.storage.from_("proyectos-usuarios").remove([model_storage_path])
            return jsonify({"success": False, "error": "Error guardando metadata en BD."}), 500

        # --- PASO 7: RESPUESTA EXITOSA ---
        elapsed = round(time.time() - start_time, 2)
        logger.info(f"[TRAIN_MODEL] Entrenamiento completado en {elapsed}s")
        return jsonify({
            "success": True,
            "message": "Modelo entrenado y guardado con éxito.",
            "model_id": model_id,
            "project_id": project_id,
            "model_display_name": model_display_name,
            "storage_path": model_storage_path,
            "results": training_data,
            "elapsed_time_sec": elapsed,
            "is_experiment": False
        }), 201

    except Exception as e:
        logger.error(f"[TRAIN_MODEL] Error fatal: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno durante el entrenamiento."}), 500

# ================================================================
#   ENDPOINT: LISTAR MODELOS ENTRENADOS DEL USUARIO ACTUAL
# ================================================================
@app.route('/api/models', methods=['GET'])
@token_required
def list_models_endpoint(current_user):
    """
    Devuelve una lista de todos los modelos entrenados por el usuario actual.
    Respuesta estandarizada para frontend.
    """
    try:
        # --- Validación de usuario ---
        if not current_user or not hasattr(current_user, "id"):
            return jsonify({"success": False, "error": "Usuario no autenticado o inválido."}), 401

        # --- Query principal ---
        models_query = (
            supabase_admin.table('trained_models')
            .select(
                'id, created_at, project_id, project_name, source_dataset_id, '
                'model_display_name, target_col, evaluation_results, problem_type'
            )
            .eq('user_id', current_user.id)
            .order('created_at', desc=True)
            .execute()
        )

        if not models_query.data:
            return jsonify({"success": True, "data": []}), 200

        models_list = []
        for model in models_query.data:
            # --- Lectura robusta del tipo de problema ---
            results = model.get('evaluation_results') or {}
            problem_type = model.get('problem_type') or results.get('problem_type')

            # --- Determinación de métrica principal ---
            main_metric_value = "N/A"
            if problem_type == 'clasificacion':
                accuracy = results.get('accuracy')
                if isinstance(accuracy, (int, float)):
                    main_metric_value = round(accuracy, 3)
            elif problem_type == 'regresion':
                r2 = results.get('r2_score')
                if isinstance(r2, (int, float)):
                    main_metric_value = round(r2, 3)
            elif problem_type == 'vision_classification':
                accuracy = results.get('metrics', {}).get('accuracy')
                if isinstance(accuracy, (int, float)):
                    main_metric_value = round(accuracy, 3)
            elif problem_type == 'clustering':
                n_clusters = results.get('n_clusters')
                if n_clusters is not None:
                    main_metric_value = f"{n_clusters} Grupos"
            elif problem_type == 'etiquetado':
                f1 = results.get('f1_score')
                if isinstance(f1, (int, float)):
                    main_metric_value = round(f1, 3)

            # --- Construcción del objeto para frontend ---
            models_list.append({
                "id": model.get('id'),
                "projectId": model.get('project_id'),
                "sourceDatasetId": model.get('source_dataset_id'),
                "createdAt": model.get('created_at'),
                "projectName": model.get('project_name'),
                "modelName": model.get('model_display_name'),
                "targetColumn": model.get('target_col'),
                "problemType": problem_type or "desconocido",
                "mainMetric": main_metric_value
            })

        return jsonify({"success": True, "data": models_list}), 200

    except Exception as e:
        logging.error(f"[list_models_endpoint] Error: {str(e)}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500



        
@app.route('/api/models/<string:model_id>/sensitivity', methods=['POST'])
@token_required
def sensitivity_analysis_endpoint(current_user, model_id):
    """
    Endpoint robusto para realizar análisis de sensibilidad de un modelo guardado.
    """
    try:
        # --- 1. Validar payload del frontend ---
        config = request.get_json()
        required_keys = ['feature_to_vary', 'variation_range', 'base_data_point']
        if not config or not all(k in config for k in required_keys):
            missing = [k for k in required_keys if k not in (config or {})]
            return jsonify({"success": False, "error": f"Faltan parámetros: {', '.join(missing)}"}), 400

        feature_to_vary = config['feature_to_vary']
        variation_range = config['variation_range']
        base_data_point = config['base_data_point']

        # --- 2. Obtener metadata del modelo ---
        model_record = supabase_admin.table('trained_models') \
            .select('model_storage_path') \
            .eq('id', model_id) \
            .eq('user_id', current_user.id) \
            .single() \
            .execute()

        if not model_record.data or 'model_storage_path' not in model_record.data:
            return jsonify({"success": False, "error": "Modelo no encontrado o sin ruta de almacenamiento."}), 404

        model_storage_path = model_record.data['model_storage_path']

        # --- 3. Cargar modelo con handler seguro ---
        model_bytes = supabase_handler.load_file(
            user_id=current_user.id,
            path=model_storage_path
        )
        if model_bytes is None:
            return jsonify({"success": False, "error": "No se pudo descargar el modelo."}), 500

        # --- 4. Ejecutar análisis de sensibilidad ---
        analysis_result = prediction_service.analyze_sensitivity(
            serialized_model_bytes=model_bytes,
            feature_to_vary=feature_to_vary,
            variation_range=variation_range,
            base_data_point=base_data_point
        )

        if not analysis_result.get("success"):
            logging.warning(f"[Sensitivity] Error en análisis: {analysis_result.get('error')}")
            return jsonify(analysis_result), 500

        # --- 5. Responder con éxito ---
        return jsonify(analysis_result), 200

    except Exception as e:
        logging.error(f"[Sensitivity Endpoint] Error inesperado: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500




@app.route('/api/models/<string:model_id>/batch-predict', methods=['POST'])
@token_required
def batch_predict_endpoint(current_user, model_id):
    """
    Endpoint robusto para predicción batch con modelo guardado.
    Lee datasets soportados (CSV, XLSX, Parquet) de forma segura usando supabase_handler.
    """
    try:
        # 1. Validar payload JSON
        payload = request.get_json()
        if not payload or 'dataset_id' not in payload:
            return jsonify({"success": False, "error": "Falta 'dataset_id' en la solicitud."}), 400
        dataset_id = payload['dataset_id']

        # 2. Descargar dataset (solo metadatos, no el archivo en bruto)
        try:
            dataset_record = supabase_admin.table('datasets') \
                .select('storage_path, dataset_type') \
                .eq('dataset_id', dataset_id) \
                .eq('user_id', current_user.id) \
                .single().execute()
        except Exception as db_err:
            logging.error(f"[Dataset DB] Error consultando dataset: {db_err}", exc_info=True)
            return jsonify({"success": False, "error": "Error al buscar el dataset."}), 500

        if not dataset_record.data:
            return jsonify({"success": False, "error": "Dataset no encontrado."}), 404

        file_storage_path = dataset_record.data['storage_path']

        # 3. Leer dataset con handler seguro
        try:
            df_to_predict = supabase_handler.load_file_as_dataframe(
                user_id=current_user.id,
                path=file_storage_path
            )
            if df_to_predict is None or df_to_predict.empty:
                return jsonify({"success": False, "error": "Archivo corrupto, inválido o vacío."}), 400
        except Exception as read_err:
            logging.error(f"[Dataset Read] Error leyendo dataset: {read_err}", exc_info=True)
            return jsonify({"success": False, "error": "No se pudo procesar el dataset."}), 400

        # 4. Descargar modelo (con handler seguro)
        try:
            model_record = supabase_admin.table('trained_models') \
                .select('model_storage_path') \
                .eq('id', model_id) \
                .eq('user_id', current_user.id) \
                .single().execute()
        except Exception as db_err:
            logging.error(f"[Model DB] Error consultando modelo: {db_err}", exc_info=True)
            return jsonify({"success": False, "error": "Error al buscar el modelo."}), 500

        if not model_record.data:
            return jsonify({"success": False, "error": "Modelo no encontrado."}), 404

        model_storage_path = model_record.data['model_storage_path']

        model_bytes = supabase_handler.load_file(
            user_id=current_user.id,
            path=model_storage_path
        )
        if model_bytes is None:
            return jsonify({"success": False, "error": "No se pudo descargar el modelo."}), 500

        # 5. Ejecutar predicción
        try:
            result = prediction_service.make_batch_prediction(model_bytes, df_to_predict)
            if not result.get("success"):
                return jsonify(result), 500
            df_with_predictions = result['data']
        except Exception as pred_err:
            logging.error(f"[Prediction] Error ejecutando predicción: {pred_err}", exc_info=True)
            return jsonify({"success": False, "error": "Error ejecutando la predicción."}), 500

        # 6. Convertir DataFrame a JSON
        try:
            results_list = df_with_predictions.to_dict(orient='records')
        except Exception as convert_err:
            logging.error(f"[Convert] Error convirtiendo DataFrame a JSON: {convert_err}", exc_info=True)
            return jsonify({"success": False, "error": "Error formateando resultados."}), 500

        return jsonify({"success": True, "data": results_list}), 200

    except Exception as e:
        logging.error(f"[Unexpected] Error inesperado: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500



# ================================================================
#   ENDPOINT: OBTENER DETALLES DE UN MODELO ESPECÍFICO
# ================================================================
@app.route('/api/models/<string:model_id>', methods=['GET'])
@token_required
def get_model_details_endpoint(current_user, model_id):
    """
    Devuelve todos los detalles de un único modelo entrenado,
    incluyendo la lista de features y su tipo inferido.
    Informa si hay valores nulos en las columnas.
    """
    try:
        # --- 1. Validación de entrada ---
        if not model_id or not isinstance(model_id, str):
            return jsonify({"success": False, "error": "ID de modelo inválido."}), 400

        # --- 2. Consultar el modelo en la base de datos ---
        try:
            model_query = (
                supabase_admin
                .table('trained_models')
                .select('*, source_dataset_id')
                .eq('id', model_id)
                .eq('user_id', current_user.id)
                .single()
                .execute()
            )
        except Exception as db_err:
            logging.error(f"[get_model_details_endpoint] Error consultando la base: {db_err}", exc_info=True)
            return jsonify({"success": False, "error": "Error al acceder a la base de datos."}), 500

        if not model_query or not model_query.data:
            return jsonify({"success": False, "error": "Modelo no encontrado o acceso denegado."}), 404

        model_data = model_query.data

        # --- 3. Inferir tipos de features desde el dataset original ---
        feature_details = []
        dataset_id = model_data.get("source_dataset_id")

        if dataset_id:
            try:
                dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
                if dataset_info.get("success"):
                    storage_path = dataset_info["data"].get("storage_path")

                    df_sample = supabase_handler.load_file_as_dataframe(
                        user_id=current_user.id,
                        path=storage_path,
                        nrows=5000
                    )

                    if df_sample is not None and not df_sample.empty:

                        # ------ INICIO: BLOQUE FINAL (VERSIÓN QUE RESPETA NULOS) ------
                        print("\n[DEBUG] INICIANDO ANÁLISIS DE CARACTERÍSTICAS - VERSIÓN FINAL (Respeta Nulos)\n" + "="*50)
                        for feature_name in model_data.get("feature_cols", []):
                            if feature_name not in df_sample.columns:
                                print(f"\n[DEBUG]   - COLUMNA '{feature_name}' NO ENCONTRADA. SALTANDO.")
                                continue

                            col_series = df_sample[feature_name]  # Serie original
                            dtype = col_series.dtype
                            has_nulls = bool(col_series.isnull().any())
                            col_type = None

                            unique_count = col_series.nunique()  # únicos sin NaN

                            print(f"\n[DEBUG] PROCESANDO: '{feature_name}'")
                            print(f"[DEBUG]   - Dtype leído por Pandas: {dtype}")
                            print(f"[DEBUG]   - Cantidad de valores únicos (sin contar NaN): {unique_count}")
                            print(f"[DEBUG]   - Contiene nulos: {has_nulls}")

                            # --- LÓGICA DE DECISIÓN (ROBUSTA Y SEGURA) ---
                            if pd.api.types.is_object_dtype(dtype) or pd.api.types.is_categorical_dtype(dtype):
                                col_type = "categorical"
                            elif pd.api.types.is_numeric_dtype(dtype):
                                col_series_no_na = col_series.dropna()
                                is_float_but_actually_int = False
                                if pd.api.types.is_float_dtype(dtype) and not col_series_no_na.empty:
                                    is_float_but_actually_int = (col_series_no_na == col_series_no_na.astype(int)).all()

                                if pd.api.types.is_integer_dtype(dtype) or is_float_but_actually_int:
                                    col_type = "categorical" if unique_count <= 20 else "numeric"
                                else:
                                    col_type = "numeric"

                            if col_type is None:
                                col_type = "categorical"

                            print(f"[DEBUG]   - TIPO DECIDIDO: '{col_type}'")

                            detail = {
                                "name": feature_name,
                                "type": col_type,
                                "hasNulls": has_nulls
                            }

                            if col_type == "categorical":
                                unique_values = col_series.dropna().unique()
                                if len(unique_values) < 50:
                                    if pd.api.types.is_numeric_dtype(unique_values.dtype):
                                        detail["options"] = sorted(map(str, unique_values.astype(int).tolist()))
                                    else:
                                        detail["options"] = sorted(map(str, unique_values.tolist()))
                                    print(f"[DEBUG]   - OPCIONES AÑADIDAS: {detail['options']}")

                            feature_details.append(detail)

                        print("\n" + "="*50 + "\n[DEBUG] FIN DEL ANÁLISIS\n")
                        # ------ FIN DEL BLOQUE FINAL ------
            except Exception as feature_err:
                logging.warning(
                    f"[get_model_details_endpoint] No se pudieron enriquecer features del dataset {dataset_id}: {feature_err}",
                    exc_info=True
                )

        # --- 4. Construir respuesta limpia para frontend ---
        evaluation_results = model_data.get("evaluation_results")
        if isinstance(evaluation_results, dict):
            problem_type = evaluation_results.get("problem_type")
        else:
            problem_type = None
            evaluation_results = {}

        final_response = {
            "id": model_data.get("id"),
            "modelName": model_data.get("model_display_name"),
            "projectName": model_data.get("project_name"),
            "projectId": model_data.get("project_id"),
            "createdAt": model_data.get("created_at"),
            "targetColumn": model_data.get("target_col"),
            "problemType": problem_type,
            "evaluationResults": evaluation_results,
            "features": feature_details
        }

        return jsonify({"success": True, "data": final_response}), 200

    except Exception as e:
        logging.error(f"[get_model_details_endpoint] Error inesperado: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500

# ================================================================
#   ENDPOINT: ELIMINAR UN MODELO ENTRENADO
# ================================================================
@app.route('/api/models/<string:model_id>', methods=['DELETE'])
@token_required
def delete_model_endpoint(current_user, model_id):
    """
    Elimina un modelo entrenado: su registro en la base de datos y su
    archivo correspondiente en el almacenamiento.
    """
    logger = app.logger
    logger.info(f"[DELETE MODEL] Petición del usuario {current_user.id} para eliminar el modelo {model_id}")

    try:
        # --- 1. Buscar el modelo para obtener su ruta de almacenamiento ---
        #    Esto también verifica que el modelo realmente pertenezca al usuario.
        model_record_query = (
            supabase_admin.table('trained_models')
            .select('model_storage_path')
            .eq('id', model_id)
            .eq('user_id', current_user.id)
            .single()  # .single() asegura que solo obtenemos un resultado
            .execute()
        )

        # Si .execute() no encuentra datos, no devuelve un error, sino que .data está vacío.
        if not model_record_query.data:
            logger.warning(f"Intento de eliminar un modelo no encontrado o no autorizado: {model_id}")
            return jsonify({"success": False, "error": "Modelo no encontrado o acceso denegado."}), 404

        model_storage_path = model_record_query.data.get('model_storage_path')

        # --- 2. Eliminar el registro de la base de datos ---
        #    Se hace primero. Si esto falla, no dejamos un archivo huérfano.
        (
            supabase_admin.table('trained_models')
            .delete()
            .eq('id', model_id)
            .eq('user_id', current_user.id)
            .execute()
        )
        logger.info(f"Registro del modelo {model_id} eliminado de la base de datos.")

        # --- 3. Eliminar el archivo del almacenamiento ---
        if model_storage_path:
            try:
                supabase_admin.storage.from_("proyectos-usuarios").remove([model_storage_path])
                logger.info(f"Archivo {model_storage_path} eliminado del almacenamiento.")
            except Exception as storage_error:
                # Si la eliminación del archivo falla, solo lo registramos. La base de datos
                # es la fuente de verdad, y ya está limpia.
                logger.warning(f"No se pudo eliminar el archivo {model_storage_path} del almacenamiento: {storage_error}")

        # --- 4. Enviar respuesta de éxito ---
        return jsonify({
            "success": True,
            "message": "Modelo eliminado con éxito."
        }), 200

    except Exception as e:
        # Captura cualquier otro error inesperado durante el proceso
        logger.error(f"[DELETE MODEL] Error inesperado al eliminar el modelo {model_id}: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Ocurrió un error interno en el servidor."}), 500

# ===================================================================
# === NUEVO ENDPOINT: /evaluate (VERSIÓN CORREGIDA)                ===
# ===================================================================
@app.route('/api/models/<string:model_id>/evaluate', methods=['POST'])
@token_required
def evaluate_model_endpoint(current_user, model_id):
    """
    Endpoint para evaluar un modelo guardado contra un nuevo dataset.
    Devuelve un reporte de métricas en formato JSON.
    """
    try:
        # 1. Validar payload
        payload = request.get_json()
        if not payload or 'dataset_id' not in payload:
            return jsonify({"success": False, "error": "Falta 'dataset_id' en la solicitud."}), 400
        dataset_id = payload['dataset_id']

        # 2. Consultar dataset (solo metadatos)
        try:
            dataset_record = supabase_admin.table('datasets') \
                .select('storage_path, dataset_type') \
                .eq('dataset_id', dataset_id) \
                .eq('user_id', current_user.id) \
                .single().execute()
            
            if not dataset_record.data:
                return jsonify({"success": False, "error": "Dataset no encontrado o no autorizado."}), 404

            file_storage_path = dataset_record.data.get("storage_path")

            # --- CORRECCIÓN: lectura segura con handler ---
            df_to_evaluate = supabase_handler.load_file_as_dataframe(
                user_id=current_user.id,
                path=file_storage_path
            )
            if df_to_evaluate is None or df_to_evaluate.empty:
                return jsonify({"success": False, "error": "Archivo corrupto, inválido o vacío."}), 400

        except Exception as db_err:
            logging.error(f"[Dataset] Error procesando dataset: {db_err}", exc_info=True)
            return jsonify({"success": False, "error": "Error al buscar o procesar el dataset."}), 500

        # 3. Consultar y descargar modelo (con handler seguro)
        try:
            model_record = supabase_admin.table('trained_models') \
                .select('model_storage_path') \
                .eq('id', model_id) \
                .eq('user_id', current_user.id) \
                .single().execute()

            if not model_record.data:
                return jsonify({"success": False, "error": "Modelo no encontrado o no autorizado."}), 404

            model_storage_path = model_record.data.get("model_storage_path")

            # --- CORRECCIÓN: lectura segura ---
            model_bytes = supabase_handler.load_file(
                user_id=current_user.id,
                path=model_storage_path
            )
            if model_bytes is None:
                return jsonify({"success": False, "error": "No se pudo descargar el modelo."}), 500

        except Exception as db_err:
            logging.error(f"[Model] Error procesando modelo: {db_err}", exc_info=True)
            return jsonify({"success": False, "error": "Error al buscar o descargar el modelo."}), 500

        # 4. Ejecutar evaluación
        try:
            result = prediction_service.evaluate_model(model_bytes, df_to_evaluate)
            if not result.get("success"):
                return jsonify(result), 400
        except Exception as eval_err:
            logging.error(f"[Evaluate] Error ejecutando evaluación: {eval_err}", exc_info=True)
            return jsonify({"success": False, "error": "Error ejecutando la evaluación del modelo."}), 500

        # 5. Respuesta final
        return jsonify(result), 200

    except Exception as e:
        logging.error(f"[Unexpected] Error inesperado en /evaluate: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500



# ========================================
# FUNCIÓN DE TRADUCCIÓN (mock / reemplazar)
# ========================================
def translate_text(text: str, source_lang: str, target_lang: str) -> str:
    """
    Función dummy de traducción. 
    Reemplázala con tu motor real (ej: DeepL, Google Translate, HuggingFace).
    """
    return f"[{source_lang}->{target_lang}] {text}"

# ========================================
# ENDPOINT ROBUSTO
# ========================================
@app.route('/api/translate/es-en', methods=['POST'])
def translate_endpoint():
    """
    Endpoint para traducir texto de español a inglés.
    Espera JSON con:
        {
            "texto": "Texto en español a traducir"
        }
    """
    try:
        data = request.get_json()
        if not data or 'texto' not in data:
            return jsonify({"success": False, "error": "Falta la clave 'texto' en la solicitud."}), 400

        texto_es = data['texto']
        if not isinstance(texto_es, str) or texto_es.strip() == "":
            return jsonify({"success": False, "error": "El campo 'texto' debe ser un string no vacío."}), 400

        # --- Traducción con cache ---
        traduccion = translate_es_en(texto_es)

        return jsonify({
            "success": True,
            "original": texto_es,
            "translation": traduccion
        })

    except Exception as e:
        logger.exception("Error inesperado en /api/translate/es-en")
        return jsonify({"success": False, "error": "Error interno del servidor"}), 500
    


def get_system_prompt_for_cluster_labeling():
    return (
        "Eres un experto analista de datos. Tu tarea es analizar un resumen estadístico de un segmento "
        "o clúster de datos. Basándote en sus características más distintivas, debes proponer 3 "
        "posibles etiquetas o nombres cortos y descriptivos para este grupo. Las etiquetas deben ser "
        "concisas e intuitivas, capturando la esencia del segmento."
        "\n\n"
        "Devuelve únicamente las 3 etiquetas propuestas, separadas por comas, sin ninguna otra explicación."
        "\n\n"
        "Ejemplo de un resumen que podrías recibir:\n"
        "-----------------------------------------\n"
        "Estadísticas Numéricas (Promedios):\n"
        "- precio: 2500.50\n"
        "- cantidad_vendida: 1.2\n"
        "Características Categóricas (Valores más comunes):\n"
        "- categoria: 'Electrónica de Alta Gama'\n"
        "- marca: 'TechCorp'\n"
        "-----------------------------------------\n\n"
        "Ejemplo de la salida que deberías generar para el caso anterior:\n"
        "Productos Caros, Electrónica Premium, Artículos de Lujo TechCorp"
    )

def get_system_prompt_for_simplification():
    return "Eres un comunicador experto. Reescribe el siguiente texto en un lenguaje claro y simple, accesible para un público general, sin perder el significado esencial. Devuelve únicamente el texto simplificado."

def get_system_prompt_for_grammar_check():
    return "Eres un asistente de escritura experto. Corrige cualquier error de gramática, ortografía o puntuación en el siguiente texto. Devuelve únicamente el texto corregido."

def get_system_prompt_for_keyword_extraction():
    return "Analiza el siguiente texto y extrae las 5 a 10 palabras o frases clave más importantes. Devuelve solo una lista de las palabras clave, separadas por comas."

def get_system_prompt_for_title_generation():
    return "Eres un experto en redacción creativa. Genera un título breve, claro y atractivo para el siguiente texto, asegurando que resuma su esencia y sea llamativo. Devuelve únicamente el título."

def get_system_prompt_for_tone_change(tone: str):
    # --- DICCIONARIO CORREGIDO ---
    # Usamos claves en inglés para que coincida con el frontend
    # y comentarios de Python (#)
    tone_instructions = {
        "formal": "Reescribe el texto con un tono formal, utilizando un lenguaje profesional y respetuoso.",
        "informal": "Reescribe el texto con un tono informal y amigable, utilizando expresiones cotidianas y un lenguaje cercano.",
        "persuasive": "Reescribe el texto con un tono persuasivo, orientado a convencer al lector.",
        "emotional": "Reescribe el texto con un tono emocional, apelando a los sentimientos del lector.",
        "neutral": "Reescribe el texto con un tono neutral, evitando opiniones y emociones."
    }
    return tone_instructions.get(tone, tone_instructions['neutral'])

PROMPT_FUNCTIONS = {
    "simplify": get_system_prompt_for_simplification,
    "correct_grammar": get_system_prompt_for_grammar_check,
    "extract_keywords": get_system_prompt_for_keyword_extraction,
    "generate_title": get_system_prompt_for_title_generation,
    "change_tone": get_system_prompt_for_tone_change,
    "label_cluster": get_system_prompt_for_cluster_labeling, 
}

# ==========================================================
#   FUNCIÓN AUXILIAR PARA EJECUTAR EL MODELO
# ==========================================================

def _execute_llm_generation(tool_response: Dict, system_prompt: str, text_content: str) -> str:
    """
    Función robusta para manejar la lógica de generación LLM
    compatible con proveedores API (Google, OpenAI, Anthropic) y modelos locales.
    """

    # --- VALIDACIONES INICIALES ---
    if not isinstance(tool_response, dict):
        raise TypeError("tool_response debe ser un diccionario.")
    if not system_prompt or not isinstance(system_prompt, str):
        raise ValueError("system_prompt debe ser un string no vacío.")
    if not text_content or not isinstance(text_content, str):
        raise ValueError("text_content debe ser un string no vacío.")
    
    if "tool_type" not in tool_response:
        raise KeyError("tool_response debe contener 'tool_type'.")

    tool_type = tool_response["tool_type"]

    # ========================
    # CASO 1: MODELOS DE API
    # ========================
    if tool_type == "api":
        try:
            client = tool_response.get("tool_object")
            model_id = tool_response.get("model_id")
            provider = tool_response.get("provider")

            if not client:
                raise ValueError("tool_response debe contener 'tool_object' para tipo 'api'.")
            if not provider or not model_id:
                raise ValueError("tool_response debe contener 'provider' y 'model_id' para tipo 'api'.")

            provider = provider.lower()

            # --- GOOGLE GENERATIVE AI ---
            if provider == "google":
                try:
                    model = client.GenerativeModel(
                        model_id,
                        system_instruction=system_prompt  # instrucción de sistema nativa
                    )
                    response = model.generate_content(text_content)

                    if hasattr(response, 'text'):
                        return response.text
                    elif hasattr(response, 'parts'):
                        return "".join(part.text for part in response.parts)
                    else:
                        raise RuntimeError("Formato de respuesta inesperado del modelo de Google.")

                except Exception as google_err:
                    logger.error(f"Error en generación con Google: {str(google_err)}", exc_info=True)
                    raise

            # --- OPENAI ---
            elif provider == "openai":
                try:
                    response = client.chat.completions.create(
                        model=model_id,
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": text_content}
                        ]
                    )
                    return response.choices[0].message.content.strip()

                except Exception as openai_err:
                    logger.error(f"Error en generación con OpenAI: {str(openai_err)}", exc_info=True)
                    raise

            # --- ANTHROPIC ---
            elif provider == "anthropic":
                try:
                    response = client.messages.create(
                        model=model_id,
                        max_tokens=2048,
                        system=system_prompt,
                        messages=[{"role": "user", "content": text_content}]
                    )
                    if response and hasattr(response, "content"):
                        return response.content[0].text.strip()
                    else:
                        raise RuntimeError("Formato de respuesta inesperado del modelo Anthropic.")

                except Exception as anthropic_err:
                    logger.error(f"Error en generación con Anthropic: {str(anthropic_err)}", exc_info=True)
                    raise

            else:
                raise NotImplementedError(f"Proveedor API '{provider}' no soportado actualmente.")

        except Exception as e:
            logger.error(f"Error en generación API ({tool_response.get('provider')}): {str(e)}", exc_info=True)
            raise

    # ========================
    # CASO 2: MODELOS LOCALES
    # ========================
    elif tool_type == "local":
        try:
            model = tool_response.get("tool_object")
            if not model or not callable(getattr(model, "generate", None)):
                raise ValueError("El objeto del modelo local debe tener un método 'generate'.")

            # Suponemos que el método `generate` recibe el prompt completo
            prompt_final = f"{system_prompt}\n\n{text_content}"
            output = model.generate(prompt_final)

            if isinstance(output, str):
                return output.strip()
            elif hasattr(output, "text"):
                return output.text.strip()
            else:
                raise RuntimeError("Formato de salida inesperado del modelo local.")

        except Exception as e:
            logger.error(f"Error en generación local: {str(e)}", exc_info=True)
            raise

    else:
        raise ValueError(f"Tipo de herramienta '{tool_type}' no reconocido. Debe ser 'api' o 'local'.")



# ==========================================================
#   ENDPOINT DE RECETAS (GENERACIÓN DE TEXTO CON IA)
# ==========================================================

@app.route("/api/analysis/prompt-recipe", methods=["POST"])
@token_required
def execute_prompt_recipe(current_user):
    request_id = str(uuid.uuid4())
    logger.info(f"[{request_id}] Iniciando solicitud para /prompt-recipe")
    
    try:
        # --- 1. Validar input ---
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "Cuerpo de la petición JSON inválido."}), 400

        text = data.get("text", "").strip()
        recipe = data.get("recipe")
        requested_model = data.get("model_name")  # Modelo solicitado por el usuario
        options = data.get("options", {})

        if not all([text, recipe]):
            return jsonify({"success": False, "error": "Los campos 'text' y 'recipe' son obligatorios."}), 400

        if recipe not in PROMPT_FUNCTIONS:
            return jsonify({"success": False, "error": f"Receta '{recipe}' no válida."}), 400

        # --- 2. Lógica de Tiers y Selección de Modelo (MODIFICADA) ---
        user_tier = current_user.user_metadata.get("tier", "free")
        logger.info(f"[{request_id}] Usuario con tier: {user_tier}")

        # Ahora los usuarios 'free' tienen acceso a ambos modelos
        allowed_models = {
            "free": [
                #"Local: TinyLlama (1.1B)",
                "API: Google Gemini Pro"  # Se agrega para usuarios free
            ],
            "pro": [
                # Aquí puedes añadir modelos premium si los habilitas después
            ]
        }
        allowed_models["pro_plus"] = allowed_models["pro"]

        # Determinar el modelo a usar
        model_to_use = None
        user_allowed_list = allowed_models.get(user_tier, allowed_models["free"])

        if requested_model:
            if requested_model in user_allowed_list:
                model_to_use = requested_model
                logger.info(f"[{request_id}] Usuario solicitó modelo válido: {model_to_use}")
            else:
                logger.warning(f"[{request_id}] Modelo no permitido para tier '{user_tier}': {requested_model}")
                return jsonify({"success": False, "error": f"Tu plan '{user_tier}' no permite usar el modelo '{requested_model}'."}), 403
        else:
            model_to_use = user_allowed_list[0]  # Por defecto, el primero
            logger.info(f"[{request_id}] Usando modelo por defecto para tier '{user_tier}': {model_to_use}")

        # --- 3. Obtener herramienta del ModelManager ---
        llm_manager = get_llm_manager()
        logger.info(f"[{request_id}] Solicitando herramienta para modelo '{model_to_use}'")
        tool_response = llm_manager.get_tool(model_to_use)

        if not tool_response.get("success"):
            error_msg = tool_response.get('error', 'Modelo no disponible temporalmente.')
            logger.error(f"[{request_id}] Error obteniendo herramienta: {error_msg}")
            return jsonify({"success": False, "error": error_msg}), 503

        # --- 4. Generar prompt del sistema ---
        if recipe == "change_tone":
            tone = options.get("tone")
            if not tone:
                return jsonify({"success": False, "error": "La opción 'tone' es obligatoria para la receta 'change_tone'."}), 400
            system_prompt = PROMPT_FUNCTIONS[recipe](tone)
        else:
            system_prompt = PROMPT_FUNCTIONS[recipe]()

        # --- 5. Ejecutar modelo en un hilo separado ---
        logger.info(f"[{request_id}] Ejecutando modelo...")
        future = executor.submit(_execute_llm_generation, tool_response, system_prompt, text)

        try:
            ai_response = future.result(timeout=EXECUTION_TIMEOUT_SEC)
        except FuturesTimeout:
            logger.error(f"[{request_id}] Timeout ejecutando modelo '{model_to_use}'")
            return jsonify({"success": False, "error": "La solicitud tardó demasiado. Intente nuevamente."}), 504

        # --- 6. Respuesta exitosa ---
        response_data = {
            "success": True,
            "data": {
                "request_id": request_id,
                "recipe": recipe,
                "model_used": model_to_use,
                "ai_response": ai_response.strip()
            }
        }
        logger.info(f"[{request_id}] Solicitud completada con éxito.")
        return jsonify(response_data), 200

    except ValueError as e:
        logger.warning(f"[{request_id}] Error de valor: {e}")
        return jsonify({"success": False, "error": str(e)}), 422
    except NotImplementedError as e:
        logger.error(f"[{request_id}] Funcionalidad no implementada: {e}")
        return jsonify({"success": False, "error": str(e)}), 501
    except Exception as e:
        logger.error(f"[{request_id}] Error inesperado: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Ocurrió un error interno inesperado."}), 500

# ================================================================
#     ENDPOINTS PARA EL LABORATORIO DE PROMPTS (ROBUSTOS)
# ================================================================

# --- ENDPOINT 2: EJECUTAR UN PROMPT ---


@app.route('/api/promptlab/execute', methods=['POST'])
@token_required
def prompt_lab_execute(current_user):
    """
    Ejecuta un prompt del laboratorio usando lógica directa similar a /prompt-recipe.
    """
    request_id = str(uuid.uuid4())
    logger.info(f"[{request_id}] Iniciando solicitud directa para /promptlab/execute")

    try:
        # --- 1. Validar input ---
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "Cuerpo de la petición JSON inválido."}), 400

        # Campos requeridos
        model_name = data.get("model_name")
        system_prompt = data.get("system_prompt")
        context = data.get("context", "")
        user_prompt = data.get("prompt")
        dataset_id = data.get("dataset_id")

        missing_fields = [k for k, v in {"model_name": model_name, "system_prompt": system_prompt, "prompt": user_prompt}.items() if not v]
        if missing_fields:
            return jsonify({"success": False, "error": f"Faltan campos obligatorios: {', '.join(missing_fields)}"}), 400

        # --- 2. Obtener herramienta del ModelManager ---
        llm_manager = get_llm_manager()
        tool_response = llm_manager.get_tool(model_name)
        if not tool_response.get("success"):
            error_msg = tool_response.get('error', 'Modelo no disponible temporalmente.')
            logger.error(f"[{request_id}] Error obteniendo herramienta: {error_msg}")
            return jsonify({"success": False, "error": error_msg}), 503

        # --- 3. Preparar contenido y ejecutar modelo ---
        full_user_content = f"Contexto:\n---\n{context if context else 'N/A'}\n---\n\nPregunta:\n{user_prompt}"

        future = executor.submit(
            _execute_llm_generation,
            tool_response,
            system_prompt,
            full_user_content.strip()
        )

        try:
            ai_response = future.result(timeout=EXECUTION_TIMEOUT_SEC)
        except FuturesTimeout:
            logger.error(f"[{request_id}] Timeout ejecutando modelo '{model_name}'")
            return jsonify({"success": False, "error": "La solicitud tardó demasiado. Intente nuevamente."}), 504
        except Exception as e:
            logger.error(f"[{request_id}] Error en la ejecución del LLM: {str(e)}", exc_info=True)
            return jsonify({"success": False, "error": f"Error interno del modelo: {str(e)}"}), 500

       
        # --- 5. Respuesta final ---
        response_data = {
            "success": True,
            "data": {
                "request_id": request_id,
                "model_used": model_name,
                "ai_response": ai_response.strip(),
            }
        }
        logger.info(f"[{request_id}] Solicitud completada con éxito.")
        return jsonify(response_data), 200

    except Exception as e:
        logger.error(f"[{request_id}] Error inesperado en el endpoint: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Ocurrió un error interno inesperado."}), 500


# --- ENDPOINT 3: DONAR/COMPARTIR UN PROMPT ---
@app.route('/api/promptlab/donate', methods=['POST'])
@token_required
def donate_prompt_entry(current_user):
    """
    Marca una entrada del historial para ser compartida anónimamente.
    Requiere: history_entry_id.
    """
    try:
        data = request.get_json(force=True, silent=True)
        if not data:
            return jsonify({"success": False, "error": "Cuerpo de la petición vacío"}), 400

        entry_id = data.get("history_entry_id")
        if not entry_id:
            return jsonify({"success": False, "error": "Falta el ID de la entrada del historial"}), 400

        if not hasattr(history_service, "set_sharing_status"):
            logger.critical("history_service no tiene el método 'set_sharing_status'.")
            return jsonify({"success": False, "error": "Servicio no disponible"}), 500

        # --- PROCESO ---
        result = history_service.set_sharing_status(
            entry_id=entry_id,
            share=True,
            user_id=current_user.id
        )

        status_code = 200 if result.get("success") else 403
        return jsonify(result), status_code

    except BadRequest as br_err:
        logger.error(f"Error en formato JSON: {br_err}", exc_info=True)
        return jsonify({"success": False, "error": "JSON mal formado"}), 400
    except Exception as e:
        logger.error(f"Error en donate_prompt_entry: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno al procesar la donación"}), 500

@app.route('/api/datasets/<string:dataset_id>/column/<string:column_name>', methods=['GET'])
@token_required
def get_column_content_as_text(current_user, dataset_id, column_name):
    """
    Extrae todos los valores de una columna específica de un dataset tabular
    y los devuelve como un solo string, separados por saltos de línea.
    Ahora es robusto y funciona con CSV, Parquet, Excel, etc.
    """
    try:
        # --- 1. Validaciones iniciales (Sin cambios) ---
        if not dataset_id or not isinstance(dataset_id, str):
            return jsonify({"success": False, "error": "ID del dataset inválido."}), 400
        if not column_name or not isinstance(column_name, str):
            return jsonify({"success": False, "error": "Nombre de columna inválido."}), 400

        # --- 2. Obtener la información del dataset desde Supabase (Sin cambios) ---
        dataset_info = (
            supabase_admin.table("datasets")
            .select("storage_path").eq("dataset_id", dataset_id)
            .eq("user_id", current_user.id).single().execute()
        )

        if not dataset_info or not dataset_info.data:
            return jsonify({"success": False, "error": "Dataset no encontrado o acceso denegado."}), 404

        storage_path = dataset_info.data.get("storage_path")
        if not storage_path:
            return jsonify({"success": False, "error": "El dataset no tiene una ruta de almacenamiento válida."}), 500

        # --- 3. Cargar el DataFrame usando el manejador genérico (ESTA ES LA MODIFICACIÓN) ---
        df = supabase_handler.load_file_as_dataframe(
            user_id=current_user.id,
            path=storage_path
        )

        # Verificamos si la carga fue exitosa
        if df is None:
            logging.error(f"[get_column_content_as_text] supabase_handler no pudo cargar el DataFrame desde {storage_path}")
            return jsonify({"success": False, "error": "Error al procesar el archivo del dataset. Formato no compatible o archivo corrupto."}), 500

        if df.empty:
            return jsonify({"success": False, "error": "El dataset está vacío."}), 400

        # --- 4. Validar que la columna existe (Sin cambios en la lógica) ---
        if column_name not in df.columns:
            return jsonify({"success": False, "error": f"La columna '{column_name}' no existe en el dataset."}), 404

        # --- 5. Construir el texto de la columna (Sin cambios en la lógica) ---
        column_series = df[column_name].dropna().astype(str)
        if column_series.empty:
            return jsonify({"success": False, "error": f"La columna '{column_name}' no contiene datos válidos."}), 400
        
        column_text = '\n'.join(column_series.tolist())
        
        # --- 6. Respuesta final (Sin cambios) ---
        return jsonify({"success": True, "textContent": column_text}), 200

    except Exception as e:
        logging.critical(f"[get_column_content_as_text] Error inesperado: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500


@app.route('/api/promptlab/available-models', methods=['GET'])
@token_required
def get_public_available_models(current_user):
    """
    Devuelve la lista de modelos de IA disponibles usando la función auxiliar
    get_llm_manager para asegurar que la instancia está inicializada.
    """
    user_id = getattr(current_user, 'id', 'unknown')
    request_id = f"req_{user_id}"
    logger.info(f"[{request_id}] Solicitud de modelos iniciada por usuario {user_id}")

    try:
        # --- 👇 ¡AQUÍ ESTÁ LA LÓGICA CORRECTA, USANDO TU IDEA! 👇 ---
        # 1. Obtenemos la instancia segura y única del ModelManager
        llm_manager = get_llm_manager()

        # 2. Obtenemos el diccionario de herramientas desde esa instancia
        available_tools = getattr(llm_manager, "AVAILABLE_TOOLS", {})
        if not available_tools:
            logger.warning(f"[{request_id}] No se encontraron herramientas en la instancia de ModelManager.")
            return jsonify({"success": False, "error": "No hay modelos configurados en el servidor."}), 500

        # 3. Transformamos el diccionario al formato que el frontend necesita
        #    (Esta parte se queda igual porque ya era correcta)
        available_models = []
        for key, value in available_tools.items():
            clean_name = key.replace("API:", "").replace("Local:", "").strip()
            available_models.append({
                "id": key,       # El ID es la clave exacta que necesita el backend
                "name": clean_name # El nombre es lo que ve el usuario
            })
        
        # --------------------------------------------------------

        logger.info(f"[{request_id}] Modelos devueltos desde ModelManager: {len(available_models)}")
        return jsonify({"success": True, "models": available_models}), 200

    except Exception as e:
        logger.error(f"[{request_id}] Error en get_public_available_models: {e}", exc_info=True)
        return jsonify({
            "success": False,
            "error": "No se pudo obtener la lista de modelos."
        }), 500


@app.route('/api/promptlab/text-models', methods=['GET']) # <-- Nombre explícito
@token_required
def get_available_text_models_for_promptlab(current_user):
    """
    Devuelve la lista de modelos de IA de TEXTO (LLMs) disponibles
    para ser usados en el Laboratorio de Prompts.
    """
    user_id = getattr(current_user, 'id', 'unknown')
    request_id = f"req_{user_id}"
    logger.info(f"[{request_id}] Solicitud de modelos iniciada por usuario {user_id}")

    try:
        # 1. Obtenemos la instancia única y segura del LLM Manager
        llm_manager = get_llm_manager()
        available_tools = getattr(llm_manager, "AVAILABLE_TOOLS", {})

        if not available_tools:
            logger.warning(f"[{request_id}] No se encontraron herramientas en la instancia de ModelManager.")
            return jsonify({
                "success": False,
                "error": "No hay modelos configurados en el servidor."
            }), 500

        # 2. Filtro inteligente: solo incluimos modelos de TEXTO
        text_models = []
        for display_name, model_id in available_tools.items():

            # Verificamos prefijos válidos
            if display_name.startswith("API:") or display_name.startswith("Local:"):

                # Excluir explícitamente modelos de visión u otros
                if any(exclusion in display_name for exclusion in [
                    "Image Classifier", "Object Detector", "Image Clustering"
                ]):
                    continue

                clean_name = display_name.replace("API:", "").replace("Local:", "").strip()
                text_models.append({
                    "id": display_name,   # Clave completa (backend la usa como ID)
                    "name": clean_name    # Nombre limpio para el frontend
                })

        logger.info(f"[{request_id}] Modelos de TEXTO devueltos: {len(text_models)}")
        return jsonify({"success": True, "models": text_models}), 200

    except Exception as e:
        logger.error(f"[{request_id}] Error en get_public_available_models: {e}", exc_info=True)
        return jsonify({
            "success": False,
            "error": "No se pudo obtener la lista de modelos."
        }), 500








@app.route('/api/promptlab/history', methods=['GET'])
@token_required
def get_user_history(current_user):
    """
    Recupera el historial de pruebas del usuario actual con paginación segura.
    Query params: /api/promptlab/history?page=1&limit=10
    """
    request_id = f"req_{current_user.id}"
    logger.info(f"[{request_id}] Solicitud de historial iniciada por usuario {current_user.id}")

    try:
        # --- 1. Obtener y validar parámetros de paginación ---
        try:
            page = int(request.args.get('page', 1))
            limit = int(request.args.get('limit', 20))
            if page < 1:
                page = 1
            if limit < 1 or limit > 100:
                limit = 20
        except ValueError:
            logger.warning(f"[{request_id}] Parámetros de paginación inválidos, usando valores por defecto.")
            page, limit = 1, 20

        # --- 2. Obtener historial desde el servicio ---
        result = history_service.get_history_by_user(
            user_id=current_user.id,
            page=page,
            limit=limit
        )

        if not result.get("success"):
            logger.warning(f"[{request_id}] Error al recuperar historial: {result.get('error')}")
            return jsonify(result), 400

        # --- 3. Respuesta exitosa ---
        logger.info(f"[{request_id}] Historial devuelto con éxito. Items: {len(result.get('data', []))}")
        return jsonify(result), 200

    except Exception as e:
        logger.error(f"[{request_id}] Error crítico al obtener historial: {e}", exc_info=True)
        return jsonify({
            "success": False,
            "error": "Error interno al obtener el historial. Contacta soporte si persiste."
        }), 500



# ================================================================
#   ENDPOINT: GUARDAR HISTORIAL DEL LABORATORIO DE PROMPTS
# ================================================================
@app.route('/api/promptlab/save-history', methods=['POST'])
@token_required
def prompt_lab_save_history(current_user):
    """
    Guarda una ejecución del laboratorio de prompts en el historial a petición del usuario.
    """
    request_id = str(uuid.uuid4())
    logger.info(f"[{request_id}] Nueva solicitud para guardar en historial iniciada.")

    try:
        # ------------------- VALIDACIÓN DEL CUERPO -------------------
        if not request.is_json:
            logger.warning(f"[{request_id}] La petición no contiene JSON válido.")
            return jsonify({
                "success": False,
                "error": "La petición debe contener un cuerpo JSON válido."
            }), 400

        data = request.get_json(silent=True)
        if not data:
            logger.warning(f"[{request_id}] JSON vacío o malformado recibido.")
            return jsonify({
                "success": False,
                "error": "Cuerpo de la petición JSON inválido o vacío."
            }), 400

        # ------------------- PREPARAR PAYLOAD -------------------
        history_payload = {
            "user_id": current_user.id,
            "nombre_del_modelo": data.get("model_name"),
            "indicador_del_sistema": data.get("system_prompt"),
            "contexto_de_datos": data.get("context"),
            "pregunta_de_usuario": data.get("user_prompt"),
            "respuesta_ai": data.get("ai_response"),
            "titulo_personalizado": data.get("titulo_personalizado"),
            "project_id": data.get("project_id"),
            "dataset_id": data.get("dataset_id")
        }

        # ------------------- VALIDAR CAMPOS OBLIGATORIOS -------------------
        required_fields = ["nombre_del_modelo", "indicador_del_sistema", "pregunta_de_usuario", "respuesta_ai"]
        missing_fields = [field for field in required_fields if not history_payload.get(field)]

        if missing_fields:
            logger.warning(f"[{request_id}] Campos faltantes: {missing_fields}")
            return jsonify({
                "success": False,
                "error": f"Faltan campos obligatorios: {', '.join(missing_fields)}"
            }), 400

        # ------------------- GUARDAR EN EL SERVICIO -------------------
        history_result = history_service.add_entry(history_payload)

        if history_result and history_result.get("success"):
            new_history_id = history_result["data"]["id"]
            logger.info(f"[{request_id}] Historial guardado exitosamente con ID: {new_history_id}")
            return jsonify({
                "success": True,
                "message": "Prompt guardado con éxito.",
                "history_id": new_history_id
            }), 201

        else:
            error_msg = history_result.get("error", "Error desconocido al guardar.")
            logger.error(f"[{request_id}] Falló el guardado en historial: {error_msg}")
            return jsonify({
                "success": False,
                "error": error_msg
            }), 500

    except ValueError as ve:
        # Error específico de validación/conversión
        logger.error(f"[{request_id}] Error de validación: {ve}", exc_info=True)
        return jsonify({"success": False, "error": str(ve)}), 400

    except Exception as e:
        # Error inesperado
        logger.error(f"[{request_id}] Error inesperado al guardar historial: {e}", exc_info=True)
        return jsonify({
            "success": False,
            "error": "Error interno del servidor."
        }), 500

# ================================================================
#   ENDPOINT: ELIMINAR ENTRADA DEL HISTORIAL
# ================================================================
@app.route('/api/promptlab/history/<int:entry_id>', methods=['DELETE'])
@token_required
def delete_history_entry(current_user, entry_id):
    """
    Elimina una entrada específica del historial de un usuario.
    """
    try:
        logger.info(f"Usuario {current_user.id} solicita eliminar la entrada de historial {entry_id}")

        # Llamamos a la función que ya tienes en tu servicio
        result = history_service.delete_entry(
            entry_id=entry_id,
            user_id=current_user.id
        )

        if result.get("success"):
            return jsonify({
                "success": True,
                "message": "Entrada eliminada con éxito."
            }), 200
        else:
            # El servicio ya devuelve un error entendible
            return jsonify({
                "success": False,
                "error": "La entrada no fue encontrada o no tienes permiso para eliminarla."
            }), 404

    except Exception as e:
        logger.error(f"Error crítico al eliminar la entrada {entry_id}: {e}", exc_info=True)
        return jsonify({
            "success": False,
            "error": "Ocurrió un error interno al intentar eliminar la entrada."
        }), 500


# ================================================================
#   ENDPOINT: ACTUALIZAR ENTRADA DEL HISTORIAL
# ================================================================
@app.route('/api/promptlab/history/<int:entry_id>', methods=['PATCH'])
@token_required
def update_history_entry(current_user, entry_id):
    """
    Actualiza una entrada del historial (ej: título personalizado).
    """
    try:
        data = request.get_json()
        if not data or "titulo_personalizado" not in data:
            return jsonify({"success": False, "error": "Cuerpo de la petición JSON inválido o incompleto."}), 400

        nuevo_titulo = data.get("titulo_personalizado", "").strip()
        if not nuevo_titulo:
            return jsonify({"success": False, "error": "El campo 'titulo_personalizado' no puede estar vacío."}), 400

        logger.info(f"Usuario {current_user.id} solicita actualizar la entrada {entry_id}")

        result = history_service.update_entry(
            entry_id=entry_id,
            user_id=current_user.id,
            update_data={"titulo_personalizado": nuevo_titulo}
        )

        if result.get("success"):
            return jsonify({
                "success": True,
                "message": "La entrada ha sido actualizada.",
                "data": result["data"]
            }), 200
        else:
            return jsonify({
                "success": False,
                "error": "La entrada no fue encontrada o no tienes permiso para editarla."
            }), 404

    except Exception as e:
        logger.error(f"Error crítico al actualizar la entrada {entry_id}: {e}", exc_info=True)
        return jsonify({
            "success": False,
            "error": "Ocurrió un error interno."
        }), 500



# ================================================================
#   ENDPOINT: EDITAR COLUMNAS (ELIMINAR O RENOMBRAR)
# ================================================================
@app.route('/api/datasets/<string:dataset_id>/edit-columns', methods=['POST'])
@token_required
def edit_columns_history(current_user, dataset_id):
    """
    Aplica acciones de edición de columnas (eliminar o renombrar) a un dataset.
    """
    try:
        # --- 1. Validación inicial del payload ---
        payload = request.get_json()
        action = payload.get('action')
        params = payload.get('params', {})

        if action not in ["drop", "rename"]:
            return jsonify({"error": "Acción inválida. Usa 'drop' o 'rename'."}), 400

        # --- 2. Cargar dataset ---
        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        if not dataset_info:
            return jsonify({"error": "Dataset no encontrado."}), 404

        df = dataset_service.load_dataframe(dataset_info["storage_path"])

        # --- 3. Acciones ---
        if action == "drop":
            cols_to_drop = params.get("columns", [])
            if not isinstance(cols_to_drop, list):
                return jsonify({"error": "El parámetro 'columns' debe ser una lista."}), 400

            df = df.drop(columns=cols_to_drop, errors="ignore")

        elif action == "rename":
            rename_map = params.get("columns", {})
            if not isinstance(rename_map, dict):
                return jsonify({"error": "El parámetro 'columns' debe ser un diccionario {viejo:nuevo}."}), 400

            df = df.rename(columns=rename_map)

        # --- 4. Guardar dataset actualizado ---
        dataset_service.save_dataframe(df, dataset_info["storage_path"])

        return jsonify({
            "message": f"Acción '{action}' aplicada correctamente.",
            "new_columns": list(df.columns)
        }), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500



# ================================================================
#   ENDPOINT: OBTENER UNA ÚNICA ENTRADA DEL HISTORIAL (CORREGIDO)
# ================================================================
@app.route('/api/promptlab/history/<int:entry_id>', methods=['GET'])
@token_required
def get_history_entry(current_user, entry_id):
    """
    Obtiene los detalles de una entrada de historial específica para el usuario autenticado.
    """
    try:
        logger.info(f"[HISTORY][GET] Usuario {current_user.id} solicita entrada {entry_id}")

        client = supabase_handler.get_db_client()
        
        # --- LA LLAMADA A LA BASE DE DATOS NO CAMBIA ---
        response = (
            client.table("historial_pruebas")
            .select("*")
            .eq("id", entry_id)
            .eq("user_id", current_user.id)
            .single()
            .execute()
        )

        # --- CAMBIO IMPORTANTE: LA FORMA DE VALIDAR LA RESPUESTA ---
        # Si .single() no encuentra nada, .data estará vacío.
        if not response.data:
            logger.warning(f"[HISTORY][GET] Entrada {entry_id} no encontrada para usuario {current_user.id}")
            return jsonify({"success": False, "error": "Entrada no encontrada o sin permiso."}), 404

        # --- Si llegamos aquí, todo fue un éxito ---
        logger.info(f"[HISTORY][GET] Entrada {entry_id} entregada a usuario {current_user.id}")
        return jsonify({"success": True, "data": response.data}), 200

    # Si .single() falla por un error de API, lanzará una excepción que será capturada aquí.
    except Exception as e:
        logger.exception(f"[HISTORY][GET] Error crítico al obtener entrada {entry_id}: {str(e)}")
        # Ahora el mensaje de error será más informativo, mostrando la excepción real.
        return jsonify({"success": False, "error": f"Error interno del servidor: {str(e)}"}), 500



# ================================================================
#   ENDPOINT: EJECUTAR BATCH DE PROMPTS (VERSIÓN FINAL Y FUNCIONAL)
# ================================================================
@app.route('/api/promptlab/execute-batch', methods=['POST'])
@token_required
def prompt_lab_execute_batch(current_user):
    """
    Ejecuta un lote de prompts en una sola llamada al LLM, 
    y divide la respuesta en partes para cada prompt.
    """
    request_id_batch = str(uuid.uuid4())
    logger.info(f"[{request_id_batch}] Iniciando solicitud de LOTE (modo eficiente) para /promptlab/execute-batch")

    try:
        # --------------------------
        # 1. VALIDACIÓN DEL PAYLOAD
        # --------------------------
        data = request.get_json(silent=True)
        if not data:
            return jsonify({"success": False, "error": "El cuerpo de la petición debe ser un JSON válido."}), 400

        model_name = data.get("model_name")
        system_prompt = data.get("system_prompt")
        context = data.get("context", "")
        dataset_id = data.get("dataset_id")
        user_prompts = data.get("prompts")

        if not model_name or not isinstance(user_prompts, list) or not user_prompts:
            return jsonify({"success": False, "error": "Campos 'model_name' y 'prompts' (lista) son obligatorios."}), 400

        # --------------------------
        # 2. OBTENER HERRAMIENTA
        # --------------------------
        llm_manager = get_llm_manager()
        tool_response = llm_manager.get_tool(model_name)
        if not tool_response.get("success"):
            error_msg = tool_response.get('error', 'Modelo no disponible temporalmente.')
            logger.error(f"[{request_id_batch}] Error obteniendo herramienta: {error_msg}")
            return jsonify({"success": False, "error": error_msg}), 503

        # --------------------------
        # 3. EJECUCIÓN ÚNICA
        # --------------------------
        started = time.time()

        # Preparamos un input combinado con todas las preguntas
        prompts_combinados = "\n".join(user_prompts)
        full_user_content = (
            f"Contexto:\n---\n{context if context else 'N/A'}\n---\n\n"
            f"Preguntas (responde cada una en orden, separando claramente):\n{prompts_combinados}"
        )

        try:
            # Llamada única al modelo
            future = executor.submit(
                _execute_llm_generation,
                tool_response,
                system_prompt,
                full_user_content.strip()
            )
            respuesta_completa_ia = future.result(timeout=EXECUTION_TIMEOUT_SEC)
            latency_ms = int((time.time() - started) * 1000)

            # --- División de la respuesta ---
            # Ajustá el regex según el formato de salida de tu IA
            partes_respuesta = re.split(
                r'\*\*\d+\..*?\*\*',
                respuesta_completa_ia
            )
            partes_respuesta_limpias = [
                p.strip() for p in partes_respuesta if p.strip() and p.strip() not in ["Imagen", "Prompt"]
            ]

            resultados = []
            for idx, user_prompt in enumerate(user_prompts):
                if idx < len(partes_respuesta_limpias):
                    resultados.append({
                        "success": True,
                        "ai_response": partes_respuesta_limpias[idx],
                        "user_question": user_prompt,
                        "index": idx,
                        "meta": {
                            "latency_ms": latency_ms if idx == 0 else 0,  # latencia solo en el primero
                            "model_id": model_name
                        }
                    })
                else:
                    resultados.append({
                        "success": False,
                        "error": "La IA no generó una respuesta para este prompt.",
                        "user_question": user_prompt,
                        "index": idx
                    })

        except Exception as e:
            logger.error(f"[{request_id_batch}] Error en la ejecución única del lote: {e}", exc_info=True)
            resultados = [
                {"success": False, "error": str(e), "user_question": p, "index": i}
                for i, p in enumerate(user_prompts)
            ]
            latency_ms = 0

        # --------------------------
        # 4. CONSTRUIR RESUMEN
        # --------------------------
        num_exitosos = sum(1 for r in resultados if r.get("success"))
        resumen_final = {
            "total_prompts": len(user_prompts),
            "exitosos": num_exitosos,
            "fallidos": len(user_prompts) - num_exitosos,
            "latencia_promedio_ms": latency_ms if num_exitosos > 0 else 0
        }

        common_data_for_saving = {
            "model_name": model_name,
            "system_prompt": system_prompt,
            "context": context,
            "dataset_id": dataset_id
        }

        return jsonify({
            "success": True,
            "data": {
                "common_data": common_data_for_saving,
                "resultados": resultados,
                "resumen": resumen_final
            }
        }), 200

    except Exception as e:
        logger.error(f"[{request_id_batch}] Error inesperado en execute-batch: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Ocurrió un error interno al procesar la solicitud."}), 500


@app.route('/api/faq', methods=['POST'])
@token_required
def handle_faq_query(current_user): # <-- ¡Ahora acepta el usuario!
    # Puedes ignorar la variable 'current_user' si no la necesitas por ahora.
    # Simplemente tiene que estar en la firma de la función.
    
    data = request.get_json()
    # ... el resto de tu código no cambia ...
    pregunta_usuario = data.get('pregunta')
    modulo = data.get('modulo')

    if not pregunta_usuario:
        return jsonify({"success": False, "error": "El campo 'pregunta' es obligatorio."}), 400

    resultado = faq_service.responder_duda(pregunta_usuario, modulo)

    if "error" in resultado:
        return jsonify({"success": False, "error": resultado["error"]}), 500
    
    return jsonify({"success": True, "data": resultado}), 200


# --- INICIO: FUNCIONES AUXILIARES GENERALES ---

# Estas funciones son para la subida de archivos TABULARES
TABULAR_ALLOWED_EXTENSIONS = {"csv", "xlsx", "xls"}
def is_allowed_tabular_file(filename):
    """Valida extensión de archivo TABULAR."""
    return "." in filename and filename.rsplit(".", 1)[1].lower() in TABULAR_ALLOWED_EXTENSIONS

# Esta función es genérica y está bien aquí
def generate_unique_filename(filename):
    """Genera nombre único para evitar colisiones."""
    ext = filename.rsplit(".", 1)[1].lower()
    return f"{uuid.uuid4().hex}.{ext}"

# --- FIN: FUNCIONES AUXILIARES GENERALES ---


# =========================================================
# ENDPOINT 1: ENRIQUECER DATASET CON URLS DE IMÁGENES
# =========================================================

@app.route('/api/datasets/<string:dataset_id>/enrich-with-image-urls', methods=['POST'])
@token_required
def enrich_dataset_from_image_urls(current_user, dataset_id):
    """
    Lee un dataset tabular, toma URLs de una columna, analiza las imágenes y crea un nuevo dataset enriquecido.
    """
    data = request.get_json(silent=True) or {}
    column_name = data.get('column_name')
    if not column_name:
        return jsonify({"success": False, "error": "Falta el nombre de la columna ('column_name')."}), 400

    try:
        # --- PASO 0: OBTENER EL MODEL MANAGER ---
        model_manager = get_llm_manager()
        if not model_manager:
            logger.error("❌ No se pudo inicializar el model_manager.")
            return jsonify({"success": False, "error": "No se pudo inicializar el motor de análisis de imágenes."}), 500

        # --- PASO 1: Cargar el dataset original ---
        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        if not dataset_info or not dataset_info.get("success") or not dataset_info.get("data"):
            return jsonify({"success": False, "error": "Dataset no encontrado o sin permisos."}), 404

        storage_path_original = dataset_info["data"].get("storage_path")
        if not storage_path_original:
            return jsonify({"success": False, "error": "Dataset inválido, sin ruta de almacenamiento."}), 400

        original_df = supabase_handler.load_file_as_dataframe(user_id=current_user.id, path=storage_path_original)
        if original_df is None:
            return jsonify({"success": False, "error": "No se pudo cargar el dataset original."}), 500
        if column_name not in original_df.columns:
            return jsonify({"success": False, "error": f"La columna '{column_name}' no existe en el dataset."}), 400

        # --- PASO 2: Procesar cada URL ---
        results = []
        for index, row in original_df.iterrows():
            image_url = row[column_name]
            if pd.isna(image_url) or not isinstance(image_url, str) or not image_url.startswith("http"):
                results.append({})
                continue

            try:
                image_bytes = download_image_from_url(image_url)
                if image_bytes:
                    result = procesar_imagen_completa(
                        image_bytes,
                        image_url,
                        model_manager_instance=model_manager
                    )
                    results.append(result)
                else:
                    results.append({"error": "Fallo en la descarga"})
            except Exception as img_err:
                logger.warning(f"⚠️ Error procesando imagen {image_url}: {img_err}", exc_info=True)
                results.append({"error": f"Error procesando imagen: {str(img_err)}"})

        # --- PASO 3: Crear el nuevo DataFrame enriquecido ---
        results_df = pd.DataFrame(results).add_prefix("img_")
        original_df.reset_index(drop=True, inplace=True)
        results_df.reset_index(drop=True, inplace=True)
        enriched_df = pd.concat([original_df, results_df], axis=1)

        # --- PASO 4: Generar bytes CSV y nombre del archivo ---
        csv_bytes = enriched_df.to_csv(index=False).encode("utf-8")
        new_dataset_id = str(uuid.uuid4())
        original_name = dataset_info["data"].get("dataset_name", "dataset")
        file_name = f"enriched_{original_name.replace('.csv', '')}_{new_dataset_id[:8]}.csv"

        # --- PASO 5: Guardar en Supabase ---
        storage_path, _ = supabase_handler.save_file(
            file_bytes=csv_bytes,
            user_id=current_user.id,
            project_id=dataset_info["data"].get("project_id"),
            folder="datasets",
            filename=file_name
        )
        if not storage_path:
            raise Exception("No se pudo guardar el archivo enriquecido en Supabase Storage.")

        # --- PASO 6: Registrar el nuevo dataset ---
        new_dataset = dataset_service.create_dataset_record(
            dataset_id=new_dataset_id,
            user_id=current_user.id,
            project_id=dataset_info["data"].get("project_id"),
            dataset_name=f"{original_name} (Enriquecido)",
            dataset_type="tabular",
            storage_path=storage_path,
            file_size=len(csv_bytes)
        )

        return jsonify({
            "success": True,
            "message": "Dataset enriquecido creado exitosamente.",
            "new_dataset": new_dataset
        }), 201

    except Exception as e:
        logger.critical(f"❌ Error enriqueciendo dataset {dataset_id}: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500

# =========================================================
# ENDPOINT 2: CREAR DATASET DESDE SUBIDA DE IMÁGENES
# =========================================================

# --- Configuración y helpers SOLO para este endpoint ---
MAX_IMAGE_SIZE_MB = 10
ALLOWED_IMAGE_EXTENSIONS = {'png', 'jpg', 'jpeg', 'webp'}

def is_allowed_image_file(filename):
    """Valida extensión de archivo de IMAGEN."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_IMAGE_EXTENSIONS


@app.route('/api/project/<string:projectId>/vision/analyze-and-create-dataset', methods=['POST'])
@token_required
def analyze_images_and_create_dataset(current_user, projectId):
    """
    Recibe archivos de IMAGEN, los procesa con IA y crea un nuevo dataset tabular asociado al proyecto.
    """
    if 'images' not in request.files:
        return jsonify({"success": False, "error": "No se encontraron archivos de imagen."}), 400

    files = request.files.getlist('images')
    if not files or all(f.filename == '' for f in files):
        return jsonify({"success": False, "error": "No se seleccionó ningún archivo."}), 400

    results = []
    processing_errors = []

    # --- OBTÉN EL MODEL MANAGER ANTES DEL BUCLE ---
    model_manager = get_llm_manager()
    if not model_manager:
        logger.error("❌ No se pudo inicializar el model_manager.")
        return jsonify({"success": False, "error": "No se pudo inicializar el motor de análisis de imágenes."}), 500

    for file in files:
        filename = secure_filename(file.filename)

        if not is_allowed_image_file(filename):
            processing_errors.append({"nombre_archivo": filename, "error": "Tipo de archivo no permitido."})
            continue

        file.seek(0, os.SEEK_END)
        file_size_mb = file.tell() / (1024 * 1024)
        file.seek(0)

        if file_size_mb > MAX_IMAGE_SIZE_MB:
            processing_errors.append({
                "nombre_archivo": filename,
                "error": f"El archivo excede los {MAX_IMAGE_SIZE_MB} MB."
            })
            continue

        try:
            # --- LLAMADA CORREGIDA ---
            result = procesar_imagen_completa(
                file.read(),
                filename,
                model_manager_instance=model_manager
            )

            if "error" in result or "detection_error" in result:
                processing_errors.append({**result, "nombre_archivo": filename})
            else:
                results.append(result)

        except Exception as img_err:
            logger.warning(f"⚠️ Error procesando imagen {filename}: {img_err}", exc_info=True)
            processing_errors.append({"nombre_archivo": filename, "error": str(img_err)})

    if not results:
        return jsonify({
            "success": False,
            "error": "Ninguna imagen pudo ser procesada.",
            "details": processing_errors
        }), 500

    try:
        # --- Convertir resultados a CSV ---
        df = pd.DataFrame(results).fillna('N/A')
        csv_bytes = df.to_csv(index=False).encode('utf-8')

        # --- Generar identificadores ---
        new_dataset_id = str(uuid.uuid4())
        file_name = f"vision_dataset_{new_dataset_id[:8]}.csv"

        # --- Guardar en Supabase ---
        storage_path, _ = supabase_handler.save_file(
            file_bytes=csv_bytes,
            user_id=current_user.id,
            project_id=projectId,
            folder='datasets',
            filename=file_name
        )
        if not storage_path:
            raise Exception("No se pudo guardar el archivo en Supabase Storage.")

        # --- Crear registro en la tabla datasets ---
        dataset_name = f"Análisis de {len(results)} Imágenes ({pd.Timestamp.now().strftime('%Y-%m-%d')})"

        new_dataset = dataset_service.create_dataset_record(
            dataset_id=new_dataset_id,
            user_id=current_user.id,
            project_id=projectId,
            dataset_name=dataset_name,
            dataset_type='tabular',
            storage_path=storage_path,
            file_size=len(csv_bytes)
        )

        return jsonify({
            "success": True,
            "message": "Nuevo dataset creado con éxito.",
            "new_dataset": new_dataset,
            "processing_errors": processing_errors
        }), 201

    except Exception as e:
        logger.critical(f"❌ Error en guardado de Supabase para user={current_user.id}: {e}", exc_info=True)
        return jsonify({"success": False, "error": "No se pudo guardar el nuevo dataset."}), 500


# ===================================================================
#          ENDPOINT PARA GESTIÓN DE LA CUENTA DEL USUARIO
# ===================================================================


@app.route("/api/user/me", methods=["GET"])
@token_required
def get_user_profile(current_user):
    """
    Devuelve la información del perfil del usuario actualmente logueado.
    """
    try:
        user_id = current_user.id
        user_email = current_user.email

        # Hacemos la consulta a tu tabla 'profiles'
        # Seleccionamos solo 'full_name' porque es lo que tienes
        response = supabase_admin.table('profiles').select('full_name').eq('id', user_id).single().execute()
        
        profile_data = response.data
        if not profile_data:
            return jsonify({"success": False, "error": "Perfil de usuario no encontrado."}), 404

        # Combinamos la información y la devolvemos
        user_full_data = {
            "id": user_id,
            "email": user_email,
            "full_name": profile_data.get('full_name') 
            # Quitamos 'username' ya que no está en tu tabla
        }

        return jsonify({"success": True, "data": user_full_data}), 200

    except Exception as e:
        logger.error(f"Error en /api/user/me: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Ocurrió un error interno al obtener el perfil."}), 500


@app.route("/api/user/me", methods=["PUT"])
@token_required
def update_user_profile(current_user):
    """
    Actualiza la información del perfil del usuario actualmente logueado.
    """
    try:
        # 1. Obtener los datos que envía el frontend
        data_to_update = request.get_json()
        if not data_to_update:
            return jsonify({"success": False, "error": "No se enviaron datos."}), 400

        # 2. Extraer el campo que queremos actualizar
        new_full_name = data_to_update.get('full_name')

        # 3. (Opcional pero recomendado) Validar los datos
        if new_full_name is not None and len(new_full_name.strip()) < 3:
             return jsonify({"success": False, "error": "El nombre completo debe tener al menos 3 caracteres."}), 400

        # 4. Construir el objeto para la actualización
        update_payload = {}
        if new_full_name is not None:
            update_payload['full_name'] = new_full_name.strip()
        
        # Si no hay nada que actualizar, no hacemos nada
        if not update_payload:
            return jsonify({"success": False, "error": "No se proporcionaron campos válidos para actualizar."}), 400

        # 5. Actualizar la base de datos
        response = supabase_admin.table('profiles').update(update_payload).eq('id', current_user.id).execute()

        # 6. Comprobar el resultado y devolver los datos actualizados
        if response.data:
            updated_profile = response.data[0]
            # Devolvemos el perfil completo de nuevo
            full_user_data = {
                "id": current_user.id,
                "email": current_user.email,
                "full_name": updated_profile.get('full_name')
            }
            return jsonify({
                "success": True, 
                "message": "Perfil actualizado correctamente.",
                "data": full_user_data
            }), 200
        else:
            # Esto puede pasar si el RLS (Row Level Security) de Supabase lo impide
            return jsonify({"success": False, "error": "No se pudo actualizar el perfil."}), 500

    except Exception as e:
        logger.error(f"Error en PUT /api/user/me: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Ocurrió un error interno al actualizar el perfil."}), 500
    

@app.route("/api/user/password", methods=["PUT"])
@token_required
def update_user_password(current_user):
    """
    Actualiza la contraseña del usuario actualmente logueado.
    """
    try:
        # 1. Obtener las contraseñas del cuerpo de la petición
        payload = request.get_json()
        current_password = payload.get('currentPassword')
        new_password = payload.get('newPassword')

        if not current_password or not new_password:
            return jsonify({"success": False, "error": "Debes proporcionar la contraseña actual y la nueva."}), 400

        # 2. (Opcional pero recomendado) Validar la nueva contraseña en el backend
        if len(new_password) < 6:
            return jsonify({"success": False, "error": "La nueva contraseña debe tener al menos 6 caracteres."}), 400

        # 3. VERIFICAR LA CONTRASEÑA ACTUAL
        # Este es el paso de seguridad más importante.
        # Intentamos iniciar sesión con el email del usuario y la contraseña actual que nos dio.
        # Si funciona, la contraseña es correcta. Si falla, es incorrecta.
        try:
            # Usamos el cliente PÚBLICO para esta verificación, no el de admin
            cliente_public.auth.sign_in_with_password({
                "email": current_user.email,
                "password": current_password
            })
        except Exception as e:
            # Si sign_in_with_password falla, es porque la contraseña es incorrecta.
            logger.warning(f"Intento fallido de cambio de contraseña para el usuario {current_user.email}. Contraseña actual incorrecta.")
            return jsonify({"success": False, "error": "La contraseña actual es incorrecta."}), 403 # 403 Forbidden

        # 4. Si la contraseña actual es correcta, ACTUALIZAR A LA NUEVA CONTRASEÑA
        # Usamos el cliente de ADMIN para actualizar los datos del usuario por su ID.
        response = supabase_admin.auth.admin.update_user_by_id(
            uid=current_user.id,
            attributes={'password': new_password}
        )

        return jsonify({"success": True, "message": "Contraseña actualizada correctamente."}), 200

    except Exception as e:
        logger.error(f"Error en PUT /api/user/password: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Ocurrió un error interno al cambiar la contraseña."}), 500


# ===================================================================
#          ENDPOINTS PARA GESTIÓN DE CLAVES DE API
# ===================================================================

# -------------------------------------------------
# GET - Listar claves de usuario
# -------------------------------------------------
@app.route("/api/user/keys", methods=["GET"])
@token_required
def get_user_api_keys(current_user):
    """
    Devuelve los proveedores configurados por un usuario.
    Solo devuelve info ofuscada, nunca la clave completa.
    """
    try:
        response = supabase_admin.table("user_api_keys") \
            .select("provider, api_key") \
            .eq("user_id", current_user.id) \
            .execute()

        keys_info = []
        for key_data in response.data:
            decrypted = decrypt_text(key_data["api_key"])
            keys_info.append({
                "provider": key_data["provider"],
                "is_set": bool(decrypted),
                "display_key": f"****-****-{decrypted[-4:]}" if decrypted else "****"
            })

        return jsonify({"success": True, "data": keys_info}), 200

    except Exception as e:
        logger.error(f"Error en GET /api/user/keys: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error al obtener las claves."}), 500


# -------------------------------------------------
# PUT - Guardar o actualizar clave
# -------------------------------------------------
@app.route("/api/user/keys", methods=["PUT"])
@token_required
def save_user_api_key(current_user):
    """
    Guarda o actualiza una clave de API para un usuario.
    Espera JSON: { "provider": "google", "api_key": "sk-..." }
    """
    try:
        data = request.get_json() or {}
        provider = data.get("provider")
        api_key = data.get("api_key")

        if not provider or not api_key:
            return jsonify({"success": False, "error": "Faltan 'provider' o 'api_key'."}), 400

        if provider not in VALID_PROVIDERS:
            return jsonify({"success": False, "error": "Proveedor inválido."}), 400

        # Encriptamos la clave antes de guardarla
        encrypted_key = encrypt_text(api_key)

        # Upsert: inserta o actualiza según user_id + provider
        response = supabase_admin.table("user_api_keys").upsert({
            "user_id": current_user.id,
            "provider": provider,
            "api_key": encrypted_key
        }, on_conflict="user_id,provider").execute()

        if not response.data:
            return jsonify({"success": False, "error": "No se pudo guardar la clave de API."}), 500

        return jsonify({
            "success": True,
            "message": f"Clave de API para {provider.capitalize()} guardada con éxito."
        }), 200

    except Exception as e:
        logger.error(f"Error en PUT /api/user/keys: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error al guardar la clave."}), 500


# -------------------------------------------------
# DELETE - Revocar clave de un proveedor
# -------------------------------------------------
@app.route("/api/user/keys/<provider>", methods=["DELETE"])
@token_required
def delete_user_api_key(current_user, provider):
    """
    Elimina la clave de API de un proveedor para el usuario actual.
    """
    try:
        if provider not in VALID_PROVIDERS:
            return jsonify({"success": False, "error": "Proveedor inválido."}), 400

        response = supabase_admin.table("user_api_keys") \
            .delete() \
            .eq("user_id", current_user.id) \
            .eq("provider", provider) \
            .execute()

        if response.data == []:
            return jsonify({"success": False, "error": "No se encontró clave para este proveedor."}), 404

        return jsonify({
            "success": True,
            "message": f"Clave de API para {provider.capitalize()} eliminada con éxito."
        }), 200

    except Exception as e:
        logger.error(f"Error en DELETE /api/user/keys: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error al eliminar la clave."}), 500



def download_file_from_url(url: str):
    """
    Descarga un archivo desde cualquier URL pública.
    - Google Drive: maneja Docs, Sheets, Slides y archivos subidos (PDF, DOCX, XLSX)
    - Hugging Face / GitHub / URLs directas
    Devuelve: (bytes_del_archivo, content_type) o (None, None) si falla
    """
    try:
        # GET inicial
        resp = requests.get(url, stream=True, timeout=20)
        resp.raise_for_status()
        content_type = resp.headers.get("content-type", "")

        # --- Manejo especial de Google Drive ---
        if "drive.google.com" in url:
            # Si nos devuelve HTML en vez del archivo, buscar enlace de confirmación
            if "text/html" in content_type.lower():
                confirm_url = re.search(r'href="(\/uc\?export=download[^"]+)', resp.text)
                if confirm_url:
                    confirm_url = "https://drive.google.com" + confirm_url.group(1).replace("&amp;", "&")
                    resp = requests.get(confirm_url, stream=True, timeout=20)
                    resp.raise_for_status()
                    content_type = resp.headers.get("content-type", "")

        # --- Lectura de bytes (todo en memoria, cuidado con archivos gigantes) ---
        file_bytes = resp.content

        # DEBUG opcional
        logger.info(f"Archivo descargado desde {url} ({len(file_bytes)} bytes, content-type={content_type})")
        return file_bytes, content_type

    except requests.exceptions.RequestException as e:
        logger.warning(f"No se pudo descargar el archivo desde {url}: {e}")
        return None, None



@app.route('/api/projects/<string:project_id>/datasets/import-from-url', methods=['POST'])
@token_required
def import_dataset_from_url(current_user, project_id):
    data = request.get_json(silent=True) or {}
    file_url = data.get('fileUrl')

    if not file_url:
        return jsonify({"success": False, "error": "No se proporcionó la URL del archivo"}), 400

    try:
        final_download_url = file_url
        original_name_hint = ""

        # --- LÓGICA DE GOOGLE DRIVE ---
        if "google.com" in file_url:
            regex = r"/d/([a-zA-Z0-9_-]+)"
            match = re.search(regex, file_url)
            if not match:
                return jsonify({"success": False, "error": "URL de Google Drive no válida"}), 400

            file_id = match.group(1)

            if "/document/d/" in file_url:
                final_download_url = f"https://docs.google.com/document/d/{file_id}/export?format=docx"
                original_name_hint = "documento_importado.docx"
            elif "/spreadsheets/d/" in file_url:
                final_download_url = f"https://docs.google.com/spreadsheets/d/{file_id}/export?format=csv"
                original_name_hint = "hoja_de_calculo_importada.csv"
            elif "/presentation/d/" in file_url:
                final_download_url = f"https://docs.google.com/presentation/d/{file_id}/export?format=pptx"
                original_name_hint = "presentacion_importada.pptx"
            elif "/file/d/" in file_url:
                # PDFs, Word, Excel subidos
                final_download_url = f"https://drive.google.com/uc?export=download&id={file_id}"
                original_name_hint = ""  # Lo deducimos luego

        # --- Descargar archivo ---
        file_bytes, content_type = download_file_from_url(final_download_url)
        if file_bytes is None:
            return jsonify({"success": False, "error": "No se pudo descargar el archivo. Verifique el enlace y permisos."}), 400

        # --- Obtener nombre del archivo ---
        file_name = original_name_hint
        if not file_name:
            file_extension = ".dat"
            if 'csv' in content_type: file_extension = ".csv"
            elif 'json' in content_type: file_extension = ".json"
            elif 'pdf' in content_type: file_extension = ".pdf"
            elif 'vnd.openxmlformats-officedocument.wordprocessingml.document' in content_type: file_extension = ".docx"
            elif 'msword' in content_type: file_extension = ".doc"
            elif 'plain' in content_type: file_extension = ".txt"

            file_name = f"imported_file_{uuid.uuid4().hex[:8]}{file_extension}"

        # --- Delegar al dataset service ---
        file_stream = io.BytesIO(file_bytes)
        file_storage = FileStorage(stream=file_stream, filename=file_name, content_type=content_type)

        result = dataset_service.create_dataset_from_file(
            user_id=current_user.id,
            project_id=project_id,
            file=file_storage
        )

        status_code = 201 if result.get("success") else 400
        return jsonify(result), status_code

    except Exception as e:
        logger.error(f"Error fatal en la importación desde URL: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500



# =========================================================
# SECCIÓN DE API PARA LA GALERÍA DE IMÁGENES (VISUALS)
# =========================================================

# Configuración de seguridad para archivos (ejemplo: 10 MB máx.)
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  



# ---------------------------
# POST /api/visuals (Versión Definitiva)
# ---------------------------


@app.route('/api/visuals', methods=['POST'])
@token_required
def save_visual_endpoint(current_user):
    # --- 1. Validaciones iniciales ---
    if 'image' not in request.files:
        return jsonify({"success": False, "error": "Falta el archivo 'image'."}), 400

    file = request.files['image']
    filename = secure_filename(file.filename)
    if not filename:
        return jsonify({"success": False, "error": "El archivo no tiene un nombre válido."}), 400

    # --- LÓGICA MEJORADA CON VALOR POR DEFECTO ---
    ALLOWED_TYPES = ('generada', 'meme', 'subida', 'editada')
    visual_type = request.form.get('type')

    if not visual_type:
        visual_type = 'editada' # Asignamos el valor por defecto si está ausente.

    if visual_type not in ALLOWED_TYPES:
        # Rechazamos solo si se envió un tipo que no está en nuestra lista permitida.
        return jsonify({"success": False, "error": f"El tipo '{visual_type}' es inválido."}), 400
    # --- FIN DE LA LÓGICA MEJORADA ---

    # --- 2. Obtención y limpieza del project_id ---
    project_id = request.form.get('project_id')
    if project_id == 'null' or not project_id:
        project_id = None

    try:
        image_bytes = file.read()
        
        # ... (cálculo de color, análisis de Google, etc.) ...
        color_dominante_calculado = extraer_color_dominante(image_bytes)
        analisis_google, _ = analizar_contenido_imagen_google(image_bytes)
        tags_calculados = analisis_google.get('etiquetas', []) if analisis_google else []
        ocr_calculado = analisis_google.get('texto_detectado', None) if analisis_google else None

        # --- 3. Guardar el archivo físico en Supabase Storage ---
        storage_path, public_url = supabase_handler.save_visual_file(
            file_bytes=image_bytes,
            user_id=current_user.id,
            filename=filename,
        )
        if not storage_path:
            return jsonify({"success": False, "error": "Fallo en la subida del archivo a Storage."}), 500

        # --- 4. Preparar el registro completo para la base de datos ---
        record_data = {
            "user_id": current_user.id,
            "project_id": project_id,
            "storage_path": storage_path,
            "type": visual_type,
            "prompt": request.form.get('prompt'),
            "tags": tags_calculados,
            "ocr_text": ocr_calculado,
            "color_dominante": color_dominante_calculado
        }

        new_record = supabase_handler.insert_visual_record(record_data)
        if not new_record:
            return jsonify({"success": False, "error": "Fallo al insertar el registro en la base de datos."}), 500

        # --- 5. Devolver la respuesta de éxito ---
        new_record['public_url'] = public_url
        return jsonify({
            "success": True,
            "message": "Imagen guardada con éxito.",
            "visual": new_record
        }), 201

    except Exception as e:
        logger.critical(f"❌ Error en save_visual_endpoint para {current_user.id}: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500




# ---------------------------
# POST /api/remove-background
# ---------------------------

MAX_FILE_SIZE_MB = 10


@app.route("/api/advanced-edit", methods=["POST"])
def handle_advanced_edit():
    """
    Endpoint robusto para edición avanzada de imágenes en Flask:
    1. Elimina el fondo de la imagen principal.
    2. Si se proporciona un prompt, genera un fondo con IA.
    3. Si se sube un fondo personalizado, lo utiliza.
    4. Devuelve la imagen final.
    """
    try:
        imagen_principal = request.files.get("imagen_principal")
        prompt_fondo = request.form.get("prompt_fondo")
        fondo_personalizado = request.files.get("fondo_personalizado")

        # Validación: no permitir ambos a la vez
        if prompt_fondo and fondo_personalizado:
            return jsonify({"error": "Proporciona un prompt o un fondo personalizado, pero no ambos."}), 400

        # Validar que haya imagen principal
        if not imagen_principal:
            return jsonify({"error": "No se proporcionó la imagen principal."}), 400

        imagen_principal_bytes = imagen_principal.read()
        if len(imagen_principal_bytes) > MAX_FILE_SIZE_MB * 1024 * 1024:
            return jsonify({"error": "La imagen principal excede el tamaño máximo permitido (10MB)."}), 400

        fondo_personalizado_bytes = None
        if fondo_personalizado:
            fondo_personalizado_bytes = fondo_personalizado.read()
            if len(fondo_personalizado_bytes) > MAX_FILE_SIZE_MB * 1024 * 1024:
                return jsonify({"error": "El fondo personalizado excede el tamaño máximo permitido (10MB)."}), 400

        # Llamada al orquestador
        logger.info("🪄 Procesando edición avanzada...")
        imagen_resultante_bytes, error = orquestar_edicion_avanzada(
            imagen_original_bytes=imagen_principal_bytes,
            prompt_fondo_ia=prompt_fondo,
            fondo_personalizado_bytes=fondo_personalizado_bytes,
        )

        if error or not imagen_resultante_bytes:
            logger.error(f"❌ Error en edición avanzada: {error}")
            return jsonify({"error": error or "Error desconocido en la edición."}), 400

        logger.info("✅ Edición avanzada completada con éxito.")
        return Response(imagen_resultante_bytes, mimetype="image/png")

    except Exception as e:
        logger.exception(f"❌ Error inesperado en /api/advanced-edit: {e}")
        return jsonify({"error": "Error interno del servidor. Intenta de nuevo más tarde."}), 500


# ---------------------------
# POST /api/create-meme
# ---------------------------


@app.route('/api/create-meme', methods=['POST'])
@token_required
def handle_create_meme(current_user):
    try:
        imagen_bytes = None

        # --- Entrada de imagen ---
        if 'image' in request.files:
            file = request.files['image']
            imagen_bytes = file.read()
        elif 'image_url' in request.form:
            image_url = request.form.get('image_url')
            if image_url:
                imagen_bytes = download_image_from_url(image_url)

        if not imagen_bytes:
            return jsonify({
                "success": False,
                "error": "No se proporcionó imagen válida (ni como archivo ni como URL)."
            }), 400

        # --- Entrada de parámetros ---
        texto_arriba = request.form.get('topText', '').strip()
        texto_abajo = request.form.get('bottomText', '').strip()
        color_relleno = request.form.get('fontColor', 'white')
        color_borde = request.form.get('borderColor', 'black')
        nombre_fuente = request.form.get('fontName', 'impact')

        # Nuevo: tamaño de fuente
        try:
            font_size = int(request.form.get('fontSize')) if request.form.get('fontSize') else None
        except ValueError:
            font_size = None
            logger.warning("⚠️ fontSize inválido, ignorado.")

        # --- Generamos el meme ---
        resultado_meme, error = crear_meme(
            imagen_bytes, texto_arriba, texto_abajo,
            color_relleno, color_borde, nombre_fuente,
            font_size_override=font_size
        )

        if error:
            return jsonify({"success": False, "error": error}), 500

        meme_bytes = resultado_meme.get("imagen")
        if not meme_bytes:
            return jsonify({
                "success": False,
                "error": "La función crear_meme no devolvió bytes de imagen."
            }), 500

        return Response(meme_bytes, mimetype='image/png')

    except Exception as e:
        logger.exception(f"❌ Error inesperado en /api/create-meme: {e}")
        return jsonify({"success": False, "error": "Error interno del servidor"}), 500



@app.route('/api/analyze-content', methods=['POST'])
def handle_analyze_content():
    if 'image' not in request.files:
        return jsonify({"error": "No se encontró el archivo de imagen."}), 400
    
    file = request.files['image']
    imagen_bytes = file.read()

    analysis_results, error = analizar_contenido_imagen_google(imagen_bytes)

    if error:
        return jsonify({"error": error}), 500

    # Devolvemos el JSON con todas las etiquetas, objetos y texto
    return jsonify(analysis_results)



@app.route('/api/visuals', methods=['GET'])
@token_required
def get_visuals_endpoint(current_user):
    """
    Obtiene la galería paginada de imágenes de un usuario con filtros opcionales
    (búsqueda, tipo, color) y URLs firmadas temporales seguras.
    """
    # --- Parámetros seguros desde query string ---
    page = request.args.get("page", 1, type=int)
    per_page = request.args.get("per_page", 12, type=int)
    search_term = request.args.get("q", None, type=str)
    filter_type = request.args.get("type", None, type=str)
    color_filter = request.args.get("color", None, type=str)

    try:
        # --- Consulta a Supabase con filtros ---
        visuals, total_count = supabase_handler.get_visuals_by_user(
            user_id=current_user.id,
            page=page,
            per_page=per_page,
            search_term=search_term,
            filter_type=filter_type,
            color_filter=color_filter,
        )

        # --- Generación de URLs firmadas seguras ---
        for visual in visuals:
            storage_path = visual.get("storage_path")
            if not storage_path:
                visual["public_url"] = None
                continue

            try:
                signed_url_response = (
                    supabase_handler.get_db_client()
                    .storage.from_(supabase_handler._visuals_bucket_name)
                    .create_signed_url(storage_path, 3600)  # Expira en 1h
                )
                visual["public_url"] = signed_url_response.get("signedURL")
            except Exception as e:
                logger.warning(
                    f"⚠️ No se pudo generar URL firmada para visual_id={visual.get('id')}: {e}",
                    exc_info=True,
                )
                visual["public_url"] = None

        # --- Respuesta exitosa ---
        return jsonify({
            "success": True,
            "visuals": visuals,
            "total_count": total_count,
            "page": page,
            "per_page": per_page,
        }), 200

    except Exception as e:
        logger.error(
            f"❌ Error inesperado en get_visuals_endpoint (user_id={current_user.id}): {e}",
            exc_info=True,
        )
        return jsonify({
            "success": False,
            "error": "Error al obtener la galería de imágenes."
        }), 500


# ------------------------------------
# DELETE /api/visuals/<visual_id>
# ------------------------------------
@app.route('/api/visuals/<string:visual_id>', methods=['DELETE'])
@token_required
def delete_visual_endpoint(current_user, visual_id):
    """
    Elimina un activo visual (archivo en Storage y registro en la base de datos).
    Solo el dueño del recurso puede borrarlo.
    """
    try:
        # 1. Verificar que el registro exista y pertenezca al usuario
        response = (
            supabase_handler.get_db_client()
            .table("imagenes")
            .select("user_id, storage_path")
            .eq("id", visual_id)
            .single()
            .execute()
        )

        if not response.data:
            logger.info(f"⚠️ Intento de borrar imagen inexistente: {visual_id}")
            return jsonify({"success": False, "error": "La imagen no existe."}), 404

        if response.data["user_id"] != current_user.id:
            logger.warning(
                f"🚫 Usuario {current_user.id} intentó eliminar imagen ajena {visual_id}"
            )
            return jsonify({"success": False, "error": "No tienes permiso para eliminar esta imagen."}), 403

        storage_path = response.data["storage_path"]

        # 2. Eliminar archivo en Storage (si existe)
        if storage_path:
            deleted = supabase_handler.delete_visual_file(storage_path)
            if not deleted:
                logger.error(
                    f"⚠️ Falló la eliminación del archivo en Storage: {storage_path}"
                )

        # 3. Eliminar registro en la base de datos
        (
            supabase_handler.get_db_client()
            .table("imagenes")
            .delete()
            .eq("id", visual_id)
            .execute()
        )

        logger.info(f"✅ Imagen {visual_id} eliminada por usuario {current_user.id}")
        return jsonify({"success": True, "message": "Imagen eliminada correctamente."}), 200

    except Exception as e:
        logger.critical(
            f"❌ Error en delete_visual_endpoint (user={current_user.id}, visual_id={visual_id}): {e}",
            exc_info=True,
        )
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500




@app.route("/api/generate-image", methods=["POST"])
@token_required
def handle_generate_image(current_user):
    """
    Endpoint para generar una imagen a partir de un prompt de texto.
    """
    try:
        # --- 1. Obtener y Validar la Entrada ---
        data = request.get_json()
        if not data:
            logger.warning("No se recibió cuerpo JSON en /api/generate-image")
            return jsonify({"error": "No se recibió un cuerpo JSON."}), 400

        prompt = data.get("prompt")
        if not prompt or not prompt.strip():
            logger.warning("Prompt vacío recibido en /api/generate-image")
            return jsonify({"error": "El prompt no puede estar vacío."}), 400

        # Limitar longitud del prompt para no saturar la API de IA
        MAX_PROMPT_LENGTH = 500
        prompt = prompt.strip()
        if len(prompt) > MAX_PROMPT_LENGTH:
            logger.info(f"Prompt demasiado largo, truncando a {MAX_PROMPT_LENGTH} caracteres.")
            prompt = prompt[:MAX_PROMPT_LENGTH]

        # --- 2. Llamar a la Lógica de Negocio ---
        logger.info(f"✨ Iniciando generación de imagen con prompt: '{prompt[:50]}...' (usuario={current_user.id})")
        imagen_bytes, error = generar_imagen_desde_texto(prompt)

        # --- 3. Manejar la Respuesta ---
        if error:
            logger.error(f"❌ Error durante la generación de imagen: {error}")

            # Código 503 si el modelo de IA está cargando
            if "modelo de IA está cargando" in error:
                return jsonify({"error": error}), 503

            # Otros errores
            return jsonify({"error": error}), 400

        if not imagen_bytes:
            logger.error("La generación no produjo una imagen en /api/generate-image")
            return jsonify({"error": "La generación no produjo una imagen."}), 500

        logger.info(f"✅ Imagen generada con éxito para usuario={current_user.id}")
        return Response(imagen_bytes, mimetype="image/jpeg")

    except Exception as e:
        logger.exception(f"❌ Error inesperado en /api/generate-image: {e}")
        return jsonify({"error": "Error interno del servidor. Intenta nuevamente más tarde."}), 500

MAX_FILE_SIZE_MB = 10


@app.route('/api/project/<string:projectId>/vision/upload-for-gallery-and-dataset', methods=['POST'])
@token_required
def upload_for_gallery_and_dataset(current_user, projectId):
    """
    ENDPOINT HÍBRIDO DE PRODUCCIÓN (VERSIÓN FINAL):
    1. Genera un ID de dataset ANTES de procesar.
    2. Recibe un lote de imágenes.
    3. Analiza cada imagen y la GUARDA en la galería asociada al dataset.
    4. Genera un dataset tabular (CSV) con todos los análisis.
    5. Guarda el registro del dataset en la base de datos.
    """
    try:
        if 'images' not in request.files:
            return jsonify({"success": False, "error": "No se encontraron archivos de imagen."}), 400

        files = request.files.getlist('images')
        if not files or all(f.filename == '' for f in files):
            return jsonify({"success": False, "error": "No se seleccionó ningún archivo."}), 400

        new_dataset_id = str(uuid.uuid4())
        logger.info(f"Nuevo dataset de visión iniciado con ID: {new_dataset_id}")

        analysis_results_for_csv = []
        newly_created_visuals = []
        processing_errors = []

        for file in files:
            filename = secure_filename(file.filename)
            if not is_allowed_image_file(filename):
                processing_errors.append({"archivo": filename, "error": "Tipo de archivo no permitido."})
                continue
            image_bytes = file.read()
            if len(image_bytes) > MAX_FILE_SIZE_MB * 1024 * 1024:
                processing_errors.append({"archivo": filename, "error": "Archivo excede el tamaño máximo (10MB)."})
                continue
            try:
                new_image_id = str(uuid.uuid4())
                analysis_result = procesar_imagen_completa(
                    image_bytes,
                    filename,
                    model_manager_instance=model_manager,
                    image_id=new_image_id
                )
                if not analysis_result or "error" in analysis_result:
                    raise Exception(analysis_result.get("error", "Fallo desconocido en análisis."))
                analysis_result['dataset_id'] = new_dataset_id
                analysis_results_for_csv.append(analysis_result)
                storage_path, public_url = supabase_handler.save_visual_file(
                    file_bytes=image_bytes, user_id=current_user.id, filename=filename)
                if not storage_path:
                    raise Exception("Fallo al guardar en Supabase Storage.")
                record_to_insert = {
                    "id": new_image_id, "user_id": current_user.id, "project_id": projectId,
                    "dataset_id": new_dataset_id, "storage_path": storage_path,
                    "prompt": analysis_result.get("descripcion_ia", f"Imagen subida: {filename}"),
                    "type": "subida", "color_dominante": analysis_result.get("color_dominante_hex")}
                visual_record = supabase_handler.insert_visual_record(record_to_insert)
                if not visual_record:
                    raise Exception("Fallo al insertar registro en la tabla 'imagenes'.")
                newly_created_visuals.append(visual_record)
            except Exception as e:
                logger.error(f"❌ Error procesando '{filename}': {e}", exc_info=True)
                processing_errors.append({"archivo": filename, "error": str(e)})

        if not analysis_results_for_csv:
            return jsonify({"success": False, "error": "Ninguna imagen pudo ser procesada.", "processing_errors": processing_errors}), 500

        try:
            df = pd.DataFrame(analysis_results_for_csv).fillna("N/A")
            csv_bytes = df.to_csv(index=False).encode('utf-8')
            file_name = f"vision_analysis_dataset_{new_dataset_id[:8]}.csv"
            storage_path, _ = supabase_handler.save_file(
                file_bytes=csv_bytes, user_id=current_user.id, project_id=projectId, folder="datasets", filename=file_name)
            if not storage_path:
                raise Exception("No se pudo guardar el archivo del dataset en Supabase Storage.")
            dataset_name = f"Análisis de {len(analysis_results_for_csv)} imágenes ({pd.Timestamp.now().strftime('%Y-%m-%d')})"
            new_dataset = dataset_service.create_dataset_record(
                dataset_id=new_dataset_id, user_id=current_user.id, project_id=projectId, dataset_name=dataset_name,
                dataset_type="vision_analysis", storage_path=storage_path, file_size=len(csv_bytes))
            if not new_dataset:
                raise Exception("No se pudo crear el registro del dataset en la base de datos.")

            # --- INICIO DEL CAMBIO 1 ---
            logger.info("Proceso de subida completado. Esperando 2 segundos para permitir que Supabase se estabilice...")
            time.sleep(2) # Pausa de 2 segundos
            # --- FIN DEL CAMBIO 1 ---

            return jsonify({
                "success": True, "message": "Dataset creado y las imágenes fueron añadidas a tu galería.",
                "new_dataset": new_dataset, "new_visuals_count": len(newly_created_visuals),
                "processing_errors": processing_errors}), 201

        except Exception as e:
            logger.exception(f"❌ Error creando dataset tabular para user={current_user.id}, project={projectId}")
            
            # --- INICIO DEL CAMBIO 2 ---
            logger.info("Proceso de subida parcial completado. Esperando 2 segundos para permitir que Supabase se estabilice...")
            time.sleep(5) # Pausa de 2 segundos
            # --- FIN DEL CAMBIO 2 ---

            return jsonify({
                "success": True, "message": "Imágenes procesadas y añadidas a galería, pero falló la creación del dataset.",
                "new_dataset": None, "new_visuals_count": len(newly_created_visuals),
                "processing_errors": processing_errors + [{"dataset_creation_error": str(e)}]
            }), 207

    except Exception as e:
        logger.error(f"❌ Error fatal en upload-for-gallery-and-dataset: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500



@app.route('/api/vision-lab/dataset/<string:dataset_id>', methods=['GET'])
@token_required
def get_vision_lab_data(current_user, dataset_id):
    """
    Endpoint robusto para el Laboratorio de Visión.
    Devuelve:
      - Datos de análisis (CSV)
      - Metadatos de las imágenes (storage_path)
      - Etiquetas manuales ya guardadas
      - URLs firmadas para visualización
    """
    try:
        # --- 1. Cargar CSV de análisis ---
        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        path_csv = dataset_info["data"].get("storage_path")
        df = supabase_handler.load_file_as_dataframe(user_id=current_user.id, path=path_csv)
        if df is None:
            return jsonify({"success": False, "error": "Error cargando archivo de análisis."}), 500

        # Parseo seguro de JSON dentro del CSV
        def safe_parse(value):
            if isinstance(value, str):
                try:
                    return json.loads(value.replace("'", '"'))
                except:
                    return value
            return value

        for col in ['metadata', 'tags_ia', 'objetos_detectados']:
            if col in df.columns:
                df[col] = df[col].apply(safe_parse)
        analysis_list = df.to_dict(orient='records')

        # --- 2. Cargar metadatos de imágenes ---
        image_records_response = supabase_admin.table("imagenes") \
            .select("id, storage_path") \
            .eq("dataset_id", dataset_id) \
            .eq("user_id", current_user.id) \
            .execute()
        image_metadata_list = image_records_response.data or []

        # --- 3. Obtener etiquetas manuales para todas las imágenes ---
        image_ids = [img['id'] for img in image_metadata_list]
        tags_response = supabase_admin.table("image_tags") \
            .select("image_id, tags(name)") \
            .in_("image_id", image_ids) \
            .execute()
        
        tags_map = {}
        for row in (tags_response.data or []):
            tags_map.setdefault(row['image_id'], []).append(row['tags']['name'])

        # --- 4. Combinar análisis, metadatos y etiquetas ---
        image_metadata_map = {img['id']: img for img in image_metadata_list}
        combined_data = [
            {
                **analysis_item,
                'storage_path': image_metadata_map.get(analysis_item.get('image_id'), {}).get('storage_path'),
                'custom_tags': tags_map.get(analysis_item.get('image_id'), [])
            }
            for analysis_item in analysis_list
        ]

        # --- 5. Crear URLs firmadas ---
        image_records_for_signing = [{"id": img['id'], "storage_path": img['storage_path']} for img in image_metadata_list]
        image_data_with_urls = supabase_handler.create_signed_urls_for_visuals(image_records_for_signing)

        # --- 6. Respuesta final ---
        return jsonify({
            "success": True,
            "analysisData": combined_data,
            "imageData": image_data_with_urls
        }), 200

    except Exception as e:
        logger.exception(f"❌ Error en get_vision_lab_data para dataset_id={dataset_id}: {e}")
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500




@app.route('/api/models/<string:model_id>', methods=['PATCH'])
@token_required
def update_visualmodel_details_endpoint(current_user, model_id):
    """
    Actualiza los detalles de un modelo entrenado (ej: su nombre).
    Sirve tanto para modelos tabulares como de visión.
    """
    try:
        data = request.get_json()
        if not data or 'model_display_name' not in data:
            return jsonify({"success": False, "error": "Falta el parámetro 'model_display_name'."}), 400

        new_name = data['model_display_name']
        if not isinstance(new_name, str) or not new_name.strip():
            return jsonify({"success": False, "error": "El nombre del modelo no puede estar vacío."}), 400

        # Actualizamos en la base de datos, asegurándonos que el modelo pertenece al usuario
        update_query = (
            supabase_admin.table('trained_models')
            .update({'model_display_name': new_name.strip()})
            .eq('id', model_id)
            .eq('user_id', current_user.id)
            .execute()
        )

        # La respuesta de update contiene los datos actualizados si fue exitoso
        if not update_query.data:
            return jsonify({"success": False, "error": "Modelo no encontrado o acceso denegado."}), 404

        return jsonify({
            "success": True,
            "message": "Modelo renombrado con éxito.",
            "data": update_query.data[0]
        }), 200

    except Exception as e:
        logger.error(f"[UPDATE MODEL] Error inesperado al actualizar {model_id}: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Ocurrió un error interno en el servidor."}), 500


# ================================================================
def _get_clustering_data(user_id: str, model_id: str) -> Optional[dict]:
    """
    Recupera los resultados de clustering de un modelo guardado en Supabase Storage.
    """
    try:
        model_query = (
            supabase_admin.table('trained_models')
            .select('model_storage_path')
            .eq('id', model_id)
            .eq('user_id', user_id)
            .single()
            .execute()
        )

        if not model_query.data or not model_query.data.get('model_storage_path'):
            logger.warning(f"[ClusteringData] No se encontró storage_path para modelo {model_id}")
            return None

        storage_path = model_query.data['model_storage_path']
        file_bytes = supabase_admin.storage.from_("proyectos-usuarios").download(storage_path)

        if not file_bytes:
            logger.error(f"[ClusteringData] No se pudo descargar el archivo en {storage_path}")
            return None

        return json.loads(file_bytes.decode("utf-8"))

    except Exception as e:
        logger.error(f"[ClusteringData] Error recuperando datos de clustering: {e}", exc_info=True)
        return None


@app.route('/api/vision/available-models', methods=['GET'])
def get_available_vision_models():
    # vision_prediction_service es tu instancia de VisionPredictionService
    available_architectures = list(vision_prediction_service.architectures.keys())
    return jsonify({
        "success": True,
        "models": available_architectures
    }), 200



# ================================================================
#   CACHE TEMPORAL Y LIMPIEZA AUTOMÁTICA (ya inicializado en train)
# ================================================================
# Estructura: { temp_id: { "path": str, "created_at": timestamp } }
TEMP_TRAINING_CACHE = {}
TEMP_TRAINING_TTL = 60 * 30  # 30 minutos

def cleanup_temp_models():
    while True:
        now = time.time()
        to_delete = []
        for temp_id, info in list(TEMP_TRAINING_CACHE.items()):
            if now - info['created_at'] > TEMP_TRAINING_TTL:
                try:
                    Path(info['path']).unlink(missing_ok=True)
                    logger.info(f"[TempCleanup] Eliminado modelo temporal caducado: {info['path']}")
                except Exception as e:
                    logger.error(f"[TempCleanup] Error eliminando {info['path']}: {e}")
                to_delete.append(temp_id)
        for temp_id in to_delete:
            TEMP_TRAINING_CACHE.pop(temp_id, None)
        time.sleep(300)

cleanup_thread = threading.Thread(target=cleanup_temp_models, daemon=True)
cleanup_thread.start()



# ================================================================
#   ENDPOINT DE GUARDADO DEFINITIVO DE MODELO VISIÓN
# ================================================================
@app.route('/api/project/<string:project_id>/vision/save-model', methods=['POST'])
@token_required
def handle_vision_save_model(current_user, project_id):
    """
    Guarda un modelo de visión entrenado usando un ID temporal del caché del servidor.
    Sube el modelo a Supabase Storage y registra metadata en la tabla `trained_models`.
    """
    try:
        # --- 0. Validación de entrada ---
        data = request.get_json(silent=True) or {}
        required_keys = ['temp_training_id', 'model_display_name', 'training_results']
        if not all(k in data for k in required_keys):
            return jsonify({
                "success": False,
                "error": f"Faltan claves requeridas: {required_keys}"
            }), 400

        temp_id = data["temp_training_id"]
        model_display_name = data["model_display_name"]
        training_results = data.get("training_results", {})

        # --- 1. Recuperar archivo temporal ---
        cache_entry = TEMP_TRAINING_CACHE.get(temp_id)
        if not cache_entry:
            return jsonify({
                "success": False,
                "error": "La sesión de entrenamiento ha expirado o el ID es inválido."
            }), 404

        temp_file_path = Path(cache_entry["path"])
        try:
            with open(temp_file_path, "rb") as f:
                model_bytes = f.read()
        except FileNotFoundError:
            return jsonify({"success": False, "error": "El archivo temporal no fue encontrado."}), 404

        # --- 2. Subir modelo a Supabase Storage ---
        project_query = supabase_admin.table('projects').select('project_name').eq('project_id', project_id).single().execute()
        if not project_query.data:
            return jsonify({"success": False, "error": "Proyecto no encontrado."}), 404
        project_name = project_query.data['project_name']

        model_id = str(uuid.uuid4())
        model_storage_path = f"{current_user.id}/{project_id}/models/{model_id}_vision.joblib"

        logger.info(f"[VisionSave] Intentando subir {len(model_bytes)} bytes a {model_storage_path}")

        supabase_admin.storage.from_("proyectos-usuarios").upload(
            path=model_storage_path,
            file=model_bytes,
            file_options={"content-type": "application/octet-stream"}
        )

        logger.info(f"[VisionSave] Subida a Storage completada exitosamente.")

        # --- 3. Insertar metadata en la BD ---
        new_model_record = {
            "id": model_id,
            "user_id": current_user.id,
            "project_id": project_id,
            "problem_type": "vision_classification",
            "project_name": project_name,
            "model_name": training_results.get("model_name", "ResNet34"),
            "model_display_name": model_display_name,
            "model_storage_path": model_storage_path,
            "source_dataset_id": data.get("source_dataset_id"),
            "target_col": "N/A (Vision)",
            "feature_cols": ["image"],
            "evaluation_results": training_results
        }

        supabase_admin.table("trained_models").insert(new_model_record).execute()
        logger.info(f"[VisionSave] Metadata del modelo '{model_display_name}' guardada en BD.")

        # --- 4. Limpieza de caché y archivo temporal ---
        TEMP_TRAINING_CACHE.pop(temp_id, None)
        temp_file_path.unlink(missing_ok=True)
        logger.info(f"[VisionSave] Limpieza completada de archivo temporal {temp_file_path}")

        return jsonify({
            "success": True,
            "message": "Modelo de visión guardado y registrado exitosamente.",
            "model_id": model_id
        }), 201

    except werkzeug.exceptions.RequestEntityTooLarge:
        return jsonify({"success": False, "error": "El cuerpo de la solicitud es demasiado grande."}), 413
    except Exception as e:
        logger.exception(f"[VisionSave] Error inesperado: {e}")
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500


# ================================================================
#   ENDPOINT DE ENTRENAMIENTO DE VISIÓN
# ================================================================


# Lock global para controlar entrenamientos concurrentes
TRAINING_LOCK = Lock()

@app.route('/api/project/<string:project_id>/vision/train', methods=['POST'])
@token_required
def handle_vision_train(current_user, project_id):
    """
    Entrena un modelo de visión. Usa un Lock para prevenir entrenamientos concurrentes.
    """

    # --- 1. Intentar adquirir lock ---
    if not TRAINING_LOCK.acquire(blocking=False):
        logger.warning("[VisionTrain] Se rechazó un nuevo entrenamiento porque ya hay uno en curso.")
        return jsonify({
            "success": False,
            "error": "Ya hay un proceso de entrenamiento en ejecución. Por favor, espera a que termine antes de iniciar uno nuevo."
        }), 429

    try:
        # --- 2. Validar usuario ---
        if not current_user or not hasattr(current_user, "id"):
            return jsonify({"success": False, "error": "Usuario inválido."}), 401

        data = request.get_json(silent=True) or {}
        training_config = data.get("training_config", {})
        model_arch = training_config.get("model_arch", "resnet34")
        epochs = int(training_config.get("epochs", 5))

        cluster_id = data.get("source_cluster_result_id")
        group_labels = data.get("group_labels")  # dict {'0': 'verano', '1': 'invierno'}

        image_records = []

        # --- 3. Selección de datos de entrenamiento ---
        if cluster_id and group_labels:
            logger.info(f"[VisionTrain] Modo clustering con cluster_id={cluster_id}")
            cluster_data = _get_clustering_data(current_user.id, cluster_id)

            if not cluster_data or 'results' not in cluster_data:
                return jsonify({"success": False, "error": "No se pudieron cargar los datos del clustering."}), 404

            for item in cluster_data['results']:
                group_id_str = str(item.get('cluster_id'))
                if group_id_str in group_labels:
                    image_records.append({
                        "storage_path": item.get('storage_path'),
                        "tags": [group_labels[group_id_str]]
                    })
        else:
            logger.info(f"[VisionTrain] Modo etiquetas generales para proyecto {project_id}")
            image_records = supabase_handler.get_tagged_images_for_project(
                user_id=current_user.id,
                project_id=project_id
            )

        # --- 4. Validación de datos ---
        if not image_records or len(image_records) < 2:
            logger.warning(f"[VisionTrain] No hay suficientes imágenes para entrenar. Total: {len(image_records)}")
            return jsonify({
                "success": False,
                "error": "No hay suficientes imágenes seleccionadas para el entrenamiento (mínimo 2)."
            }), 400

        # --- 5. Entrenamiento ---
        logger.info(f"[VisionTrain] Entrenando modelo {model_arch} con {len(image_records)} imágenes")
        result = vision_prediction_service.train_model(
            image_records=image_records,
            model_arch=model_arch,
            epochs=epochs
        )

        if not result.get("success"):
            logger.error(f"[VisionTrain] Error durante entrenamiento: {result.get('error')}")
            return jsonify({"success": False, "error": result.get('error')}), 500

        # --- 6. Manejo de artefactos temporales ---
        training_data = result.get("data", {})
        if 'artifacts_bytes' in training_data:
            model_bytes = training_data.pop('artifacts_bytes')
            temp_id = f"temp_{uuid.uuid4()}"
            temp_file_path = Path(tempfile.gettempdir()) / f"{temp_id}_vision_model.joblib"

            with open(temp_file_path, "wb") as f:
                f.write(model_bytes)

            TEMP_TRAINING_CACHE[temp_id] = {
                "path": str(temp_file_path),
                "created_at": time.time()
            }

            logger.info(f"[VisionTrain] Modelo temporal guardado en {temp_file_path} con ID {temp_id}")
            training_data['temp_training_id'] = temp_id

        # --- 7. Respuesta final ---
        return jsonify({
            "success": True,
            "message": "Entrenamiento completado exitosamente.",
            "training_results": training_data
        }), 200

    except Exception as e:
        logger.error(f"[VisionTrain] Error inesperado: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500

    finally:
        # --- 8. Liberar lock SIEMPRE ---
        TRAINING_LOCK.release()



# 1️⃣ Endpoint para AÑADIR una etiqueta a una imagen
@app.route('/api/images/<string:image_id>/tags', methods=['POST'])
@token_required
def add_tag_to_image(current_user, image_id):
    data = request.get_json()
    tag_name = (data.get('tagName') or "").strip()
    if not tag_name:
        return jsonify({"success": False, "error": "tagName no proporcionado"}), 400

    try:
        db = supabase_handler.get_db_client()

        # 1. Verificar que la imagen pertenece al usuario
        image_resp = db.table("imagenes").select("id").eq("id", image_id).eq("user_id", current_user.id).single().execute()
        if not image_resp.data:
            return jsonify({"success": False, "error": "Imagen no encontrada o no pertenece al usuario"}), 404

        # 2. Buscar si la etiqueta existe para este usuario
        tag_resp = db.table("tags").select("id").eq("name", tag_name).eq("user_id", current_user.id).execute()

        if tag_resp.data:
            # Si existe, usamos el id de la primera coincidencia
            tag_id = tag_resp.data[0]['id']
        else:
            # Si no existe, la creamos
            tag_id = str(uuid.uuid4())
            db.table("tags").insert({
                "id": tag_id,
                "name": tag_name,
                "user_id": current_user.id
            }).execute()

        # 3. Insertar en tabla de unión (ignora duplicados con upsert)
        db.table("image_tags").upsert({
            "image_id": image_id,
            "tag_id": tag_id
        }).execute()

        return jsonify({"success": True, "message": f"Etiqueta '{tag_name}' añadida"}), 201

    except Exception as e:
        logger.exception(f"Error añadiendo etiqueta '{tag_name}' a imagen {image_id}: {e}")
        return jsonify({"success": False, "error": "Error interno al añadir etiqueta"}), 500


# 2️⃣ Endpoint para ELIMINAR una etiqueta de una imagen
@app.route('/api/images/<string:image_id>/tags/<string:tag_name>', methods=['DELETE'])
@token_required
def remove_tag_from_image(current_user, image_id, tag_name):
    tag_name = tag_name.strip()
    if not tag_name:
        return jsonify({"success": False, "error": "tagName no proporcionado"}), 400

    try:
        # Verificar que la imagen pertenece al usuario
        image_resp = supabase_handler.get_db_client().table("imagenes").select("id").eq("id", image_id).eq("user_id", current_user.id).single().execute()
        if not image_resp.data:
            return jsonify({"success": False, "error": "Imagen no encontrada o no pertenece al usuario"}), 404

        # Obtener ID de la etiqueta
        tag_resp = supabase_handler.get_db_client().table("tags").select("id").eq("name", tag_name).eq("user_id", current_user.id).single().execute()
        if not tag_resp.data:
            return jsonify({"success": False, "error": "Etiqueta no encontrada"}), 404
        tag_id = tag_resp.data['id']

        # Eliminar de image_tags
        supabase_handler.get_db_client().table("image_tags").delete().eq("image_id", image_id).eq("tag_id", tag_id).execute()

        return jsonify({"success": True, "message": f"Etiqueta '{tag_name}' eliminada"}), 200

    except Exception as e:
        logger.exception(f"Error eliminando etiqueta '{tag_name}' de imagen {image_id}: {e}")
        return jsonify({"success": False, "error": "Error interno al eliminar etiqueta"}), 500


# 3️⃣ Endpoint para OBTENER las etiquetas de una imagen
@app.route('/api/images/<string:image_id>/tags', methods=['GET'])
@token_required
def get_tags_for_image(current_user, image_id):
    try:
        # Verificar que la imagen pertenece al usuario
        image_resp = supabase_handler.get_db_client().table("imagenes").select("id").eq("id", image_id).eq("user_id", current_user.id).single().execute()
        if not image_resp.data:
            return jsonify({"success": False, "error": "Imagen no encontrada o no pertenece al usuario"}), 404

        # Obtener etiquetas (JOIN implícito)
        tags_resp = supabase_handler.get_db_client().table("image_tags").select("tags(name)").eq("image_id", image_id).execute()
        tags = [t['name'] for t in tags_resp.data] if tags_resp.data else []

        return jsonify({"success": True, "tags": tags}), 200

    except Exception as e:
        logger.exception(f"Error obteniendo etiquetas para imagen {image_id}: {e}")
        return jsonify({"success": False, "error": "Error interno al obtener etiquetas"}), 500


@app.route('/api/datasets/<string:dataset_id>/bulk-tag-by-cluster', methods=['POST'])
@token_required
def handle_bulk_tagging_by_cluster(current_user, dataset_id):
    """
    Aplica etiquetas a imágenes basándose en los resultados de un clustering previo.
    """
    data = request.get_json()
    cluster_result_id = data.get('clusterResultId')
    group_labels = data.get('group_labels') # ej: {'0': 'verano', '1': 'invierno'}

    if not cluster_result_id or not isinstance(group_labels, dict):
        return jsonify({"success": False, "error": "Faltan clusterResultId o group_labels."}), 400
    
    try:
        # 1. Cargar el archivo JSON con los resultados del clustering desde Supabase Storage
        #    Necesitamos obtener el storage_path de este modelo de clustering
        model_meta_query = supabase_admin.table('trained_models').select('model_storage_path').eq('id', cluster_result_id).eq('user_id', current_user.id).single().execute()
        if not model_meta_query.data:
            return jsonify({"success": False, "error": "Resultado de clustering no encontrado."}), 404
        
        storage_path = model_meta_query.data['model_storage_path']
        cluster_file_bytes = supabase_admin.storage.from_("proyectos-usuarios").download(storage_path)
        clustering_details = json.loads(cluster_file_bytes)
        
        # Creamos un mapa de storage_path -> cluster_id para búsqueda rápida
        path_to_cluster_map = {item['storage_path']: item['cluster_id'] for item in clustering_details.get('results', [])}

        # 2. Obtener todas las imágenes del dataset que estamos etiquetando
        image_records = supabase_handler.get_visuals_by_dataset(dataset_id, current_user.id)

        # 3. Iterar y aplicar etiquetas
        tags_applied_count = 0
        for image in image_records:
            image_path = image.get('storage_path')
            image_id = image.get('id')
            
            # Si esta imagen estaba en los resultados del clustering...
            if image_path in path_to_cluster_map:
                cluster_id = str(path_to_cluster_map[image_path])
                
                # ...y si ese cluster tiene una etiqueta definida por el usuario...
                if cluster_id in group_labels and group_labels[cluster_id]:
                    tag_name = group_labels[cluster_id].strip()
                    
                    # ...insertamos la etiqueta en la base de datos (usando tu lógica existente)
                    # Esto asume que tienes una tabla `imagen_tags` y una tabla `tags`
                    # y una función para manejar la inserción.
                    supabase_handler.add_tag_to_image(image_id, tag_name, current_user.id)
                    tags_applied_count += 1

        return jsonify({
            "success": True, 
            "message": f"{tags_applied_count} etiquetas han sido guardadas exitosamente."
        }), 200

    except Exception as e:
        logger.error(f"Error en bulk-tagging para dataset {dataset_id}: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor al guardar etiquetas."}), 500

# ================================================================
#   ENDPOINT: EJECUTAR CLUSTERING K-MEANS
# ================================================================
@app.route("/api/v1/cluster/run_kmeans", methods=["POST"])
@token_required
def run_kmeans_clustering(current_user):
    """
    Recibe rutas de imágenes en Supabase, descarga las imágenes y ejecuta clustering K-Means.
    Devuelve un mapping {storage_path, cluster_id, plot_coords}.
    """
    try:
        # --- 1. Extraer parámetros del request ---
        image_storage_paths = request.form.getlist("image_storage_paths")
        n_clusters = request.form.get("n_clusters", type=int, default=5)

        if not image_storage_paths:
            return jsonify({"success": False, "error": "Debe proporcionar al menos una ruta de imagen."}), 400

        logger.info(f"📥 Recibidas {len(image_storage_paths)} imágenes para clustering en {n_clusters} clusters.")

        # --- 2. Obtener el servicio de clustering ---
        tool_response = model_manager.get_tool("Local: Image Clustering (K-Means)")
        if not tool_response["success"]:
            return jsonify({"success": False, "error": tool_response["error"]}), 500

        clustering_service = tool_response["tool_object"]

        # --- 3. Descargar imágenes desde Supabase ---
        image_bytes_list = []
        valid_storage_paths = []  # <-- Mantener solo las rutas válidas
        for path in image_storage_paths:
            try:
                image_bytes = supabase_handler.download_visual_file(path)
                if image_bytes:
                    image_bytes_list.append(image_bytes)
                    valid_storage_paths.append(path)
                else:
                    logger.warning(f"⚠️ No se pudo descargar la imagen en {path}")
            except Exception:
                logger.exception(f"❌ Error al descargar {path}")
                continue

        if not image_bytes_list:
            return jsonify({"success": False, "error": "No se pudo descargar ninguna de las imágenes."}), 400

        # --- 4. Ejecutar clustering ---
        result = clustering_service.run_kmeans_clustering(image_bytes_list, n_clusters)
        if not result.get("success"):
            return jsonify({"success": False, "error": result.get("error", "Error desconocido en clustering.")}), 500

        # --- 5. Armar respuesta (MODIFICADO: incluimos coordenadas PCA) ---
        response_data = {
            "success": True,
            "results": [
                {
                    "storage_path": path,
                    "cluster_id": cluster,
                    "plot_coords": {"x": coords[0], "y": coords[1]}
                }
                for path, cluster, coords in zip(valid_storage_paths, result["cluster_labels"], result["plot_coords"])
            ],
            "n_clusters": result["n_clusters"],
            "metrics": result.get("metrics", {})
        }

        logger.info("✅ Clustering completado correctamente.")
        return jsonify(response_data), 200

    except Exception as e:
        logger.exception("❌ Error inesperado en /api/v1/cluster/run_kmeans")
        return jsonify({"success": False, "error": str(e)}), 500



# ================================================================
#   ENDPOINT: GUARDAR RESULTADO DE CLUSTERING (APROBADO POR USUARIO)
# ================================================================


@app.route('/api/v1/cluster/approve_and_save', methods=['POST'])
@token_required
def approve_and_save_cluster(current_user):
    try:
        data = request.get_json()
        project_id = data.get("projectId")
        dataset_id = data.get("datasetId") # Asegúrate de que el frontend lo envía
        project_name = data.get("projectName")
        model_display_name = data.get("modelDisplayName")
        clustering_details = data.get("clusteringDetails")

        if not all([project_id, dataset_id, project_name, model_display_name, clustering_details]):
            return jsonify({"success": False, "error": "Faltan datos en la petición."}), 400

        # --- 1. Guardar artefacto en Supabase Storage ---
        model_id = str(uuid.uuid4())
        storage_path = f"{current_user.id}/clustering/{project_name.replace(' ', '_')}_{model_id}.json"
        
        supabase_admin.storage.from_("proyectos-usuarios").upload(
            file=json.dumps(clustering_details, indent=2).encode('utf-8'),
            path=storage_path,
            file_options={"content-type": "application/json"}
        )

        # --- 2. Crear el registro COMPLETO para la tabla trained_models ---
        new_model_record = {
            "id": model_id,
            "user_id": current_user.id,
            "project_id": project_id,
            "project_name": project_name,
            "model_name": "KMeans",
            "model_display_name": model_display_name,
            "problem_type": "clustering",
            
            # --- Columnas de relleno para satisfacer las restricciones ---
            "source_dataset_id": dataset_id,
            "target_col": "N/A (Clustering)",
            "feature_cols": [], # Una lista vacía es un valor válido
            
            # --- Columna de la ruta con el nombre correcto ---
            "model_storage_path": storage_path,
            
            # --- Resultados de evaluación específicos de clustering ---
            "evaluation_results": {
                "problem_type": "clustering",
                "n_clusters": clustering_details.get("n_clusters")
            }
        }

        insert_query = supabase_admin.table('trained_models').insert(new_model_record).execute()

        logger.info(f"✅ Resultado de clustering '{model_display_name}' guardado por usuario {current_user.id}")
        return jsonify({
            "success": True,
            "message": "Resultado de clustering guardado con éxito.",
            "model_id": insert_query.data[0]['id']
        }), 201

    except Exception as e:
        logger.error(f"❌ Error de PostgREST al guardar: {e.message}", exc_info=True)
        # Intentamos borrar el archivo que ya se subió para no dejar basura
        if 'storage_path' in locals():
            try:
                supabase_admin.storage.from_("proyectos-usuarios").remove([storage_path])
            except Exception as cleanup_error:
                logger.error(f"Error en la limpieza de storage: {cleanup_error}")
        return jsonify({"success": False, "error": f"Error de base de datos: {e.message}"}), 500
    except Exception as e:
        logger.exception("❌ Error inesperado en /api/v1/cluster/approve_and_save")
        return jsonify({"success": False, "error": str(e)}), 500



@app.route('/api/models/clustering/<string:model_id>', methods=['GET'])
@token_required
def get_clustering_result_details(current_user, model_id):
    """
    Obtiene los detalles de un resultado de clustering guardado,
    incluyendo el archivo JSON del storage.
    """
    try:
        # 1. Buscar el registro del modelo en la base de datos
        model_query = (
            supabase_admin.table('trained_models')
            .select('model_storage_path')
            .eq('id', model_id)
            .eq('user_id', current_user.id)
            .single()
            .execute()
        )

        if not model_query or not model_query.data:
            return jsonify({
                "success": False,
                "error": "Resultado no encontrado o acceso no autorizado."
            }), 404

        storage_path = model_query.data.get('model_storage_path')
        if not storage_path:
            logger.warning(f"[get_clustering_result_details] Modelo {model_id} sin ruta de almacenamiento")
            return jsonify({
                "success": False,
                "error": "El archivo de resultados no está disponible."
            }), 500

        # 2. Descargar y leer el archivo JSON desde Supabase Storage
        try:
            file_bytes = supabase_admin.storage.from_("proyectos-usuarios").download(storage_path)
            if not file_bytes:
                logger.error(f"[get_clustering_result_details] Archivo no encontrado en storage: {storage_path}")
                return jsonify({
                    "success": False,
                    "error": "No se pudo descargar el archivo de resultados."
                }), 500
        except Exception as storage_error:
            logger.error(f"[get_clustering_result_details] Error al acceder a storage: {str(storage_error)}", exc_info=True)
            return jsonify({
                "success": False,
                "error": "Error al acceder al archivo de resultados."
            }), 500

        # 3. Parsear el JSON
        try:
            # Supabase retorna bytes → decodificamos
            if isinstance(file_bytes, (bytes, bytearray)):
                file_str = file_bytes.decode("utf-8")
            else:
                file_str = str(file_bytes)

            result_data = json.loads(file_str)
        except json.JSONDecodeError as json_err:
            logger.error(f"[get_clustering_result_details] Archivo corrupto en {storage_path}: {str(json_err)}", exc_info=True)
            return jsonify({
                "success": False,
                "error": "El archivo de resultados está dañado o no es un JSON válido."
            }), 500

        # 4. Respuesta final
        return jsonify({
            "success": True,
            "data": result_data
        }), 200

    except Exception as e:
        logger.error(f"[get_clustering_result_details] Error inesperado: {str(e)}", exc_info=True)
        return jsonify({
            "success": False,
            "error": "Error interno del servidor."
        }), 500


def _load_model_from_db(user_id: str, model_id: str):
    """
    Carga un modelo de visión y sus artefactos desde Supabase.
    Abstrae la lógica de buscar metadata y descargar desde Storage.
    """
    model_meta_query = supabase_admin.table('trained_models').select('model_storage_path').eq('id', model_id).eq('user_id', user_id).single().execute()
    
    if not model_meta_query.data:
        raise FileNotFoundError(f"Modelo con id '{model_id}' no encontrado o sin permisos.")

    model_storage_path = model_meta_query.data.get('model_storage_path')
    if not model_storage_path:
        raise ValueError("La metadata del modelo no tiene una ruta de almacenamiento válida.")

    try:
        artifacts_bytes = supabase_admin.storage.from_("proyectos-usuarios").download(path=model_storage_path)
    except Exception as e:
        logger.error(f"No se pudo descargar el artefacto desde {model_storage_path}: {e}")
        raise FileNotFoundError(f"El archivo del modelo en la ruta '{model_storage_path}' no se pudo encontrar.")

    return vision_prediction_service.load_model_from_artifacts(artifacts_bytes)

def _prepare_dataset_in_tempdir(image_records: list) -> Path:
    """
    Crea un directorio temporal con la estructura (clase/imagen.png)
    necesaria para ImageFolder, descargando las imágenes desde el bucket 'visuales'.
    """
    temp_dir = Path(tempfile.mkdtemp(prefix="vision_eval_"))
    logger.info(f"📁 Creando dataset de evaluación temporal en: {temp_dir}")

    for record in image_records:
        # Asumimos que la primera etiqueta es la clase correcta
        tags = record.get("tags")
        storage_path = record.get("storage_path")

        if not tags or not isinstance(tags, list) or not tags[0] or not storage_path:
            logger.warning(f"Registro omitido por falta de tags/ruta: {record.get('id')}")
            continue

        label = tags[0].get('name') if isinstance(tags[0], dict) else tags[0]
        if not label:
            continue
            
        label_dir = temp_dir / str(label)
        label_dir.mkdir(exist_ok=True)

        try:
            # --- ✅ LA CORRECCIÓN ESTÁ AQUÍ ---
            # Le decimos explícitamente que descargue desde el bucket 'visuales'
            image_bytes = supabase_admin.storage.from_("visuales").download(path=storage_path)
            # --- FIN DE LA CORRECCIÓN ---

            if image_bytes:
                img = Image.open(BytesIO(image_bytes)).convert("RGB")
                # Usamos el nombre del archivo original para evitar colisiones
                file_name = Path(storage_path).name
                img.save(label_dir / file_name, "PNG") # Guardamos como PNG para consistencia
        except Exception as e:
            # Añadimos más detalles al log de error
            logger.error(f"❌ No se pudo descargar o procesar {storage_path} desde el bucket 'visuales': {e}")

    return temp_dir


# ================================================================
#   ENDPOINT: PREDICCIÓN INTERACTIVA CON MODELO DE VISIÓN
# ================================================================
# --- FUNCIÓN AUXILIAR ROBUSTA ---
def _load_model_from_db_or_cache(user_id: str, model_id_or_temp_id: str):
    """
    Carga un modelo de visión y sus artefactos, ya sea desde la base de datos
    (con un UUID real) o desde la caché temporal del servidor (prefijo 'temp_').

    Args:
        user_id (str): ID del usuario autenticado.
        model_id_or_temp_id (str): Puede ser un UUID de la tabla `trained_models`
                                   o un ID temporal de la sesión de entrenamiento.
    Returns:
        Modelo cargado por vision_prediction_service.
    Raises:
        FileNotFoundError: Si el modelo no existe o el artefacto no está disponible.
        ValueError: Si la metadata está incompleta o corrupta.
    """

    # --- 1. Cargar desde caché temporal ---
    if isinstance(model_id_or_temp_id, str) and model_id_or_temp_id.startswith("temp_"):
        logger.info(f"[ModelLoad] Intentando cargar modelo temporal desde caché con ID: {model_id_or_temp_id}")

        cache_entry = TEMP_TRAINING_CACHE.get(model_id_or_temp_id)
        if not cache_entry:
            raise FileNotFoundError(f"Sesión de entrenamiento '{model_id_or_temp_id}' expirada o inválida.")

        temp_file_path = Path(cache_entry.get("path", ""))
        if not temp_file_path.exists():
            raise FileNotFoundError("Archivo de modelo temporal no encontrado en el servidor.")

        try:
            with open(temp_file_path, "rb") as f:
                artifacts_bytes = f.read()
            logger.info(f"[ModelLoad] Modelo temporal cargado exitosamente desde {temp_file_path}")
            return vision_prediction_service.load_model_from_artifacts(artifacts_bytes)
        except Exception as e:
            logger.exception(f"[ModelLoad] Error leyendo archivo temporal {temp_file_path}: {e}")
            raise FileNotFoundError("No se pudo cargar el archivo del modelo temporal.")

    # --- 2. Cargar desde DB (Supabase) ---
    logger.info(f"[ModelLoad] Intentando cargar modelo persistente {model_id_or_temp_id} para usuario {user_id}")

    try:
        model_meta_query = (
            supabase_admin.table("trained_models")
            .select("model_storage_path")
            .eq("id", model_id_or_temp_id)
            .eq("user_id", user_id)
            .single()
            .execute()
        )
    except Exception as e:
        # Capturamos específicamente el error "0 rows returned"
        if getattr(e, "code", None) == "PGRST116":
            raise FileNotFoundError(f"Modelo con id '{model_id_or_temp_id}' no encontrado en la base de datos.")
        logger.exception(f"[ModelLoad] Error consultando metadata en DB: {e}")
        raise FileNotFoundError("Error al acceder a la base de datos para obtener metadata del modelo.")

    if not model_meta_query or not model_meta_query.data:
        raise FileNotFoundError(f"Modelo con id '{model_id_or_temp_id}' no encontrado o sin permisos.")

    model_storage_path = model_meta_query.data.get("model_storage_path")
    if not model_storage_path:
        raise ValueError("La metadata del modelo no contiene una ruta de almacenamiento válida.")

    try:
        artifacts_bytes = supabase_admin.storage.from_("proyectos-usuarios").download(path=model_storage_path)
        if not artifacts_bytes:
            raise FileNotFoundError(f"No se pudo descargar artefacto en {model_storage_path}")
        logger.info(f"[ModelLoad] Modelo descargado exitosamente desde {model_storage_path}")
        return vision_prediction_service.load_model_from_artifacts(artifacts_bytes)
    except Exception as e:
        logger.error(f"[ModelLoad] No se pudo descargar el artefacto en {model_storage_path}: {e}")
        raise FileNotFoundError(f"El archivo del modelo en la ruta '{model_storage_path}' no se pudo encontrar.")



@app.route('/api/models/vision/<string:model_id>/predict', methods=['POST'])
@token_required
def predict_with_vision_model_endpoint(current_user, model_id):
    """
    Realiza predicciones para una o más imágenes subidas directamente por el usuario.
    Recibe los archivos de imagen en un campo 'images' de un form-data.
    """

    MAX_IMAGES = 20
    MAX_FILE_SIZE_MB = 10

    if 'images' not in request.files:
        return jsonify({"success": False, "error": "No se encontraron archivos 'images' en la solicitud."}), 400
    
    image_files = request.files.getlist('images')
    if not image_files or image_files[0].filename == '':
        return jsonify({"success": False, "error": "No se seleccionó ninguna imagen."}), 400

    # 🚨 Validaciones de seguridad
    if len(image_files) > MAX_IMAGES:
        return jsonify({
            "success": False, 
            "error": f"Se permiten máximo {MAX_IMAGES} imágenes por solicitud."
        }), 400

    for f in image_files:
        f.seek(0, 2)  # Ir al final del archivo
        size_mb = f.tell() / (1024 * 1024)
        f.seek(0)  # Volver al inicio
        if size_mb > MAX_FILE_SIZE_MB:
            return jsonify({
                "success": False, 
                "error": f"La imagen {f.filename} supera el límite de {MAX_FILE_SIZE_MB} MB."
            }), 400

    temp_images_path = None
    try:
        # 1. Cargar el modelo entrenado y sus artefactos
        model, class_names, transforms = _load_model_from_db_or_cache(current_user.id, model_id)

        # 2. Guardar las imágenes subidas en un directorio temporal
        temp_images_path = Path(tempfile.mkdtemp(prefix="vision_predict_"))
        image_paths_to_predict = []
        for image_file in image_files:
            file_path = temp_images_path / image_file.filename
            image_file.save(file_path)
            image_paths_to_predict.append(file_path)

        # 3. Realizar predicciones usando tu servicio de visión
        predictions = vision_prediction_service.predict_images(
            model=model,
            image_paths=image_paths_to_predict,
            transforms=transforms,
            class_names=class_names
        )
        
        return jsonify({"success": True, "predictions": predictions}), 200

    except FileNotFoundError as e:
        return jsonify({"success": False, "error": str(e)}), 404
    except ValueError as e:
        return jsonify({"success": False, "error": str(e)}), 400
    except Exception as e:
        logger.error(f"Error prediciendo con modelo de visión {model_id}: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor durante la predicción."}), 500
    finally:
        # 4. Limpieza: Asegurarnos de borrar el directorio temporal de imágenes
        if temp_images_path and temp_images_path.exists():
            logger.info(f"🧹 Limpiando directorio de predicción temporal: {temp_images_path}")
            shutil.rmtree(temp_images_path)


# ================================================================
#   ENDPOINT: EVALUAR UN MODELO DE VISIÓN GUARDADO
# ================================================================
@app.route('/api/models/vision/<string:model_id>/evaluate', methods=['POST'])
@token_required
def evaluate_vision_model_endpoint(current_user, model_id):
    """
    Evalúa un modelo de visión contra un dataset existente en la plataforma.
    Recibe: { "evaluation_dataset_id": "uuid-del-dataset" }
    """

    MAX_IMAGES = 500         # Límite más alto que predicción porque suelen ser datasets
    MAX_FILE_SIZE_MB = 10    # Igual que en predicción

    data = request.get_json()
    if not data or 'evaluation_dataset_id' not in data:
        return jsonify({"success": False, "error": "Falta el parámetro 'evaluation_dataset_id'."}), 400

    evaluation_dataset_id = data['evaluation_dataset_id']
    temp_dataset_path = None
    preview_list = []  # <-- Para almacenar URLs firmadas de imágenes evaluadas

    try:
        # 1. Cargar el modelo entrenado
        model, class_names, transforms = _load_model_from_db_or_cache(current_user.id, model_id)

        # 2. Obtener los registros de imágenes del dataset de evaluación
        image_records = supabase_handler.get_visuals_by_dataset(evaluation_dataset_id, current_user.id)

        if not image_records:
            return jsonify({"success": False, "error": "El dataset de evaluación no contiene imágenes etiquetadas."}), 404

        # 🚨 Validaciones
        if len(image_records) > MAX_IMAGES:
            return jsonify({
                "success": False,
                "error": f"El dataset excede el máximo permitido de {MAX_IMAGES} imágenes."
            }), 400

        for rec in image_records:
            storage_path = rec.get("storage_path")
            try:
                # Descargar imagen desde bucket "visuales"
                image_bytes = supabase_admin.storage.from_("visuales").download(path=storage_path)
                if not image_bytes:
                    continue

                # Validar tamaño máximo permitido
                size_mb = len(image_bytes) / (1024 * 1024)
                if size_mb > MAX_FILE_SIZE_MB:
                    return jsonify({
                        "success": False,
                        "error": f"La imagen en {storage_path} supera el límite de {MAX_FILE_SIZE_MB} MB."
                    }), 400

                # Generar URL firmada (preview)
                signed_url_response = supabase_admin.storage.from_("visuales").create_signed_url(storage_path, 600)
                preview_list.append({
                    "storage_path": storage_path,
                    "image_url": signed_url_response.get("signedURL")
                })

            except Exception as e:
                logger.warning(f"⚠️ No se pudo verificar/firmar imagen {storage_path}: {e}")

        # 3. Preparar el dataset temporal
        temp_dataset_path = _prepare_dataset_in_tempdir(image_records)

        if not any(temp_dataset_path.iterdir()):
            return jsonify({"success": False, "error": "No se pudo preparar el dataset. Verifique las etiquetas de las imágenes."}), 500

        # 4. Ejecutar la evaluación
        results = vision_prediction_service.evaluate_model(
            model=model,
            dataset_path=temp_dataset_path,
            transforms=transforms,
            class_names=class_names
        )

        return jsonify({
            "success": True,
            "evaluation_results": results,
            "previews": preview_list  # <-- Se devuelven métricas + URLs firmadas
        }), 200

    except FileNotFoundError as e:
        return jsonify({"success": False, "error": str(e)}), 404
    except ValueError as e:
        return jsonify({"success": False, "error": str(e)}), 400
    except Exception as e:
        logger.error(f"Error evaluando modelo {model_id}: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor durante la evaluación."}), 500
    finally:
        # 5. Limpieza del directorio temporal
        if temp_dataset_path and temp_dataset_path.exists():
            logger.info(f"🧹 Limpiando directorio temporal: {temp_dataset_path}")
            shutil.rmtree(temp_dataset_path)


# ================================================================
#   ENDPOINT: PREDICCIÓN BATCH CON MODELO DE VISIÓN
# ================================================================

@app.route('/api/models/vision/<string:model_id>/batch-predict', methods=['POST'])
@token_required
def batch_predict_vision_model_endpoint(current_user, model_id):
    """
    Realiza predicciones para todas las imágenes de un dataset de visión existente.
    Recibe un JSON con el ID del dataset.
    """
    try:
        data = request.get_json(silent=True) or {}
        dataset_id = data.get("dataset_id")
        if not dataset_id:
            return jsonify({"success": False, "error": "Falta el parámetro 'dataset_id'."}), 400

        # 1. Cargar el modelo (desde DB o caché)
        try:
            model, class_names, transforms = _load_model_from_db_or_cache(current_user.id, model_id)
            model.eval()
            device = next(model.parameters()).device
        except FileNotFoundError as e:
            return jsonify({"success": False, "error": str(e)}), 404
        except Exception as e:
            logger.error(f"[BatchPredict] Error cargando modelo {model_id}: {e}", exc_info=True)
            return jsonify({"success": False, "error": "No se pudo cargar el modelo."}), 500

        # 2. Obtener imágenes del dataset
        image_records = supabase_handler.get_visuals_by_dataset(dataset_id, current_user.id)
        if not image_records:
            return jsonify({"success": False, "error": "El dataset no contiene imágenes."}), 404

        predictions_list = []

        # 3. Iterar imágenes y predecir
        for record in image_records:
            storage_path = record.get("storage_path")
            if not storage_path:
                predictions_list.append({"error": "Imagen sin ruta válida."})
                continue

            try:
                # Descargar imagen desde el bucket correcto
                image_bytes = supabase_admin.storage.from_("visuales").download(path=storage_path)
                if not image_bytes:
                    raise ValueError("No se pudo descargar la imagen.")

                # Preprocesar imagen
                img = Image.open(BytesIO(image_bytes)).convert("RGB")
                input_tensor = transforms(img).unsqueeze(0).to(device)

                # Predicción
                with torch.no_grad():
                    outputs = model(input_tensor)
                    probs = torch.softmax(outputs, dim=1)
                    confidence, pred_idx = torch.max(probs, 1)
                    predicted_class = class_names[pred_idx.item()]

                # Generar URL firmada (del bucket 'visuales')
                signed_url_response = supabase_admin.storage.from_("visuales").create_signed_url(storage_path, 600)
                image_url = signed_url_response.get("signedURL")

                predictions_list.append({
                    "image_url": image_url,
                    "original_path": storage_path,
                    "predicted_class": predicted_class,
                    "confidence": round(float(confidence.item()), 4)
                })

            except Exception as e:
                logger.error(f"[BatchPredict] Error procesando imagen {storage_path}: {e}", exc_info=True)
                predictions_list.append({
                    "original_path": storage_path,
                    "error": "No se pudo procesar la imagen."
                })

        return jsonify({"success": True, "data": predictions_list}), 200

    except Exception as e:
        logger.error(f"[BatchPredict] Error inesperado para modelo {model_id}: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500




# Config
MAX_IMAGE_SIZE_MB = 10
ALLOWED_EXTENSIONS = {"jpg", "jpeg", "png", "bmp", "tiff", "webp"}

def is_allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


# --- Endpoint 1: Análisis Profundo de Imagen ---
@app.route("/api/analyze-image-deep", methods=["POST"])
@token_required
def handle_analyze_image_deep(current_user):
    """Análisis profundo y completo de una ÚNICA imagen."""
    request_id = str(uuid.uuid4())
    try:
        # --- 1. Validar entrada ---
        if "image" not in request.files:
            raise BadRequest("No se proporcionó ningún archivo de imagen.")

        file = request.files["image"]
        filename = secure_filename(file.filename)

        if not is_allowed_file(filename):
            raise BadRequest("Tipo de archivo no permitido.")

        image_bytes = file.read()
        if len(image_bytes) > MAX_IMAGE_SIZE_MB * 1024 * 1024:
            raise BadRequest(f"La imagen excede el tamaño máximo de {MAX_IMAGE_SIZE_MB}MB.")

        # --- 2. Llamar al Orquestador ---
        logger.info(f"[{request_id}] 🔬 Análisis profundo de imagen '{filename}' (usuario={current_user.id})")
        analysis_result = analisis_completo_de_imagen(
            imagen_bytes=image_bytes,
            filename=filename,
            model_manager_instance=model_manager
        )

        if not analysis_result:
            raise InternalServerError("El análisis no produjo resultados.")

        logger.info(f"[{request_id}] ✅ Análisis completado para '{filename}'")

        # --- 3. Respuesta ---
        return jsonify({
            "success": True,
            "analysis": analysis_result,
            "request_id": request_id
        })

    except BadRequest as e:
        logger.warning(f"[{request_id}] ⚠️ Error de validación: {e}")
        return jsonify({"success": False, "error": str(e), "request_id": request_id}), 400
    except Exception as e:
        logger.exception(f"[{request_id}] ❌ Error inesperado en /api/analyze-image-deep: {e}")
        return jsonify({"success": False, "error": "Error interno del servidor.", "request_id": request_id}), 500


# --- Endpoint 2: Análisis Profundo de PDF ---
@app.route("/api/analyze-pdf", methods=["POST"])
@token_required
def handle_analyze_pdf(current_user):
    """Análisis profundo de un archivo PDF, página por página."""
    request_id = str(uuid.uuid4())
    try:
        if "pdf_file" not in request.files:
            raise BadRequest("No se proporcionó ningún archivo PDF.")

        file = request.files["pdf_file"]
        filename = secure_filename(file.filename)

        if not filename.lower().endswith(".pdf"):
            raise BadRequest("El archivo no es un PDF válido.")

        pdf_bytes = file.read()
        logger.info(f"[{request_id}] 📄 Iniciando análisis del PDF '{filename}' (usuario={current_user.id})")

        pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        pages_analysis = []

        for page_num in range(len(pdf_doc)):
            page = pdf_doc.load_page(page_num)
            pix = page.get_pixmap()
            image_bytes = pix.tobytes("png")

            analysis_result = analisis_completo_de_imagen(
                imagen_bytes=image_bytes,
                filename=f"{filename}_page_{page_num + 1}",
                model_manager_instance=model_manager
            )

            pages_analysis.append({
                "page_number": page_num + 1,
                "analysis": analysis_result
            })

        pdf_doc.close()

        logger.info(f"[{request_id}] ✅ Análisis del PDF '{filename}' completado con {len(pages_analysis)} páginas.")

        return jsonify({
            "success": True,
            "pdf_analysis": {
                "filename": filename,
                "total_pages": len(pages_analysis),
                "pages": pages_analysis
            },
            "request_id": request_id
        })

    except BadRequest as e:
        logger.warning(f"[{request_id}] ⚠️ Error de validación PDF: {e}")
        return jsonify({"success": False, "error": str(e), "request_id": request_id}), 400
    except Exception as e:
        logger.exception(f"[{request_id}] ❌ Error inesperado en /api/analyze-pdf: {e}")
        return jsonify({"success": False, "error": "Error interno al procesar el PDF.", "request_id": request_id}), 500


# ================================================================
#  extraccion text de pdf
# ================================================================


# --- Función Auxiliar (la que propusimos antes) ---
def analizar_pagina_de_catalogo(texto_pagina):
    items_encontrados = []
    # Este patrón es un ejemplo, puedes hacerlo tan complejo como necesites
    patron = re.compile(r"(.*?)\s*SKU:\s*([\w-]+)\s*.*?Precio:\s*\$([\d.]+)", re.DOTALL | re.IGNORECASE)
    matches = patron.findall(texto_pagina)
    for match in matches:
        descripcion, sku, precio = match
        items_encontrados.append({
            "item_id": f"item_{uuid.uuid4()}",
            "descripcion": descripcion.strip(),
            "sku": sku.strip(),
            "precio": precio.strip()
        })
    return items_encontrados

# --- TU ENDPOINT MODIFICADO ---





MAX_PDF_SIZE_MB = 20  # límite de tamaño de PDF en MB

@app.route("/api/pdf/extract-images", methods=["POST"])
@token_required
def handle_extract_images(current_user):
    """
    Endpoint robusto para extraer todas las imágenes de un PDF.
    Retorna las imágenes codificadas en base64, con número de página y extensión.
    """
    request_id = str(uuid.uuid4())
    try:
        # --- Validaciones ---
        if "pdf_file" not in request.files:
            raise BadRequest("No se proporcionó ningún archivo PDF.")

        file = request.files["pdf_file"]
        file.seek(0, 2)  # ir al final del archivo
        size_mb = file.tell() / (1024 * 1024)
        file.seek(0)  # volver al inicio

        if size_mb > MAX_PDF_SIZE_MB:
            raise BadRequest(f"El archivo supera el límite permitido de {MAX_PDF_SIZE_MB} MB.")

        pdf_bytes = file.read()
        logger.info(f"[{request_id}] 🖼️ Extracción de imágenes iniciada para PDF '{file.filename}'")

        # --- Extracción de imágenes ---
        pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        extracted_images = []

        for page_num in range(len(pdf_doc)):
            page = pdf_doc.load_page(page_num)
            image_list = page.get_images(full=True)

            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = pdf_doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                image_base64 = base64.b64encode(image_bytes).decode("utf-8")

                extracted_images.append({
                    "id": f"page{page_num+1}_img{img_index+1}",
                    "page_number": page_num + 1,
                    "extension": image_ext,
                    "image_base64": image_base64
                })

        pdf_doc.close()

        if not extracted_images:
            logger.info(f"[{request_id}] ⚠️ No se encontraron imágenes en el PDF '{file.filename}'")

        logger.info(f"[{request_id}] ✅ Extracción completada: {len(extracted_images)} imágenes encontradas")

        return jsonify({
            "success": True,
            "filename": file.filename,
            "image_count": len(extracted_images),
            "images": extracted_images,
            "request_id": request_id
        })

    except BadRequest as e:
        logger.warning(f"[{request_id}] ⚠️ Error de validación: {e}")
        return jsonify({"success": False, "error": str(e), "request_id": request_id}), 400
    except Exception as e:
        logger.exception(f"[{request_id}] ❌ Error interno: {e}")
        return jsonify({"success": False, "error": "Error interno al extraer las imágenes.", "request_id": request_id}), 500



# ================================================================
#   ENDPOINT: Guardar o actualizar análisis (Versión Optimizada)
# ================================================================
@app.route("/api/save-analysis", methods=["PUT"])
@token_required
def handle_save_analysis_optimized(current_user):
    """
    Guarda o actualiza un análisis de un dataset en Supabase usando upsert.
    """
    request_id = str(uuid.uuid4())
    try:
        data = request.get_json()
        if not data:
            raise BadRequest("No se proporcionaron datos.")

        dataset_id = data.get("dataset_id")
        name = data.get("name") or "Análisis sin nombre"
        pages = data.get("pages") or []

        if not dataset_id:
            raise BadRequest("dataset_id es obligatorio.")

        supabase = supabase_handler.get_db_client()

        # 🚀 Operación UPSERT en una sola llamada
        # on_conflict le dice a la DB qué hacer si la restricción de unicidad (user_id, dataset_id) falla.
        # En este caso, simplemente actualizará el registro existente.
        result = (
            supabase.table("analyses")
            .upsert(
                {
                    "dataset_id": dataset_id,
                    "user_id": current_user.id,
                    "name": name,
                    "pages": pages,
                    # 'updated_at' podría manejarse automáticamente por la base de datos (ver sugerencia #2)
                },
                on_conflict="user_id,dataset_id", 
            )
            .execute()
        )

        if not result.data:
            raise Exception("Falló la operación de guardado/actualización en Supabase.")

        analysis_id = result.data[0]["id"]
        logger.info(
            f"[{request_id}] ✅ Análisis guardado/actualizado (id={analysis_id}, usuario={current_user.id})"
        )

        return jsonify({
            "success": True,
            "message": "Análisis guardado correctamente.",
            "analysis_id": analysis_id,
            "request_id": request_id,
        })
    # ... (el resto del manejo de errores es el mismo)
    except BadRequest as e:
        logger.warning(f"[{request_id}] ⚠️ Error de validación: {e}")
        return jsonify({"success": False, "error": str(e), "request_id": request_id}), 400
    except Exception as e:
        logger.exception(f"[{request_id}] ❌ Error inesperado en /api/save-analysis: {e}")
        return jsonify({
            "success": False,
            "error": "Error interno al guardar el análisis.",
            "request_id": request_id
        }), 500




# ============================================================
#  ENDPOINT: DESCARGA DE ARCHIVO DE UN DATASET
# ============================================================
@app.route('/api/datasets/<string:dataset_id>/download', methods=['GET'])
@token_required
def download_dataset_file(current_user, dataset_id):
    """
    Descarga el archivo físico de un dataset desde el storage.
    Esta pieza permite que el editor de PDF cargue el contenido real.
    """
    request_id = str(uuid.uuid4())
    BUCKET_NAME = "proyectos-usuarios"  # Verifica que coincida con tu bucket real en Supabase

    try:
        # 1. Validación básica del dataset_id
        if not dataset_id or not dataset_id.strip():
            logger.warning(f"[{request_id}] ❌ dataset_id vacío o inválido.")
            return jsonify({"success": False, "error": "El ID del dataset es inválido."}), 400

        # 2. Buscar el dataset en la base de datos
        response = (
            supabase_admin.table("datasets")
            .select("storage_path")
            .eq("dataset_id", dataset_id)
            .eq("user_id", current_user.id)
            .single()
            .execute()
        )

        if not response or not response.data:
            logger.warning(f"[{request_id}] ❌ Dataset {dataset_id} no encontrado para usuario {current_user.id}.")
            return jsonify({"success": False, "error": "Dataset no encontrado."}), 404

        storage_path = response.data.get("storage_path")
        if not storage_path:
            logger.warning(f"[{request_id}] ❌ El dataset {dataset_id} no tiene un storage_path registrado.")
            return jsonify({"success": False, "error": "Ruta de archivo no registrada en el dataset."}), 404

        logger.info(f"[{request_id}] 📂 Intentando descargar archivo desde: {storage_path}")

        # 3. Descargar el archivo desde Supabase Storage
        try:
            file_bytes = supabase_admin.storage.from_(BUCKET_NAME).download(storage_path)
        except Exception as e:
            logger.error(f"[{request_id}] ⚠️ Error al descargar archivo desde Supabase Storage: {e}", exc_info=True)
            return jsonify({"success": False, "error": "No se pudo descargar el archivo del almacenamiento."}), 502

        if not file_bytes:
            logger.warning(f"[{request_id}] ❌ Archivo vacío o inexistente en storage_path={storage_path}.")
            return jsonify({"success": False, "error": "El archivo no existe o está vacío en el almacenamiento."}), 404

        # 4. Responder con el archivo
        filename = storage_path.split("/")[-1]
        logger.info(f"[{request_id}] ✅ Descarga lista para el archivo: {filename}")

        return send_file(
            BytesIO(file_bytes),
            mimetype="application/octet-stream",  # Tipo genérico
            as_attachment=True,
            download_name=filename,
        )

    except Exception as e:
        logger.error(f"[{request_id}] 💥 Error interno al descargar dataset {dataset_id}: {e}", exc_info=True)
        return jsonify({
            "success": False,
            "error": "Error interno al descargar el archivo.",
            "request_id": request_id
        }), 500


MAX_PDF_SIZE_MB = 20  # ejemplo de límite de tamaño

@app.route("/api/pdf/extract-structured-data", methods=["POST"])
@token_required
def handle_extract_structured_data(current_user):
    """
    Orquesta el análisis de un PDF, llamando a la función especializada de Gemini para cada página.
    """
    # ... (tus validaciones de archivo PDF no cambian)
    file = request.files["pdf_file"]
    pdf_bytes = file.read()
    
    pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    pages_analysis = []

    for page_num in range(len(pdf_doc)):
        page = pdf_doc.load_page(page_num)
        image_bytes = page.get_pixmap().tobytes("png")
        
        # --- LLAMAMOS A NUESTRA NUEVA HERRAMIENTA ---
        ia_result = extraer_datos_estructurados_con_gemini(image_bytes)
        
        # Codificamos la imagen para el frontend
        image_base64 = base64.b64encode(image_bytes).decode('utf-8')
        
        pages_analysis.append({
            "page_number": page_num + 1,
            "analysis": {
                "image_base64": image_base64,
                # El resultado de la IA ya tiene la clave "extracted_items"
                "extracted_items": ia_result.get("extracted_items", [])
            }
        })

    pdf_doc.close()

    return jsonify({
        "success": True,
        "pdf_analysis": {
            "filename": file.filename,
            "total_pages": len(pages_analysis),
            "pages": pages_analysis
        }
    })


# --- Directorio temporal global y compartido ---
TEMP_DIR = "/tmp/pdf_service"
os.makedirs(TEMP_DIR, exist_ok=True)

MAX_PDF_SIZE = 20 * 1024 * 1024  # 20MB


# ================================================================
#   ENDPOINT DE RECONSTRUCCIÓN DE PDF
# ================================================================
@app.route("/api/pdf/reconstruct", methods=["POST", "GET"])
@token_required
def reconstruct_pdf(current_user):
    """
    Reconstruye un PDF en páginas:
      - POST: sube un PDF y devuelve su doc_id
      - GET: recibe ?doc_id=xxx&page=n y devuelve la estructura de esa página
    """
    logger.info("--- Endpoint /reconstruct (POST) INICIADO ---") # <-- Log de vida
    try:
        # --- SUBIDA DEL PDF (POST) ---
        if request.method == "POST":
            if "pdf_file" not in request.files:
                raise BadRequest("No se proporcionó PDF.")

            file = request.files["pdf_file"]
            pdf_bytes = file.read()

            if len(pdf_bytes) > MAX_PDF_SIZE:
                raise BadRequest("El PDF supera el límite de 20MB.")

            # Generar ID único y guardar en disco
            doc_id = str(uuid.uuid4())
            filepath = os.path.join(TEMP_DIR, f"{doc_id}.pdf")
            with open(filepath, "wb") as f:
                f.write(pdf_bytes)

            # Abrimos el documento solo para obtener metadatos
            pdf_doc = fitz.open(filepath)
            logger.info("Procesamiento del PDF completado.") # <-- Log de vida
            return jsonify({
                "success": True,
                "doc_id": doc_id,
                "pages": len(pdf_doc),
                "width": pdf_doc[0].rect.width,
                "height": pdf_doc[0].rect.height
            })
        
        # --- RECONSTRUCCIÓN DE PÁGINA (GET) ---
        elif request.method == "GET":
            doc_id = request.args.get("doc_id")
            page_num = int(request.args.get("page", 1)) - 1

            filepath = os.path.join(TEMP_DIR, f"{doc_id}.pdf")
            if not doc_id or not os.path.exists(filepath):
                raise BadRequest("doc_id inválido o no encontrado.")

            pdf_doc = fitz.open(filepath)

            if page_num < 0 or page_num >= len(pdf_doc):
                raise BadRequest("Número de página fuera de rango.")

            page = pdf_doc.load_page(page_num)
            elements = []

            # A. Extraer texto
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block["type"] == 0:  # texto
                    for line in block["lines"]:
                        for span in line["spans"]:
                            elements.append({
                                "id": f"text_{uuid.uuid4()}",
                                "type": "text",
                                "content": span["text"],
                                "bbox": span["bbox"],
                                "font_size": round(span["size"]),
                                "font_family": span["font"]
                            })

            # B. Extraer imágenes
            images = page.get_images(full=True)
            for img in images:
                xref = img[0]
                image_bbox = page.get_image_bbox(img)
                elements.append({
                    "id": f"img_{uuid.uuid4()}",
                    "type": "image",
                    "xref": xref,
                    "doc_id": doc_id,
                    "bbox": tuple(image_bbox) 
                })
            logger.info("Iniciando procesamiento del PDF...") # <-- Log de vida
            return jsonify({
                "success": True,
                "page_number": page_num + 1,
                "width": page.rect.width,
                "height": page.rect.height,
                "elements": elements
            })
            logger.info("--- Endpoint /reconstruct (POST) FINALIZADO ---") 
    except Exception as e:
        print(f"[ERROR reconstruct_pdf] {e}")
        return jsonify({"success": False, "error": str(e)}), 500


# ================================================================
#   ENDPOINT PARA OBTENER IMÁGENES DE UN PDF
# ================================================================
@app.route("/api/pdf/image/<string:doc_id>/<int:xref>", methods=["GET"])
@token_required
def get_pdf_image(current_user, doc_id, xref):
    """
    Devuelve una imagen de un PDF:
    - Como stream JPEG optimizado (más eficiente)
    - (opcional: devolver en base64 si lo pide el front)
    """
    try:
        filepath = os.path.join(TEMP_DIR, f"{doc_id}.pdf")
        if not os.path.exists(filepath):
            raise BadRequest("doc_id inválido o no encontrado.")

        pdf_doc = fitz.open(filepath)
        base_image = pdf_doc.extract_image(xref)
        image_bytes = base_image["image"]

        # Comprimir con Pillow
        image = Image.open(io.BytesIO(image_bytes))
        output = io.BytesIO()
        image.save(output, format="JPEG", optimize=True, quality=80)
        output.seek(0)

        # --- Opción A: devolver como stream ---
        return send_file(output, mimetype="image/jpeg")

        # --- Opción B: devolver en base64 ---
        # image_base64 = base64.b64encode(output.getvalue()).decode("utf-8")
        # return jsonify({"success": True, "image": image_base64})

    except Exception as e:
        print(f"[ERROR get_pdf_image] {e}")
        return jsonify({"success": False, "error": str(e)}), 500


# ================================================================
#   LIMPIEZA AUTOMÁTICA DE ARCHIVOS TEMPORALES
# ================================================================
def cleanup_temp_files():
    print("🧹 Ejecutando limpieza de archivos temporales...")
    now = time.time()
    for filename in os.listdir(TEMP_DIR):
        filepath = os.path.join(TEMP_DIR, filename)
        try:
            # Archivos de más de 24h → borrar
            if now - os.path.getmtime(filepath) > 86400:
                os.remove(filepath)
                print(f"🗑️ Eliminado archivo antiguo: {filename}")
        except Exception as e:
            print(f"[CLEANUP ERROR] {e}")

# Scheduler cada 6h
scheduler = BackgroundScheduler()
scheduler.add_job(func=cleanup_temp_files, trigger="interval", hours=6)
scheduler.start()
atexit.register(lambda: scheduler.shutdown())



# ============================================================
#  ENDPOINT: ACTUALIZAR ELEMENTO INDIVIDUAL
# ============================================================
@app.route('/api/documents/<string:dataset_id>/element/<string:element_id>', methods=['PUT'])
@token_required
def update_document_element(current_user, dataset_id, element_id):
    """
    Actualiza parcialmente un elemento del documento.
    Ejemplo de payload:
    {
      "text": "nueva palabra",
      "bbox": [100, 120, 200, 140]
    }
    """
    request_id = str(uuid.uuid4())
    try:
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "No se proporcionaron datos."}), 400

        supabase = supabase_handler.get_db_client()
        result = (
            supabase.table("document_elements")
            .update(data)
            .eq("dataset_id", dataset_id)
            .eq("element_id", element_id)
            .eq("user_id", current_user.id)
            .execute()
        )

        if not result.data:
            return jsonify({
                "success": False,
                "error": f"Elemento {element_id} no encontrado.",
                "request_id": request_id
            }), 404

        return jsonify({
            "success": True,
            "message": "Elemento actualizado correctamente.",
            "element": result.data[0],
            "request_id": request_id,
        })

    except Exception as e:
        logger.exception(f"[{request_id}] ❌ Error en update_document_element: {e}")
        return jsonify({
            "success": False,
            "error": "Error interno al actualizar elemento.",
            "request_id": request_id,
        }), 500


# ============================================================
#  ENDPOINT: CARGA DE ELEMENTOS DE UN DOCUMENTO (con bbox)
# ============================================================
@app.route('/api/documents/<string:dataset_id>/elements', methods=['GET'])
@token_required
def get_document_elements(current_user, dataset_id):
    """
    Devuelve la lista de elementos de un documento con metadatos:
    - id
    - texto
    - bounding box (coordenadas)
    - page_index
    - atributos adicionales
    """
    request_id = str(uuid.uuid4())
    try:
        supabase = supabase_handler.get_db_client()
        response = (
            supabase.table("document_elements")
            .select("*")
            .eq("dataset_id", dataset_id)
            .eq("user_id", current_user.id)
            .execute()
        )

        if not response.data:
            return jsonify({
                "success": True,
                "elements": [],
                "message": "Documento sin elementos almacenados.",
                "request_id": request_id,
            }), 200

        return jsonify({
            "success": True,
            "elements": response.data,
            "request_id": request_id,
        })

    except Exception as e:
        logger.exception(f"[{request_id}] ❌ Error en get_document_elements: {e}")
        return jsonify({
            "success": False,
            "error": "Error interno al cargar elementos del documento.",
            "request_id": request_id,
        }), 500



# ================================================================
# ENDPOINT: RUN CLUSTERING ANALYSIS (KMEANS / DBSCAN)
# ================================================================
@app.route('/api/datasets/<string:dataset_id>/run-clustering-analysis', methods=['POST'])
@token_required
def run_clustering_analysis_endpoint(current_user, dataset_id):
    """
    Ejecuta clustering (K-Means o DBSCAN) sobre un dataset tabular.
    Devuelve métricas, etiquetas y resumen sin guardar el modelo.
    """
    logger = app.logger
    logger.info(f"[CLUSTERING_ANALYSIS] Usuario={current_user.id} Dataset={dataset_id}")

    try:
        config = request.get_json(silent=True) or {}
        algorithm = config.get("algorithm")
        params = config.get("parameters", {})

        if algorithm not in {"kmeans", "dbscan"}:
            return jsonify({"success": False, "error": "Debe especificar 'algorithm' (kmeans o dbscan)."}), 400

        # --- 2. CARGAR DATASET ---
        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        if not dataset_info or "data" not in dataset_info:
            return jsonify({"success": False, "error": "Dataset no encontrado."}), 404

        storage_path = dataset_info["data"].get("storage_path")
        df = supabase_handler.load_file_as_dataframe(user_id=current_user.id, path=storage_path)

        if df is None or df.empty:
            return jsonify({"success": False, "error": "No se pudo cargar el dataset o está vacío."}), 400

        tabular_data = df.to_dict(orient='records')

        # --- 3. EJECUTAR SERVICIO ---
        result = None
        if algorithm == "kmeans":
            n_clusters = params.get("n_clusters", 5)
            if not isinstance(n_clusters, int) or n_clusters < 2:
                return jsonify({"success": False, "error": "'n_clusters' debe ser un entero >= 2."}), 400

            result = tabular_clustering_service.run_kmeans_clustering(
                tabular_data=tabular_data, n_clusters=n_clusters
            )

        elif algorithm == "dbscan":
            eps = params.get("eps", 0.5)
            min_samples = params.get("min_samples", 5)

            if not isinstance(eps, (int, float)) or eps <= 0:
                return jsonify({"success": False, "error": "'eps' debe ser un número positivo."}), 400
            if not isinstance(min_samples, int) or min_samples < 2:
                return jsonify({"success": False, "error": "'min_samples' debe ser un entero >= 2."}), 400

            result = tabular_clustering_service.run_dbscan_clustering(
                tabular_data=tabular_data, eps=eps, min_samples=min_samples
            )

        if not result or not result.get("success"):
            return jsonify(result or {"success": False, "error": "Error desconocido en el clustering."}), 500

        return jsonify(result), 200

    except Exception as e:
        logger.error(f"[CLUSTERING_ANALYSIS] Error fatal: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno durante el análisis de clustering."}), 500


# ================================================================
# ENDPOINT: ASISTENTE KMEANS (MÉTODO DEL CODO)
# ================================================================
@app.route('/api/datasets/<string:dataset_id>/suggest-kmeans-k', methods=['POST'])
@token_required
def suggest_kmeans_k_endpoint(current_user, dataset_id):
    """
    Sugiere valores de k para KMeans usando el método del codo.
    """
    logger = app.logger
    logger.info(f"[KMEANS_ASSIST_EP] Usuario={current_user.id} Dataset={dataset_id}")

    try:
        config = request.get_json(silent=True) or {}
        max_k = config.get("max_k", 15)

        if not isinstance(max_k, int) or max_k < 3:
            return jsonify({"success": False, "error": "'max_k' debe ser un entero >= 3."}), 400

        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        if not dataset_info or "data" not in dataset_info:
            return jsonify({"success": False, "error": "Dataset no encontrado."}), 404

        storage_path = dataset_info["data"].get("storage_path")
        df = supabase_handler.load_file_as_dataframe(user_id=current_user.id, path=storage_path)

        if df is None or df.empty:
            return jsonify({"success": False, "error": "No se pudo cargar el dataset o está vacío."}), 400

        tabular_data = df.to_dict(orient='records')
        result = tabular_clustering_service.suggest_kmeans_k(tabular_data=tabular_data, max_k=max_k)

        return jsonify(result), 200 if result.get("success") else 500

    except Exception as e:
        logger.error(f"[KMEANS_ASSIST_EP] Error fatal: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno al sugerir k para KMeans."}), 500


# ================================================================
# ENDPOINT: ASISTENTE DBSCAN (GRÁFICO DE DISTANCIA)
# ================================================================
@app.route('/api/datasets/<string:dataset_id>/suggest-dbscan-eps', methods=['POST'])
@token_required
def suggest_dbscan_eps_endpoint(current_user, dataset_id):
    """
    Sugiere un valor eps para DBSCAN en base al gráfico k-distances.
    """
    logger = app.logger
    logger.info(f"[DBSCAN_ASSIST_EP] Usuario={current_user.id} Dataset={dataset_id}")

    try:
        config = request.get_json(silent=True) or {}
        min_samples = config.get("min_samples")

        if not isinstance(min_samples, int) or min_samples < 2:
            return jsonify({"success": False, "error": "'min_samples' debe ser un entero >= 2."}), 400

        dataset_info = dataset_service.get_dataset_info_by_id(dataset_id, current_user.id)
        if not dataset_info or "data" not in dataset_info:
            return jsonify({"success": False, "error": "Dataset no encontrado."}), 404

        storage_path = dataset_info["data"].get("storage_path")
        df = supabase_handler.load_file_as_dataframe(user_id=current_user.id, path=storage_path)

        if df is None or df.empty:
            return jsonify({"success": False, "error": "No se pudo cargar el dataset o está vacío."}), 400

        tabular_data = df.to_dict(orient='records')
        result = tabular_clustering_service.suggest_dbscan_eps(
            tabular_data=tabular_data, min_samples=min_samples
        )

        return jsonify(result), 200 if result.get("success") else 500

    except Exception as e:
        logger.error(f"[DBSCAN_ASSIST_EP] Error fatal: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno al sugerir eps para DBSCAN."}), 500


# ================================================================
#   NUEVO ENDPOINT: ENTRENAR Y GUARDAR MODELO DE CLUSTERING
# ================================================================
@app.route('/api/datasets/<string:dataset_id>/train-clustering-model', methods=['POST'])
@token_required
def train_clustering_model_endpoint(current_user, dataset_id):
    """
    Entrena y guarda un modelo de clustering (KMeans o DBSCAN) en la base de datos
    y en el almacenamiento, asociándolo a un proyecto.
    """
    logger = app.logger
    logger.info(f"[CLUSTERING_TRAIN] Usuario={current_user.id} Dataset={dataset_id}")
    try:
        config = request.get_json()
        if not config:
            return jsonify({"success": False, "error": "No se recibió configuración JSON."}), 400

        # --- OBTENEMOS PARÁMETROS NECESARIOS ---
        algorithm = config.get("algorithm")
        parameters = config.get("parameters", {})
        project_id = config.get("project_id")
        model_display_name = config.get("model_display_name")
        project_name = config.get("project_name")

        # --- VALIDACIONES ---
        if not all([algorithm, project_id, model_display_name, project_name]):
            return jsonify({"success": False, "error": "Faltan parámetros obligatorios: 'algorithm', 'project_id', 'model_display_name', 'project_name'."}), 400

        # --- LLAMADA AL SERVICIO CON EL NOMBRE CORRECTO ---
        result = tabular_clustering_service.train_and_save_model( # <-- NOMBRE CORREGIDO
            user_id=current_user.id,
            dataset_id=dataset_id,
            algorithm=algorithm,
            parameters=parameters,
            project_id=project_id,
            model_display_name=model_display_name,
            project_name=project_name
        )

        status_code = 201 if result.get("success") else 400
        return jsonify(result), status_code

    except Exception as e:
        logger.error(f"[CLUSTERING_TRAIN] Error fatal: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno al entrenar el modelo de clustering."}), 500



# ==============================================================================
# ENDPOINT: PREDICCIÓN DE CLÚSTERES CON UN DATASET YA GUARDADO EN LA PLATAFORMA
# ==============================================================================
@app.route('/api/models/<string:model_id>/predict-with-saved-dataset', methods=['POST'])
@token_required
def predict_cluster_with_saved_dataset(current_user, model_id):
    """
    Asigna clústeres a un dataset que ya está guardado en Supabase,
    identificado por su dataset_id.
    """
    try:
        # --- 1. OBTENER EL ID DEL DATASET A USAR ---
        # La petición ahora no envía un archivo, sino un JSON con el ID del dataset.
        data = request.get_json()
        if not data or 'dataset_id' not in data:
            return jsonify({
                "success": False, 
                "error": "Debe proporcionar el 'dataset_id' en el cuerpo de la petición."
            }), 400
        
        dataset_id = data['dataset_id']

        # --- 2. DESCARGAR EL MODELO (Lógica que ya tienes) ---
        # Esta parte es idéntica a tu endpoint de predicción original.
        try:
            model_record = supabase_admin.table("trained_models") \
                .select("model_storage_path") \
                .eq("id", model_id).eq("user_id", current_user.id) \
                .single().execute()

            if not model_record.data:
                return jsonify({"success": False, "error": "Modelo no encontrado."}), 404

            model_storage_path = model_record.data.get("model_storage_path")
            model_bytes = supabase_handler.load_file(user_id=current_user.id, path=model_storage_path)
            
            if model_bytes is None:
                return jsonify({"success": False, "error": "No se pudo descargar el modelo."}), 500
        except Exception as db_err:
            logger.error(f"[PREDICT-SAVED] Error al buscar modelo: {db_err}", exc_info=True)
            return jsonify({"success": False, "error": "Error al procesar el modelo."}), 500

        # --- 3. DESCARGAR EL DATASET (Lógica de tu endpoint de descarga) ---
        # ¡Aquí combinamos la lógica! Usamos el dataset_id para ir a buscar el archivo.
        try:
            dataset_record = supabase_admin.table("datasets") \
                .select("storage_path") \
                .eq("dataset_id", dataset_id).eq("user_id", current_user.id) \
                .single().execute()

            if not dataset_record.data:
                return jsonify({"success": False, "error": "Dataset no encontrado."}), 404
            
            dataset_storage_path = dataset_record.data.get("storage_path")
            dataset_bytes = supabase_handler.load_file(user_id=current_user.id, path=dataset_storage_path)

            if dataset_bytes is None:
                return jsonify({"success": False, "error": "No se pudo descargar el archivo del dataset."}), 500
        except Exception as db_err:
            logger.error(f"[PREDICT-SAVED] Error al buscar dataset: {db_err}", exc_info=True)
            return jsonify({"success": False, "error": "Error al procesar el dataset."}), 500

        # --- 4. PREPARAR DATOS Y EJECUTAR PREDICCIÓN ---
        # Convertimos los bytes del archivo descargado en un DataFrame
        from io import BytesIO
        df_new = pd.read_csv(BytesIO(dataset_bytes))
        new_data = df_new.to_dict(orient="records")

        result = tabular_clustering_service.predict_with_saved_model(
            model_bytes=model_bytes,
            new_data=new_data
        )

        if not result.get("success"):
            return jsonify(result), 400

        # --- 5. FORMAR RESPUESTA (Lógica que ya tienes) ---
        predicted_labels = result["data"]["labels"]
        df_new["Segmento_Predicho"] = predicted_labels
        response_data = df_new.to_dict(orient="records")

        return jsonify({
            "success": True,
            "message": "Predicciones generadas correctamente.",
            "data": response_data
        }), 200

    except Exception as e:
        logger.error(f"[PREDICT-SAVED] Error fatal: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno durante la predicción."}), 500


# # ================================================================
# ENDPOINT: PREDICCIÓN DE CLÚSTERES (VERSIÓN FINAL Y ROBUSTA)
# ================================================================
@app.route('/api/models/<string:model_id>/predict-cluster', methods=['POST'])
@token_required
def predict_cluster_endpoint(current_user, model_id):
    """
    Asigna clústeres a un dataset nuevo usando un modelo guardado.
    El dataset se recibe como archivo (form-data).
    """
    try:
        # --- 1. VALIDACIÓN DEL ARCHIVO SUBIDO ---
        if 'file' not in request.files:
            return jsonify({
                "success": False,
                "error": "Debe subir un archivo con la clave 'file'."
            }), 400

        file = request.files['file']
        if file.filename.strip() == "":
            return jsonify({
                "success": False,
                "error": "El archivo subido no tiene nombre válido."
            }), 400

        # --- 2. LEER EL DATAFRAME ---
        try:
            df_new = pd.read_csv(file)  # soporte inicial: CSV
        except Exception as e:
            logger.warning(f"[CLUSTERING][PREDICT] Error leyendo dataset: {e}")
            return jsonify({
                "success": False,
                "error": f"No se pudo leer el archivo como CSV: {str(e)}"
            }), 400

        if df_new.empty:
            return jsonify({
                "success": False,
                "error": "El archivo está vacío o no tiene datos válidos."
            }), 400

        new_data = df_new.to_dict(orient="records")

        # --- 3. CONSULTAR Y DESCARGAR MODELO ---
        try:
            model_record = supabase_admin.table("trained_models") \
                .select("model_storage_path") \
                .eq("id", model_id) \
                .eq("user_id", current_user.id) \
                .single() \
                .execute()

            if not model_record.data:
                return jsonify({
                    "success": False,
                    "error": "Modelo no encontrado o no autorizado."
                }), 404

            model_storage_path = model_record.data.get("model_storage_path")
            if not model_storage_path:
                return jsonify({
                    "success": False,
                    "error": "El modelo no tiene ruta de almacenamiento."
                }), 500

            model_bytes = supabase_handler.load_file(
                user_id=current_user.id,
                path=model_storage_path
            )
            if model_bytes is None:
                return jsonify({
                    "success": False,
                    "error": "No se pudo descargar el modelo desde el storage."
                }), 500

        except Exception as db_err:
            logger.error(f"[CLUSTERING][PREDICT] Error procesando modelo: {db_err}", exc_info=True)
            return jsonify({
                "success": False,
                "error": "Error al buscar o descargar el modelo."
            }), 500

        # --- 4. EJECUTAR PREDICCIÓN ---
        result = tabular_clustering_service.predict_with_saved_model(
            model_bytes=model_bytes,
            new_data=new_data
        )

        if not result.get("success"):
            return jsonify(result), 400

        # --- 5. FORMAR RESPUESTA ---
        predicted_labels = result["data"]["labels"]
        df_new["Segmento_Predicho"] = predicted_labels
        response_data = df_new.to_dict(orient="records")

        return jsonify({
            "success": True,
            "message": "Predicciones generadas correctamente.",
            "data": response_data
        }), 200

    except Exception as e:
        logger.error(f"[CLUSTERING][PREDICT] Error fatal: {e}", exc_info=True)
        return jsonify({
            "success": False,
            "error": "Error interno durante la predicción."
        }), 500

# ================================================================
# ENDPOINT: GUARDAR DATASET ENRIQUECIDO (VERSIÓN FINAL ROBUSTA)
# ================================================================
@app.route('/api/datasets/save-enriched-dataset', methods=['POST'])
@token_required
def save_enriched_dataset_endpoint(current_user):
    logger.info(f"[SAVE_ENRICHED] Usuario={current_user.id} iniciando guardado.")
    temp_path = None  # inicializamos fuera del try para usar en finally

    try:
        # 1. VALIDAR PAYLOAD
        payload = request.get_json(silent=True)
        if not payload:
            return jsonify({"success": False, "error": "El cuerpo de la solicitud no es válido JSON."}), 400

        original_dataset_id = payload.get("original_dataset_id")
        new_column_name = payload.get("new_column_name")
        labels = payload.get("labels")
        new_dataset_name = payload.get("new_dataset_name")

        if not all([original_dataset_id, new_column_name, labels, new_dataset_name]):
            return jsonify({"success": False, "error": "Faltan parámetros en la solicitud."}), 400
        if not isinstance(labels, list):
            return jsonify({"success": False, "error": "El campo 'labels' debe ser una lista."}), 400

        # 2. CARGAR EL DATAFRAME ORIGINAL
        dataset_info = dataset_service.get_dataset_info_by_id(original_dataset_id, current_user.id)
        if not dataset_info.get("success"):
            return jsonify(dataset_info), 404

        storage_path = dataset_info["data"].get("storage_path")
        project_id = dataset_info["data"].get("project_id")
        df = supabase_handler.load_file_as_dataframe(user_id=current_user.id, path=storage_path)

        if df is None or df.empty:
            return jsonify({"success": False, "error": "No se pudo cargar el dataset original."}), 500

        if len(df) != len(labels):
            return jsonify({
                "success": False,
                "error": "La longitud de las etiquetas no coincide con el número de filas del dataset."
            }), 400

        # 3. AÑADIR LA NUEVA COLUMNA
        df[new_column_name] = labels

        # 4. GUARDAR EL NUEVO DATAFRAME COMO ARCHIVO TEMPORAL
        with tempfile.NamedTemporaryFile(mode='w+', delete=False, suffix='.csv', encoding='utf-8') as temp_f:
            df.to_csv(temp_f, index=False)
            temp_path = temp_f.name

        file_size = os.path.getsize(temp_path)

        # 5. SUBIR EL NUEVO ARCHIVO AL STORAGE
        new_dataset_id = str(uuid.uuid4())
        new_storage_path = f"{current_user.id}/{project_id}/datasets/{new_dataset_id}.csv"

        with open(temp_path, 'rb') as f:
            supabase_handler._client.storage.from_("proyectos-usuarios").upload(
                path=new_storage_path,
                file=f,
                file_options={"content-type": "text/csv"}
            )

        # 6. CREAR EL NUEVO REGISTRO EN LA TABLA 'datasets'
        new_dataset_record = {
            "dataset_id": new_dataset_id,
            "user_id": current_user.id,
            "project_id": project_id,
            "dataset_name": new_dataset_name,
            "storage_path": new_storage_path,
            "dataset_type": "tabular",
            "file_size": file_size,
            "status": "processed"
        }
        supabase_handler._client.table("datasets").insert(new_dataset_record).execute()

        # 7. RESPUESTA FINAL
        return jsonify({
            "success": True,
            "message": "Dataset enriquecido guardado con éxito.",
            "new_dataset_id": new_dataset_id
        }), 201

    except Exception as e:
        logger.error(f"[SAVE_ENRICHED] Error fatal: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno al guardar el nuevo dataset."}), 500

    finally:
        # LIMPIEZA DEL ARCHIVO TEMPORAL
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
                logger.info(f"Archivo temporal {temp_path} eliminado.")
            except Exception as cleanup_err:
                logger.warning(f"No se pudo eliminar archivo temporal {temp_path}: {cleanup_err}")

# ================================================================
#   NUEVO ENDPOINT: ASISTENTE DE IA PARA ETIQUETADO DE CLÚSTERES
# ================================================================
@app.route('/api/clustering/suggest-labels', methods=['POST'])
@token_required
def suggest_cluster_labels_endpoint(current_user):
    logger.info(f"[SUGGEST_LABELS] Usuario={current_user.id} solicitando sugerencias.")
    try:
        # 1. VALIDAR INPUT
        # El frontend nos enviará el JSON del `summary` de un clúster
        cluster_summary = request.get_json()
        if not cluster_summary or "numeric_summary" not in cluster_summary:
            return jsonify({"success": False, "error": "Falta el resumen del clúster."}), 400

        # 2. CONSTRUIR EL USER PROMPT (convertir el JSON a texto legible)
        text_content = "He analizado un grupo de datos y estas son sus características principales:\n\n"
        
        # Formatear resumen numérico
        numeric_summary = cluster_summary.get("numeric_summary", {})
        if numeric_summary:
            text_content += "Estadísticas Numéricas (Promedios):\n"
            for col, stats in numeric_summary.items():
                if 'mean' in stats:
                    text_content += f"- {col}: {stats['mean']:.2f}\n"
            text_content += "\n"
            
        # Formatear resumen categórico
        categorical_summary = cluster_summary.get("categorical_summary", {})
        if categorical_summary:
            text_content += "Características Categóricas (Valores más comunes):\n"
            for col, value in categorical_summary.items():
                text_content += f"- {col}: {value}\n"
        
        # 3. SELECCIONAR MODELO Y OBTENER HERRAMIENTA
        # (Reutilizamos la misma lógica de tu endpoint de recetas)
        llm_manager = get_llm_manager()
        model_to_use = "API: Google Gemini Pro" # O el modelo que definas para esta tarea
        tool_response = llm_manager.get_tool(model_to_use)
        if not tool_response.get("success"):
            return jsonify(tool_response), 503

        # 4. OBTENER EL SYSTEM PROMPT DE LA NUEVA RECETA
        system_prompt = get_system_prompt_for_cluster_labeling()

        # 5. EJECUTAR EL LLM
        # (Reutilizamos tu función _execute_llm_generation, ¡esto es lo genial!)
        future = executor.submit(
            _execute_llm_generation, 
            tool_response, 
            system_prompt, 
            text_content.strip()
        )
        ai_response = future.result(timeout=EXECUTION_TIMEOUT_SEC)

        # 6. PROCESAR Y DEVOLVER LA RESPUESTA
        # El LLM devuelve "Etiqueta 1, Etiqueta 2, Etiqueta 3"
        # Lo convertimos en un array de strings para que el frontend lo use fácilmente
        suggested_labels = [label.strip() for label in ai_response.split(',')]

        return jsonify({
            "success": True,
            "data": {
                "suggested_labels": suggested_labels
            }
        }), 200

    except Exception as e:
        logger.error(f"[SUGGEST_LABELS] Error fatal: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno al generar sugerencias."}), 500

@app.route("/api/generate-save-image", methods=["POST"])
@token_required
def handle_generate_and_save_image(current_user):
    """
    Endpoint de producción que genera una imagen, la guarda en el storage
    de forma segura y devuelve una URL firmada para su uso inmediato.
    """
    try:
        # --- 1. Obtener y Validar la Entrada ---
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "No se recibió un cuerpo JSON."}), 400

        prompt = data.get("prompt", "").strip()
        if not prompt:
            return jsonify({"success": False, "error": "El prompt no puede estar vacío."}), 400

        prompt = prompt[:500]  # Limitar longitud

        # --- 2. Llamar a la Lógica de Negocio ---
        logger.info(f"✨ Iniciando generación de imagen para usuario={current_user.id}")
        image_bytes, error = generar_imagen_desde_texto(prompt)

        if error:
            logger.error(f"❌ Error durante la generación de imagen: {error}")
            return jsonify({"success": False, "error": error}), 503
        if not image_bytes:
            return jsonify({"success": False, "error": "La generación no produjo una imagen."}), 500

        # =========================================================
        # --- 3. Validar la imagen y preparar el nombre de archivo ---
        # =========================================================
        try:
            with Image.open(BytesIO(image_bytes)) as img:
                ext = img.format.lower()
        except Exception as e:
            logger.warning(f"⚠️ No se pudo detectar el formato de la imagen: {e}")
            ext = "jpeg"  # Valor por defecto seguro

        unique_id = uuid4().hex
        filename = f"ai_{current_user.id}_{unique_id}.{ext}"

        # --- 4. Guardar en Supabase Storage ---
        storage_path, _ = supabase_handler.save_visual_file(
            file_bytes=image_bytes,
            user_id=current_user.id,
            filename=filename,
        )
        if not storage_path:
            return jsonify({"success": False, "error": "Fallo en la subida a Storage."}), 500

        # --- 5. Crear una URL Firmada Segura ---
        signed_url = None
        try:
            signed_url_response = (
                supabase_handler.get_db_client()
                .storage.from_("visuales")  # nombre de bucket
                .create_signed_url(storage_path, 3600)  # expira en 1 hora
            )

            # Blindaje para distintas versiones de supabase-py
            signed_url = (
                signed_url_response.get("signedURL")
                or signed_url_response.get("signed_url")
            )
            if not signed_url:
                raise Exception(f"Respuesta inválida al generar URL firmada: {signed_url_response}")

        except Exception as e:
            logger.error(f"⚠️ Error al generar URL firmada para {storage_path}: {e}")
            return jsonify({
                "success": False,
                "error": "La imagen se guardó, pero no se pudo generar una URL de acceso."
            }), 500

        # --- 6. Guardar el registro en la Base de Datos ---
        record_data = {
            "user_id": current_user.id,
            "storage_path": storage_path,
            "type": "generada",
            "prompt": prompt,
        }
        try:
            supabase_handler.insert_visual_record(record_data)
        except Exception as e:
            logger.warning(f"⚠️ Falló el guardado del registro en DB: {e}")

        # --- 7. Devolver la respuesta JSON al frontend ---
        logger.info(f"✅ Imagen generada y guardada con éxito en {storage_path}")
        return jsonify({
            "success": True,
            "message": "Imagen generada con éxito.",
            "imageUrl": signed_url
        }), 201

    except Exception as e:
        logger.exception(f"❌ Error inesperado en /api/generate-save-image: {e}")
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500



@app.route('/api/visuals', methods=['GET'])
@token_required
def get_visuals_meme_endpoint(current_user):
    """
    Obtiene la galería paginada de imágenes de un usuario con filtros opcionales
    y URLs firmadas temporales seguras generadas directamente desde el handler.
    """
    # --- Parámetros seguros desde query string (sin cambios) ---
    page = request.args.get("page", 1, type=int)
    per_page = request.args.get("per_page", 12, type=int)
    search_term = request.args.get("q", None, type=str)
    filter_type = request.args.get("type", None, type=str)
    color_filter = request.args.get("color", None, type=str)

    try:
        # --- ¡ESTE ES EL ÚNICO CAMBIO! ---
        # Llamamos a tu nueva y mejorada función que ya hace todo el trabajo.
        visuals, total_count = supabase_handler.get_visuals_by_user_meme(
            user_id=current_user.id,
            page=page,
            per_page=per_page,
            search_term=search_term,
            filter_type=filter_type,
            color_filter=color_filter,
        )

        # --- Respuesta exitosa (sin cambios) ---
        return jsonify({
            "success": True,
            "visuals": visuals,
            "total_count": total_count,
            "page": page,
            "per_page": per_page,
        }), 200

    except Exception as e:
        logger.error(
            f"❌ Error inesperado en get_visuals_endpoint (user_id={current_user.id}): {e}",
            exc_info=True,
        )
        return jsonify({
            "success": False,
            "error": "Error al obtener la galería de imágenes."
        }), 500




@app.route("/api/datasets/from-analysis", methods=["POST"])
@token_required
def handle_create_dataset_from_analysis(current_user):
    """
    Recibe items tabulares desde un análisis (ej: extracción de PDFs),
    los normaliza y crea un nuevo dataset tabular en Supabase.
    """
    try:
        # --- 1. Parsear y validar datos de entrada ---
        data = request.get_json(silent=True)
        if not data:
            return jsonify({"success": False, "error": "El cuerpo de la petición no es JSON válido."}), 400

        project_id = data.get("projectId")
        source_dataset_id = data.get("sourceDatasetId")
        dataset_name = data.get("name")
        items_data = data.get("data")

        if not all([project_id, dataset_name, items_data]):
            return jsonify({
                "success": False,
                "error": "Faltan datos requeridos (projectId, name o data)."
            }), 400

        if not isinstance(items_data, list) or not all(isinstance(x, dict) for x in items_data):
            return jsonify({
                "success": False,
                "error": "El campo 'data' debe ser una lista de objetos/diccionarios."
            }), 400

        if len(items_data) == 0:
            return jsonify({
                "success": False,
                "error": "No hay datos para procesar."
            }), 400

        # --- 2. Convertir a DataFrame y CSV ---
        try:
            df = pd.DataFrame(items_data).fillna("N/A")
            csv_bytes = df.to_csv(index=False).encode("utf-8")
        except Exception as e:
            logger.error(f"Error convirtiendo JSON a CSV: {e}")
            return jsonify({
                "success": False,
                "error": "Los datos extraídos tienen un formato inconsistente."
            }), 400

        # --- 3. Generar identificadores ---
        new_dataset_id = str(uuid.uuid4())
        file_name = f"analysis_dataset_{new_dataset_id[:8]}.csv"

        # --- 4. Guardar en Supabase Storage ---
        try:
            storage_path, _ = supabase_handler.save_file(
                file_bytes=csv_bytes,
                user_id=current_user.id,
                project_id=project_id,
                folder="datasets",
                filename=file_name
            )
            if not storage_path:
                raise Exception("Supabase devolvió un storage_path vacío.")
            logger.info(f"✅ CSV subido correctamente a Supabase en {storage_path}")
        except Exception as e:
            logger.error(f"❌ Error al subir CSV a Supabase Storage: {e}", exc_info=True)
            return jsonify({"success": False, "error": "No se pudo guardar el archivo en el almacenamiento."}), 500

        # --- 5. Crear registro en tabla datasets ---
        try:
            new_dataset = dataset_service.create_dataset_record(
                dataset_id=new_dataset_id,
                user_id=current_user.id,
                project_id=project_id,
                dataset_name=dataset_name,
                dataset_type="tabular",
                storage_path=storage_path,
                file_size=len(csv_bytes),
               
            )
        except Exception as e:
            logger.error(f"❌ Error al insertar dataset en Supabase DB: {e}", exc_info=True)
            return jsonify({"success": False, "error": "No se pudo crear el registro en la base de datos."}), 500

        # --- 6. Respuesta final ---
        return jsonify({
            "success": True,
            "message": "Dataset tabular creado con éxito.",
            "new_dataset": new_dataset
        }), 201

    except Exception as e:
        logger.critical(f"❌ Error inesperado en creación de dataset desde análisis: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Error interno del servidor."}), 500


def validar_targets(targets):
    valid_destinations = {"power_bi", "drive", "email", "social"}

    for t in targets:
        if "destination" not in t or "options" not in t:
            return False, "Cada target debe tener 'destination' y 'options'"

        dest = t["destination"]
        opts = t["options"]

        if dest not in valid_destinations:
            return False, f"Destino '{dest}' no soportado"

        if dest == "power_bi":
            if "datasetId" not in opts or "tableName" not in opts:
                return False, "Power BI requiere 'datasetId' y 'tableName'"

        if dest == "drive":
            if "provider" not in opts or "folderId" not in opts:
                return False, "Drive requiere 'provider' y 'folderId'"

        if dest == "email":
            if "to" not in opts or "subject" not in opts:
                return False, "Email requiere 'to' y 'subject'"

        if dest == "social":
            if "platform" not in opts or "message" not in opts:
                return False, "Social requiere 'platform' y 'message'"

    return True, None



@app.route("/api/v1/export", methods=["POST"])
@token_required
def export_file():
    """
    Endpoint de exportación genérico.
    Valida la entrada, construye un payload y reenvía a n8n para procesar el archivo.
    Devuelve una respuesta inmediata al frontend.
    """
    try:
        # --- 1. Validar multipart ---
        if "file" not in request.files or "exportConfig" not in request.form:
            return jsonify({"error": "Faltan parámetros: file o exportConfig"}), 400

        file = request.files["file"]
        export_config_raw = request.form["exportConfig"]

        # --- 2. Validar archivo ---
        if not file.filename:
            return jsonify({"error": "Archivo vacío"}), 400

        filename = secure_filename(file.filename)

        # --- 3. Validar JSON de configuración ---
        try:
            export_config = json.loads(export_config_raw)
        except json.JSONDecodeError:
            return jsonify({"error": "exportConfig no es un JSON válido"}), 400

        if "targets" not in export_config or not isinstance(export_config["targets"], list):
            return jsonify({"error": "exportConfig debe contener 'targets' como lista"}), 400

        # Validar cada destino
        is_valid, error_msg = validar_targets(export_config["targets"])
        if not is_valid:
            return jsonify({"error": error_msg}), 400

        # --- 4. Construir payload para n8n ---
        payload_n8n = {
            "user": getattr(request, "user", {}).get("user_id"),  # opcional
            "filename": filename,
            "targets": export_config["targets"]
        }

        # --- 5. Forward hacia n8n ---
        # Obtener URL de n8n desde entorno con fallback seguro
        n8n_url = os.environ.get("N8N_WEBHOOK_URL", "http://n8n:5678/webhook-placeholder")
        if "placeholder" in n8n_url:
            # Log para devops, no rompe la ejecución
            print("⚠️  ADVERTENCIA: N8N_WEBHOOK_URL no está configurada. Usando valor por defecto.")

        files = {"file": (filename, file.stream, file.mimetype)}
        data = {"payload": json.dumps(payload_n8n)}

        try:
            response = requests.post(
                n8n_url,
                files=files,
                data=data,
                timeout=5
            )
            # Opcional: loguear si la respuesta no fue 2xx
            if not response.ok:
                print(f"⚠️  n8n devolvió status {response.status_code}: {response.text}")
        except requests.RequestException as e:
            return jsonify({"error": f"No se pudo contactar n8n: {str(e)}"}), 502

        # --- 6. Respuesta inmediata al frontend ---
        return jsonify({"status": "Proceso de exportación iniciado"}), 202

    except Exception as e:
        # Error inesperado
        print(f"❌ Error interno en export_file: {str(e)}")
        return jsonify({"error": f"Error interno: {str(e)}"}), 500




# --- Healthcheck ---
@app.route("/api/v1/health", methods=["GET"])
def health():
    return jsonify({"status": "ok"}), 200


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5001, debug=False)
