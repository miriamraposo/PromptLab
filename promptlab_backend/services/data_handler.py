# data_handler.py
import io
import logging
from typing import Tuple, Optional, Dict, Any
from werkzeug.datastructures import FileStorage
import pandas as pd
import docx
from werkzeug.utils import secure_filename 


try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import chardet
except ImportError:
    chardet = None

# --- Logger profesional ---
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)
handler = logging.StreamHandler()
formatter = logging.Formatter('%(asctime)s [%(levelname)s] %(message)s')
handler.setFormatter(formatter)
logger.addHandler(handler)

# --- Configuración global ---
ALLOWED_EXTENSIONS = {'csv', 'xlsx', 'xls', 'txt', 'pdf', 'docx'}
MAX_PREVIEW_ROWS = 500

# --- Funciones auxiliares ---

ALLOWED_EXTENSIONS = {'csv', 'xlsx', 'xls', 'pdf', 'docx', 'txt'}
MAX_FILE_SIZE_MB = 10
MAX_ROWS = 10000  # Limite filas para csv/xls para evitar saturar memoria


def _is_allowed_extension(filename: str) -> bool:
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def _check_file_size(file_bytes: bytes) -> bool:
    size_mb = len(file_bytes) / (1024 * 1024)
    return size_mb <= MAX_FILE_SIZE_MB


# En data_handler.py

def detect_file_type(filename: str) -> str:
    extension = filename.rsplit('.', 1)[-1].lower()

    if extension in {'csv', 'xls', 'xlsx'}:
        return 'tabular'
    elif extension == 'txt':
        return 'text' # <-- Solo .txt es de tipo 'text'
    elif extension == 'pdf':
        return 'pdf'  # <-- PDF tiene su propio tipo
    elif extension == 'docx':
        return 'docx' # <-- DOCX tiene su propio tipo
    elif extension == 'json':
        return 'json' # <-- JSON tiene su propio tipo
    elif extension == 'md':
        return 'md'   # <-- Markdown podría ser editable también
        
    return 'unknown'


def _extract_text_from_pdf(file_stream: io.BytesIO) -> Tuple[Optional[str], str, Dict[str, Any]]:
    if not PyPDF2:
        return None, "PyPDF2 no instalado.", {}
    try:
        reader = PyPDF2.PdfReader(file_stream)
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
        if not text.strip():
            return None, "PDF vacío o solo imágenes.", {}
        metadata = {
            'page_count': len(reader.pages),
            'char_count': len(text),
            'word_count': len(text.split()),
            'preview': text[:250].strip() + '...'
        }
        return text, "Texto extraído PDF.", metadata
    except Exception as e:
        logger.error(f"Error leyendo PDF: {e}", exc_info=True)
        return None, f"Error leyendo PDF: {e}", {}


def _extract_text_from_docx(file_stream: io.BytesIO) -> Tuple[Optional[str], str, Dict[str, Any]]:
    try:
        document = docx.Document(file_stream)
        text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
        if not text.strip():
            return None, "DOCX vacío.", {}
        metadata = {
            'char_count': len(text),
            'word_count': len(text.split()),
            'preview': text[:250].strip() + '...'
        }
        return text, "Texto extraído DOCX.", metadata
    except Exception as e:
        logger.error(f"Error leyendo DOCX: {e}", exc_info=True)
        return None, f"Error leyendo DOCX: {e}", {}


def _detect_encoding(file_bytes: bytes) -> Optional[str]:
    if not chardet:
        return None
    try:
        result = chardet.detect(file_bytes)
        encoding = result.get('encoding')
        logger.info(f"Encoding detectado: {encoding}")
        return encoding
    except Exception as e:
        logger.warning(f"No se pudo detectar encoding: {e}")
        return None


def _extract_text_from_txt(file_bytes: bytes) -> Tuple[Optional[str], str, Dict[str, Any]]:
    encodings_to_try = []
    encoding_detected = _detect_encoding(file_bytes)
    if encoding_detected:
        encodings_to_try.append(encoding_detected)
    encodings_to_try.extend(['utf-8', 'latin1'])

    for enc in encodings_to_try:
        try:
            text = file_bytes.decode(enc)
            if not text.strip():
                return None, "TXT vacío.", {}
            metadata = {
                'char_count': len(text),
                'word_count': len(text.split()),
                'preview': text[:250].strip() + '...'
            }
            return text, f"Texto extraído TXT con encoding {enc}.", metadata
        except Exception:
            continue
    return None, "No se pudo decodificar TXT con los encodings probados.", {}


def _read_tabular(file_stream: io.BytesIO, extension: str) -> Tuple[Optional[pd.DataFrame], str]:
    try:
        if extension == 'csv':
            df = pd.read_csv(file_stream, nrows=MAX_ROWS, on_bad_lines='warn', low_memory=False)
        else:
            df = pd.read_excel(file_stream, nrows=MAX_ROWS)
        if df.empty or df.shape[1] == 0:
            return None, "Archivo tabular vacío o mal formado."
        return df, "Archivo tabular procesado correctamente."
    except Exception as e:
        logger.error(f"Error leyendo archivo tabular: {e}", exc_info=True)
        return None, f"Error leyendo archivo tabular: {e}"


def process_uploaded_file(file: FileStorage) -> Dict[str, Any]:
    if not file or not file.filename:
        return {"success": False, "error": "No se proporcionó archivo."}
    filename = secure_filename(file.filename).replace('#', '_')

    if not _is_allowed_extension(filename):
        return {"success": False, "error": f"Extensión no permitida: {filename}"}

    file_bytes = file.read()
    if not _check_file_size(file_bytes):
        return {"success": False, "error": f"Archivo excede {MAX_FILE_SIZE_MB} MB."}

    extension = filename.rsplit('.', 1)[-1].lower()
    file_type = detect_file_type(filename)
    file_stream = io.BytesIO(file_bytes)

    try:
        if file_type == 'tabular':
            df, msg = _read_tabular(file_stream, extension)
            if df is None:
                return {"success": False, "error": msg}
            return {
                "success": True,
                "project_type": "tabular",
                "message": msg,
                "filename": filename,
                "columns": list(df.columns),
                "n_rows": df.shape[0],
                "n_columns": df.shape[1],
                "content": df.to_dict(orient="records")
            }
        elif file_type == 'text':
            if extension == 'pdf':
                text, message, metadata = _extract_text_from_pdf(file_stream)
            elif extension == 'docx':
                text, message, metadata = _extract_text_from_docx(file_stream)
            elif extension == 'txt':
                text, message, metadata = _extract_text_from_txt(file_bytes)
            else:
                return {"success": False, "error": "Tipo de archivo de texto no reconocido."}
            if text is None:
                return {"success": False, "error": message}
            return {
                "success": True,
                "project_type": "text",
                "message": message,
                "filename": filename,
                "metadata": metadata,
                "content": text
            }
        else:
            return {"success": False, "error": "Tipo de archivo no soportado."}
    except Exception as e:
        logger.error(f"Error procesando archivo {filename}: {e}", exc_info=True)
        return {"success": False, "error": f"Error inesperado: {e}"}




def is_allowed_file(filename: str) -> bool:
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def read_text(file: FileStorage) -> str:
    return file.read().decode('utf-8', errors='ignore')

def read_csv(file: FileStorage) -> pd.DataFrame:
    return pd.read_csv(file, encoding='utf-8', nrows=MAX_PREVIEW_ROWS)

def read_excel(file: FileStorage) -> pd.DataFrame:
    return pd.read_excel(file, nrows=MAX_PREVIEW_ROWS)

def read_pdf(file: FileStorage) -> str:
    if PyPDF2 is None:
        raise ImportError("PyPDF2 no está instalado.")
    try:
        reader = PyPDF2.PdfReader(file)
        text = ''
        for page in reader.pages:
            text += page.extract_text() or ''
        return text
    except Exception as e:
        logger.error(f"Error al leer PDF: {e}")
        raise

def read_docx(file: FileStorage) -> str:
    try:
        doc = docx.Document(file)
        return '\n'.join([para.text for para in doc.paragraphs])
    except Exception as e:
        logger.error(f"Error al leer DOCX: {e}")
        raise

# --- Función principal ---

def process_file(file: FileStorage) -> Dict[str, Any]:
    """
    Procesa un archivo subido y devuelve una vista previa o texto extraído.

    :param file: Archivo recibido desde un formulario
    :return: Diccionario con tipo de archivo, contenido o errores
    """
    if not file or file.filename == '':
        logger.warning("No se subió ningún archivo.")
        return {'error': 'No se subió ningún archivo.'}

    if not is_allowed_file(file.filename):
        logger.warning(f"Extensión no permitida: {file.filename}")
        return {'error': 'Extensión de archivo no permitida.'}

    filename = file.filename.lower()
    logger.info(f"Procesando archivo: {filename}")

    try:
        if filename.endswith(('.csv', '.txt')):
            content = read_csv(file)
            return {'type': 'table', 'preview': content.to_dict(orient='records')}
        elif filename.endswith(('.xlsx', '.xls')):
            content = read_excel(file)
            return {'type': 'table', 'preview': content.to_dict(orient='records')}
        elif filename.endswith('.pdf'):
            content = read_pdf(file)
            return {'type': 'text', 'content': content}
        elif filename.endswith('.docx'):
            content = read_docx(file)
            return {'type': 'text', 'content': content}
        else:
            text = read_text(file)
            return {'type': 'text', 'content': text}

    except Exception as e:
        logger.exception(f"Error al procesar el archivo: {e}")
        return {'error': f'Error al procesar el archivo: {str(e)}'}

